import base64
import json
import re
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, Body, BackgroundTasks, Request, Response
from sqlalchemy.orm import Session, defer
from sqlalchemy import func, extract
from datetime import datetime, date, timedelta
from typing import List, Optional

from database import get_db
from security import get_teacher, get_current_user
import grading
from models import (
    User, TeacherProfile, ClassEntry, ClassStatus,
    RescheduleRequest, RescheduleStatus, DPP, Test, Doubt,
    DoubtStatus, Timetable, Notification, TestStatus,
    Exam, ExamQuestion, ExamAttempt, ExamResult,
    Material, Lecture, TeacherAttendance, ist_iso
)

def _r2img(v):
    """Exam figure base64 -> R2 URL (naye uploads auto R2 par). Pehle se URL/khaali ho
    to waise ka waisa. R2 na ho to base64 fallback (kabhi na toote)."""
    try:
        return __import__("r2_storage").normalize(v, "exam-q", "image/jpeg") if v else v
    except Exception:
        return v
from security import get_admin
from schemas import (
    ClassEntryCreate, ClassEntryUpdate, ClassEntryOut,
    TimetableCreate, TimetableOut,
    RescheduleCreate, RescheduleOut,
    DPPCreate, DPPOut,
    TestCreate, TestPaperUpload, TestOut,
    DoubtResolve, DoubtOut,
    TeacherDashboard
)

router = APIRouter(prefix="/api/teacher", tags=["Teacher"])


# ---- per-teacher "can see students" access control (app_settings, no migration) ----
STUDENTS_ACCESS_KEY = "teacher_students_ids"


def _students_allowed_ids(db):
    """Set of teacher PROFILE ids jinhe 'My Students' + phone dikhana allowed hai."""
    try:
        from models import AppSetting
        row = db.query(AppSetting).filter(AppSetting.key == STUDENTS_ACCESS_KEY).first()
        if not row or not row.value:
            return set()
        return {int(x) for x in str(row.value).replace(",", " ").split() if x.strip().isdigit()}
    except Exception:
        return set()


def _teacher_sees_students(tp, db):
    """True agar is teacher ko students ki full details (My Students + phone) dikhana allowed hai.
    Default: OFF (admin panel se per-teacher enable hota hai)."""
    try:
        return bool(tp) and tp.id in _students_allowed_ids(db)
    except Exception:
        return False


def get_teacher_profile(user, db):
    profile = db.query(TeacherProfile).filter(TeacherProfile.user_id == user.id).first()
    if not profile:
        raise HTTPException(status_code=404, detail="Teacher profile not found")
    return profile

def notify(db, user_id: int, title: str, message: str, notif_type: str,
           sender_id=None, sender_role=None, batch_key=None, batch_label=None):
    """Helper to create notification (v93: sender/batch fields optional — views tracking)"""
    n = Notification(user_id=user_id, title=title, message=message, notif_type=notif_type,
                     sender_id=sender_id, sender_role=sender_role,
                     batch_key=batch_key, batch_label=batch_label)
    db.add(n)

# ===== DASHBOARD =====
@router.get("/dashboard", response_model=TeacherDashboard)
def teacher_dashboard(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    now = datetime.now()
    month_start = date(now.year, now.month, 1)
    week_start  = date.today() - timedelta(days=date.today().weekday())

    q = db.query(ClassEntry).filter(ClassEntry.teacher_id == tp.id)

    # Classes done = teacher ki submit ki gayi LECTURES (naya system). Pending/rescheduled
    # abhi ClassEntry se (agar use ho); warna 0.
    from models import Lecture as _Lec
    _lq = db.query(_Lec).filter(_Lec.teacher_id == tp.id)
    total_done      = _lq.count()
    total_pending   = q.filter(ClassEntry.status == ClassStatus.pending).count()
    total_rescheduled = q.filter(ClassEntry.status == ClassStatus.rescheduled).count()
    monthly_done    = _lq.filter(_Lec.lecture_date >= month_start).count()
    monthly_pending = q.filter(ClassEntry.status == ClassStatus.pending, ClassEntry.scheduled_date >= month_start).count()
    weekly_done     = _lq.filter(_Lec.lecture_date >= week_start).count()

    # Reset monthly reschedule counter if new month
    if tp.reschedule_reset_month != now.month:
        tp.reschedule_count_this_month = 0
        tp.reschedule_reset_month = now.month
        db.commit()

    total_dpps  = db.query(DPP).filter(DPP.teacher_id == tp.id).count()
    total_tests = db.query(Test).filter(Test.teacher_id == tp.id).count()
    # unresolved = not-resolved PLUS resolved-but-new-follow-up (thread ka last msg student ka)
    from models import DoubtResponse as _DR
    unresolved = 0
    for d in db.query(Doubt).filter(Doubt.teacher_id == tp.id).all():
        is_resolved = (getattr(d.status, "value", str(d.status)) == "resolved")
        if not is_resolved:
            unresolved += 1
        else:
            r = (db.query(_DR).filter(_DR.doubt_id == d.id)
                 .order_by(_DR.created_at.desc(), _DR.id.desc()).first())
            if r and r.role == "student":
                unresolved += 1

    return TeacherDashboard(
        total_done=total_done, total_pending=total_pending,
        total_rescheduled=total_rescheduled, monthly_done=monthly_done,
        monthly_pending=monthly_pending, weekly_done=weekly_done,
        reschedule_this_month=tp.reschedule_count_this_month,
        total_dpps=total_dpps, total_tests=total_tests,
        unresolved_doubts=unresolved
    )

# ===== TIMETABLE =====
@router.post("/timetable", response_model=TimetableOut)
def add_timetable(req: TimetableCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    entry = Timetable(teacher_id=tp.id, **req.model_dump())
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return entry

@router.get("/timetable", response_model=List[TimetableOut])
def get_timetable(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(Timetable).filter(Timetable.teacher_id == tp.id, Timetable.is_active == True).all()

# ===== CLASSES =====
@router.post("/classes", response_model=ClassEntryOut)
def create_class(req: ClassEntryCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    entry = ClassEntry(teacher_id=tp.id, **req.model_dump())
    db.add(entry)
    db.commit()
    db.refresh(entry)
    return entry

@router.get("/classes", response_model=List[ClassEntryOut])
def get_classes(
    status: Optional[str] = None,
    subject: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    q = db.query(ClassEntry).filter(ClassEntry.teacher_id == tp.id)
    if status:
        q = q.filter(ClassEntry.status == status)
    if subject:
        q = q.filter(ClassEntry.subject == subject)
    return q.order_by(ClassEntry.scheduled_date, ClassEntry.scheduled_time).all()

@router.get("/classes/today", response_model=List[ClassEntryOut])
def get_today_classes(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(ClassEntry).filter(
        ClassEntry.teacher_id == tp.id,
        ClassEntry.scheduled_date == date.today()
    ).order_by(ClassEntry.scheduled_time).all()

@router.patch("/classes/{class_id}/upload", response_model=ClassEntryOut)
def upload_class_pdf(
    class_id: int,
    drive_link: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    """Teacher uploads PDF → status auto-Done"""
    tp = get_teacher_profile(current_user, db)
    entry = db.query(ClassEntry).filter(
        ClassEntry.id == class_id,
        ClassEntry.teacher_id == tp.id
    ).first()
    if not entry:
        raise HTTPException(status_code=404, detail="Class not found")
    entry.drive_link = drive_link
    entry.status = ClassStatus.done
    db.commit()
    db.refresh(entry)
    return entry

# ===== RESCHEDULE =====
@router.post("/reschedule", response_model=RescheduleOut)
def request_reschedule(req: RescheduleCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    now = datetime.now()

    # Reset if new month
    if tp.reschedule_reset_month != now.month:
        tp.reschedule_count_this_month = 0
        tp.reschedule_reset_month = now.month

    # NOTE: Monthly reschedule limit hata di gayi hai. Teacher jitni baar chahe
    # reschedule request bhej sakta hai — har request admin approval par hi
    # apply hoti hai. Count sirf tracking/reporting ke liye rakha gaya hai.

    class_entry = db.query(ClassEntry).filter(
        ClassEntry.id == req.class_entry_id,
        ClassEntry.teacher_id == tp.id
    ).first()
    if not class_entry:
        raise HTTPException(status_code=404, detail="Class not found")

    # Check existing pending request
    existing = db.query(RescheduleRequest).filter(
        RescheduleRequest.class_entry_id == req.class_entry_id,
        RescheduleRequest.status == RescheduleStatus.pending
    ).first()
    if existing:
        raise HTTPException(status_code=400, detail="A request for this class is already pending")

    # Mark class as rescheduled (pending admin approval)
    class_entry.status = ClassStatus.rescheduled

    rs = RescheduleRequest(
        class_entry_id=req.class_entry_id,
        teacher_id=tp.id,
        original_date=class_entry.scheduled_date,
        original_time=class_entry.scheduled_time,
        new_date=req.new_date,
        new_time=req.new_time,
        reason=req.reason,
        status=RescheduleStatus.pending
    )
    db.add(rs)

    # Notify all admins
    admins = db.query(User).filter(User.role == "admin").all()
    for admin in admins:
        notify(db, admin.id,
               f"Reschedule Request — {current_user.name}",
               f"{class_entry.subject} ({class_entry.class_name}) ko {req.new_date} pe reschedule karna chahte hain. Reason: {req.reason}",
               "reschedule_request")

    db.commit()
    db.refresh(rs)
    return rs

@router.get("/reschedule", response_model=List[RescheduleOut])
def get_my_reschedules(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(RescheduleRequest).filter(RescheduleRequest.teacher_id == tp.id).all()

# ===== DPP =====
@router.post("/dpp", response_model=DPPOut)
def upload_dpp(req: DPPCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    dpp = DPP(teacher_id=tp.id, **req.model_dump())
    db.add(dpp)
    db.commit()
    db.refresh(dpp)
    # Notify students of this subject — har change ka update students tak pahunchna chahiye
    try:
        from models import StudentProfile
        for sp in db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all():
            if sp.subjects and (dpp.subject or "").strip() in sp.subjects and sp.user:
                notify(db, sp.user.id,
                       f"📝 New DPP: {(dpp.subject or '').strip()}",
                       f"{current_user.name} ne {(dpp.subject or '').strip()} ({dpp.reference or 'General'}) ka naya DPP diya hai. DPP section mein dekho!",
                       "new_dpp")
        db.commit()
    except Exception:
        db.rollback()
    return dpp

@router.get("/dpp", response_model=List[DPPOut])
def get_dpps(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(DPP).filter(DPP.teacher_id == tp.id).all()

# ===== TESTS =====
@router.post("/tests", response_model=TestOut)
def create_test(req: TestCreate, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    test = Test(teacher_id=tp.id, **req.model_dump())
    db.add(test)
    db.commit()
    db.refresh(test)
    # Notify students of this subject about the new test
    try:
        from models import StudentProfile
        when = f"{test.test_date} {test.test_time or ''}".strip()
        for sp in db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all():
            if sp.subjects and (test.subject or "").strip() in sp.subjects and sp.user:
                notify(db, sp.user.id,
                       f"🧪 New Test Scheduled: {(test.subject or '').strip()}",
                       f"{current_user.name} ne {(test.subject or '').strip()} ka test schedule kiya hai — {when}. Tests section mein dekho!",
                       "test_reminder")
        db.commit()
    except Exception:
        db.rollback()
    return test

@router.patch("/tests/{test_id}/upload-paper")
def upload_question_paper(
    test_id: int,
    drive_link: str,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    """Upload question paper — must be 15 min before test"""
    tp = get_teacher_profile(current_user, db)
    test = db.query(Test).filter(Test.id == test_id, Test.teacher_id == tp.id).first()
    if not test:
        raise HTTPException(status_code=404, detail="Test not found")
    test.question_paper_link = drive_link
    test.status = TestStatus.active
    db.commit()
    return {"message": "Question paper uploaded! Students now have access.", "drive_link": drive_link}

@router.get("/tests", response_model=List[TestOut])
def get_tests(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return db.query(Test).filter(Test.teacher_id == tp.id).all()

# ===== DOUBTS =====
def _doubt_resp_json(db, did, my_role, my_teacher_id=None):
    """v93: doubt thread responses (oldest first) — mine flag viewer ke hisaab se."""
    from models import DoubtResponse
    out = []
    for r in (db.query(DoubtResponse).filter(DoubtResponse.doubt_id == did)
              .order_by(DoubtResponse.created_at.asc(), DoubtResponse.id.asc()).all()):
        mine = (r.role == my_role) and (my_role != "teacher" or (my_teacher_id is not None and r.author_teacher_id == my_teacher_id))
        out.append({"id": r.id, "role": r.role, "author_name": r.author_name,
                    "body": r.body, "mine": bool(mine),
                    "author_tid": (r.author_teacher_id if r.role == "teacher" else None),
                    "created_at": ist_iso(r.created_at)})
    return out

def _doubt_needs_attention(db, did):
    """FOLLOW-UP SYSTEM HATA DIYA: resolved doubt ab kabhi reopen nahi hota. Student
    dobara poochhna chahe to naya doubt banata hai. (Pehle student ka last message hone
    par 'New Follow-up' aata tha — ab nahi.)"""
    return False

def _doubt_owner_name(db, d):
    """v93: ab ye doubt kiski responsibility hai — uska display naam."""
    if getattr(d, "assigned_to_admin", False):
        return "MVS Foundation"
    tp = db.query(TeacherProfile).filter(TeacherProfile.id == d.teacher_id).first() if d.teacher_id else None
    return (tp.user.name if tp and tp.user else "Unassigned")

def _norm_subj(s):
    return (s or "").strip().lower()


def _teacher_teaches(tp, subject):
    subs = getattr(tp, "subjects", None) or []
    if isinstance(subs, str):
        subs = [p for chunk in subs.split(",") for p in chunk.split("|")]
    return _norm_subj(subject) in {_norm_subj(x) for x in subs}


@router.get("/doubts")
def get_doubts(
    status: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    # Unassigned doubts (creation ke waqt koi teacher match nahi hua) — agar is
    # teacher ke subject ke hain to claim kar lo taaki panel pe dikhein aur
    # answer/respond/resolve (jo teacher_id se filter karte hain) chal sakein.
    _claimed = False
    for d in db.query(Doubt).filter(Doubt.teacher_id == None,
                                    Doubt.assigned_to_admin == False).all():
        if _teacher_teaches(tp, d.subject):
            d.teacher_id = tp.id
            _claimed = True
    if _claimed:
        db.commit()
    own = db.query(Doubt).filter(Doubt.teacher_id == tp.id).all()
    away = (db.query(Doubt).filter(Doubt.assigned_by_teacher_id == tp.id, Doubt.teacher_id != tp.id).all())
    rows, seen = [], set()
    for d in own + away:
        if d.id in seen:
            continue
        seen.add(d.id)
        rows.append(d)
    if status:
        rows = [d for d in rows if (d.status.value if hasattr(d.status, "value") else d.status) == status]
    rows.sort(key=lambda d: (d.created_at or datetime.now()), reverse=True)
    out = []
    for d in rows:
        sname = d.student.user.name if d.student and d.student.user else "Student"
        is_away = (d.teacher_id != tp.id) or bool(getattr(d, "assigned_to_admin", False))
        out.append({"id": d.id, "student_name": sname, "student_id": d.student_id,
                    "subject": _subj_canon(d.subject), "topic": d.topic,
                    "question": d.question, "has_image": bool(d.image_b64),
                    "attach_mime": d.attach_mime, "attach_name": d.attach_name,
                    "has_voice": bool(d.audio_b64), "has_answer_voice": bool(d.answer_audio_b64),
                    "has_answer_file": bool(d.answer_attach_b64), "answer_attach_mime": d.answer_attach_mime,
                    "answer": d.answer, "status": d.status.value if hasattr(d.status, "value") else d.status,
                    "created_at": ist_iso(d.created_at),
                    "assigned_away": is_away,
                    "assigned_to_name": (_doubt_owner_name(db, d) if is_away else None),
                    "needs_attention": _doubt_needs_attention(db, d.id),
                    "responses": _doubt_resp_json(db, d.id, "teacher", tp.id)})
    return out

@router.post("/doubts/{doubt_id}/respond")
def teacher_doubt_respond(doubt_id: int, payload: dict, db: Session = Depends(get_db),
                          current_user=Depends(get_teacher)):
    """v93: teacher thread pe follow-up likhe (status change nahi hota)."""
    from models import DoubtResponse
    tp = get_teacher_profile(current_user, db)
    d = db.query(Doubt).filter(Doubt.id == doubt_id, Doubt.teacher_id == tp.id,
                               Doubt.assigned_to_admin == False).first()
    if not d:
        raise HTTPException(status_code=404, detail="Doubt not found")
    body = (payload.get("body") or "").strip()
    if not body:
        raise HTTPException(status_code=400, detail="Response text is required")
    db.add(DoubtResponse(doubt_id=d.id, role="teacher", author_name=current_user.name,
                         author_teacher_id=tp.id, body=body))
    if d.student and d.student.user:
        notify(db, d.student.user.id, "💬 New Reply on Your Doubt",
               f"{current_user.name} added a reply on your {d.subject or ''} doubt: {body[:120]}", "doubt")
    db.commit()
    return {"message": "Reply added", "responses": _doubt_resp_json(db, d.id, "teacher", tp.id)}


@router.get("/student-doubts")
def teacher_student_doubts(student_id: int, db: Session = Depends(get_db),
                           current_user=Depends(get_teacher)):
    """Ek student ke SAARE doubts (history) — teacher name pe click kare to dikhe.
    Kuch na ho to khaali list. Naye pehle."""
    from models import StudentProfile
    tp = get_teacher_profile(current_user, db)
    sp = db.query(StudentProfile).filter(StudentProfile.id == student_id).first()
    name = (sp.user.name if sp and sp.user else "Student")
    ds = (db.query(Doubt).filter(Doubt.student_id == student_id)
          .order_by(Doubt.created_at.desc()).all())
    out = []
    for d in ds:
        st = d.status.value if hasattr(d.status, "value") else d.status
        owner = _doubt_owner_name(db, d)
        out.append({"id": d.id, "subject": d.subject, "topic": d.topic,
                    "question": d.question or "", "answer": d.answer or "",
                    "status": st, "owner": owner,
                    "mine": (d.teacher_id == tp.id),
                    "created_at": str(d.created_at)[:16] if d.created_at else "",
                    "has_answer": bool(d.answer)})
    return {"student_id": student_id, "name": name, "total": len(out), "doubts": out}


@router.get("/doubts-assign-targets")
def doubt_assign_targets(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """v93: assign dropdown — saare active teachers (khud ko chhod kar) + Admin option."""
    from models import User
    tp = get_teacher_profile(current_user, db)
    rows = (db.query(TeacherProfile).join(User, TeacherProfile.user_id == User.id)
            .filter(User.is_active == True, TeacherProfile.id != tp.id)
            .order_by(User.name.asc()).all())
    teachers = [{"id": t.id, "name": (t.user.name if t.user else "Teacher"),
                 "subjects": (t.subjects or [])} for t in rows]
    return {"teachers": teachers, "admin": True}

@router.post("/doubts/{doubt_id}/assign")
def teacher_doubt_assign(doubt_id: int, payload: dict, db: Session = Depends(get_db),
                         current_user=Depends(get_teacher)):
    """v93: doubt ko doosre teacher ya admin ko assign karo.
    Assign karne wale ki taraf se doubt resolved maana jata hai —
    ab naye owner ki responsibility hai."""
    from models import DoubtResponse
    tp = get_teacher_profile(current_user, db)
    d = db.query(Doubt).filter(Doubt.id == doubt_id, Doubt.teacher_id == tp.id).first()
    if not d:
        raise HTTPException(status_code=404, detail="Doubt not found")
    if (d.status.value if hasattr(d.status, "value") else d.status) == "resolved":
        raise HTTPException(status_code=400, detail="This doubt is already resolved")
    target = payload.get("target")
    to_admin = (str(target).lower() == "admin")
    new_tp = None
    if not to_admin:
        try:
            new_id = int(target)
        except (TypeError, ValueError):
            raise HTTPException(status_code=400, detail="Invalid assign target")
        new_tp = db.query(TeacherProfile).filter(TeacherProfile.id == new_id).first()
        if not new_tp or new_tp.id == tp.id:
            raise HTTPException(status_code=400, detail="Invalid teacher selected")
    target_name = "MVS Foundation (Admin)" if to_admin else (new_tp.user.name if new_tp.user else "Teacher")
    d.assigned_by_teacher_id = tp.id
    d.assigned_by_name = current_user.name
    d.assigned_at = datetime.now()
    if to_admin:
        d.assigned_to_admin = True
    else:
        d.teacher_id = new_tp.id
        d.assigned_to_admin = False
    db.add(DoubtResponse(doubt_id=d.id, role="teacher", author_name=current_user.name,
                         author_teacher_id=tp.id,
                         body=f"Reassigned this doubt to {target_name}. They will take it forward from here."))
    # student ko batao ki unka doubt ab kiske paas hai
    if d.student and d.student.user:
        notify(db, d.student.user.id, "🔀 Your Doubt Has Been Reassigned",
               f"Your {d.subject or ''} doubt is now with {target_name} — you will get the answer from them.", "doubt")
    # naye owner ko notify
    if to_admin:
        from models import User
        for au in db.query(User).filter(User.is_active == True, User.role == "admin").all():
            notify(db, au.id, "📥 Doubt Assigned to Admin",
                   f"{current_user.name} assigned a {d.subject or ''} doubt by "
                   f"{d.student.user.name if d.student and d.student.user else 'a student'} to MVS Foundation. "
                   f"Please reply from the Doubts page.", "doubt")
    elif new_tp.user:
        notify(db, new_tp.user.id, "📥 Doubt Reassigned to You",
               f"{current_user.name} assigned a {d.subject or ''} doubt by "
               f"{d.student.user.name if d.student and d.student.user else 'a student'} to you. "
               f"It is now your responsibility to resolve it.", "new_doubt")
    db.commit()
    return {"message": f"Doubt assigned to {target_name}", "assigned_to": target_name}

def _t_doubt_media(b64, mime, name):
    return __import__("r2_storage").file_response(b64, mime or "application/octet-stream", (name or "file").replace(chr(34), ""), False)

def _t_own_doubt(did, db, current_user):
    tp = get_teacher_profile(current_user, db)
    d = db.query(Doubt).filter(Doubt.id == did, Doubt.teacher_id == tp.id).first()
    if not d:
        raise HTTPException(status_code=404, detail="Doubt not found")
    return d

@router.get("/doubt/{did}/image")
def teacher_doubt_image(did: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    d = _t_own_doubt(did, db, current_user)
    return _t_doubt_media(d.image_b64, d.attach_mime or "image/jpeg", d.attach_name)

@router.get("/doubt/{did}/voice")
def teacher_doubt_voice(did: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    d = _t_own_doubt(did, db, current_user)
    return _t_doubt_media(d.audio_b64, "audio/webm", "voice.webm")

@router.get("/doubt/{did}/answer-voice")
def teacher_doubt_answer_voice(did: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    d = _t_own_doubt(did, db, current_user)
    return _t_doubt_media(d.answer_audio_b64, "audio/webm", "answer.webm")

@router.get("/doubt/{did}/answer-file")
def teacher_doubt_answer_file(did: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    d = _t_own_doubt(did, db, current_user)
    return _t_doubt_media(d.answer_attach_b64, d.answer_attach_mime, d.answer_attach_name)


@router.delete("/doubt/{did}/answer")
def teacher_delete_doubt_answer(did: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher apna diya hua response delete kar sakta hai — doubt wapas
    pending ho jata hai (student ko fir se jawab mil sakta hai)."""
    d = _t_own_doubt(did, db, current_user)
    d.answer = None
    d.answer_image_link = None
    d.answer_audio_b64 = None
    d.answer_attach_b64 = None
    d.answer_attach_mime = None
    d.status = DoubtStatus.pending
    d.resolved_at = None
    db.commit()
    return {"message": "Response deleted — doubt wapas pending hai"}

@router.patch("/doubts/{doubt_id}/resolve")
def resolve_doubt(
    doubt_id: int,
    req: DoubtResolve,
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    doubt = db.query(Doubt).filter(Doubt.id == doubt_id, Doubt.teacher_id == tp.id).first()
    if not doubt:
        raise HTTPException(status_code=404, detail="Doubt not found")
    if getattr(doubt, "assigned_to_admin", False):
        raise HTTPException(status_code=400, detail="This doubt is with MVS Foundation now")
    doubt.answer = req.answer
    doubt.answer_image_link = req.answer_image_link
    if req.answer_audio_b64:
        doubt.answer_audio_b64 = __import__("r2_storage").normalize(req.answer_audio_b64, "doubt-audio", "audio/webm")
    if req.answer_attach_b64:
        doubt.answer_attach_b64 = __import__("r2_storage").normalize(req.answer_attach_b64, "doubt-attach", (getattr(req, "answer_attach_mime", None) or "application/octet-stream"))
        doubt.answer_attach_mime = req.answer_attach_mime or "application/octet-stream"
        doubt.answer_attach_name = (req.answer_attach_name or "attachment")[:250]
    doubt.status = DoubtStatus.resolved
    doubt.resolved_at = datetime.now()

    # Notify student
    student_user = db.query(User).filter(User.id == doubt.student.user_id).first()
    if student_user:
        notify(db, student_user.id,
               "Your Doubt Is Resolved! ✅",
               f"{doubt.subject} — {doubt.topic}: {current_user.name} has replied.",
               "doubt_resolved")
    db.commit()
    return {"message": "Doubt resolved! The student has been notified."}

# ===== NOTIFICATIONS =====
@router.get("/notifications")
def get_notifications(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    notifs = db.query(Notification).filter(
        Notification.user_id == current_user.id
    ).order_by(Notification.created_at.desc()).limit(20).all()
    return notifs

@router.patch("/notifications/read-all")
def mark_all_read(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Panel/bell view karte hi sab notifications ek saath read — badge/blink clear."""
    db.query(Notification).filter(Notification.user_id == current_user.id,
                                  Notification.is_read == False).update(
        {"is_read": True, "read_at": datetime.now()}, synchronize_session=False)
    db.commit()
    return {"ok": True}


@router.patch("/notifications/{notif_id}/read")
def mark_read(notif_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    n = db.query(Notification).filter(Notification.id == notif_id, Notification.user_id == current_user.id).first()
    if n:
        n.is_read = True
        if not n.read_at:
            n.read_at = datetime.now()
        db.commit()
    return {"message": "Marked as read"}

# ===== TIMETABLE ENTRIES (chapter + parts + date + day) =====
try:
    import subjects_registry as _SR
except Exception:
    _SR = None

def _subj_norm(s):
    """Subject matching key — NIOS canonical registry se (case/alias/code/prefix
    sab handle). 'PHYSICS'=='Physics', 'SCIENCE'=='Science and Technology',
    'Data Entry Op (229)'=='Data Entry Operations'."""
    if _SR is not None:
        try:
            return _SR.canon_norm(s)
        except Exception:
            pass
    import re as _re
    s = str(s or "")
    s = _re.sub(r"\s*[\(\[][^\)\]]*\d+[^\)\]]*[\)\]]\s*$", "", s)
    s = _re.sub(r"\s*[-–—_/]\s*\d{2,}\s*$", "", s)
    s = _re.sub(r"\s+", " ", s).strip().lower()
    s = _re.sub(r"\bop\b", "operations", s)
    return re.sub(r"[^a-z0-9]", "", s)

def _subj_canon(s):
    """Display ke liye official naam (registry se; fallback cleaned input)."""
    if _SR is not None:
        try:
            return _SR.canon_display(s)
        except Exception:
            pass
    return s

def _subj_eq(a, b):
    na, nb = _subj_norm(a), _subj_norm(b)
    return bool(na) and na == nb


def _tt_entry_key(subject, cls):
    """Class-AWARE identity key (canon code+class). Isse Economics 10 aur
    Economics 12 kabhi ek nahi maane jaate."""
    if _SR is not None:
        try:
            return _SR.canon_key(subject, cls)
        except Exception:
            pass
    import re as _re
    c = _re.sub(r"[^0-9]", "", str(cls or ""))
    return _subj_norm(subject) + "|" + c


def _teacher_classkeys(tp):
    """(class-aware keys ka set, subject_classes hai ya nahi). subject_classes se
    banta hai (subject+class dono). Na ho to legacy — caller class filter skip kare."""
    keys = set()
    for sc in (getattr(tp, "subject_classes", None) or []):
        try:
            nm = (sc.get("subject") or "").strip()
            cl = str(sc.get("class") or "").strip()
        except Exception:
            continue
        if nm:
            keys.add(_tt_entry_key(nm, cl))
    return keys, bool(keys)

def _subj_scope_for(db, model, subjects):
    """Assigned subjects + DB me stored unke variants (e.g. 'X (229)') ka map:
    {stored_subject: canonical_assigned_subject}. Exact naam ho to wahi."""
    scope, norm = {}, {}
    for s in subjects or []:
        n = _subj_norm(s)
        if n:
            norm[n] = _subj_canon(s)
            scope[s] = _subj_canon(s)
    if not norm:
        return scope
    try:
        for (v,) in db.query(model.subject).distinct().all():
            if v and v not in scope and _subj_norm(v) in norm:
                scope[v] = norm[_subj_norm(v)]
    except Exception:
        pass
    return scope

def _subj_class_digits(db, subject):
    """Subject ka official class '10'/'12' — admin ke AvailableSubject se pehle,
    phir NIOS registry fallback. Subject DONO classes me ho (English/Hindi/DEO)
    ya pata na chale -> None (tab caller jo class di hai use rakhe).
    v123: isi se 'Social Science (Class 12)' jaisi phantom splitting rukti hai —
    Social Science 10th-only subject hai, uska material kabhi 12 tag nahi hoga."""
    norm = _subj_norm(subject)
    if not norm:
        return None
    classes = set()
    try:
        from models import AvailableSubject
        for r in db.query(AvailableSubject).filter(AvailableSubject.is_active == True).all():
            if _subj_norm(r.name) == norm:
                c = str(r.class_level or "").strip()
                if c in ("10", "12"):
                    classes.add(c)
    except Exception:
        pass
    if len(classes) == 1:
        return next(iter(classes))
    if len(classes) > 1:
        return None                      # admin ne dono classes me rakha hai — dual-class
    if _SR is not None:
        try:
            in10 = _SR.canon_subject(subject, "10")
            in12 = _SR.canon_subject(subject, "12")
            if in10 and not in12:
                return "10"
            if in12 and not in10:
                return "12"
        except Exception:
            pass
    return None


def _serialize_tt(e, canon=None):
    return {
        "id": e.id, "subject": canon or e.subject, "class_name": e.class_name,
        "chapter": e.chapter, "part": e.part,
        "date": str(e.entry_date) if e.entry_date else None, "day": e.day,
        "time": getattr(e, "time_text", None), "type": getattr(e, "entry_type", None) or "chapter", "status": getattr(e, "status", None) or "approved",
        "completed": bool(getattr(e, "completed", False)),
        "topic_covered": getattr(e, "topic_covered", None),
        "homework": getattr(e, "homework", None),
        "dpp_given": bool(getattr(e, "dpp_given", False)),
        "remarks": getattr(e, "remarks", None),
        "start_time": getattr(e, "start_time", None),
        "end_time": getattr(e, "end_time", None),
        "resched_by": getattr(e, "resched_by", None),
        "completed_at": str(getattr(e, "completed_at", "")) if getattr(e, "completed_at", None) else None
    }

@router.post("/timetable-entry")
def add_tt_entry(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    edate = None
    d = (payload.get("entry_date") or "").strip()
    if d:
        try:
            edate = datetime.strptime(d, "%Y-%m-%d").date()
        except Exception:
            edate = None
    e = TimetableEntry(
        teacher_id=tp.id,
        subject=(payload.get("subject") or "").strip(),
        class_name=(payload.get("class_name") or "").strip(),
        chapter=(payload.get("chapter") or "").strip(),
        part=(payload.get("part") or "").strip() or None,
        entry_date=edate,
        day=(payload.get("day") or "").strip() or None,
        time_text=(payload.get("time") or "").strip() or None,
        entry_type=(payload.get("type") or "chapter").strip()
    )
    db.add(e); db.commit(); db.refresh(e)
    return _serialize_tt(e)

@router.get("/timetable-entries")
def list_tt_entries(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    es = db.query(TimetableEntry).filter(TimetableEntry.teacher_id == tp.id).order_by(
        TimetableEntry.subject, TimetableEntry.chapter, TimetableEntry.entry_date
    ).all()
    return [_serialize_tt(e) for e in es]

@router.delete("/timetable-entry/{entry_id}")
def delete_tt_entry(entry_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher apni timeline me dikhne wali koi bhi class delete kar sakta hai —
    apni uploaded ya admin-uploaded (subject match), same scope jaise my-timetable."""
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    e = db.query(TimetableEntry).filter(TimetableEntry.id == entry_id).first()
    if not e or (e.subject not in (tp.subjects or []) and e.teacher_id != tp.id):
        raise HTTPException(status_code=404, detail="Entry not found")
    db.delete(e); db.commit()
    return {"message": "Class deleted"}

@router.delete("/timetable-entries/all")
def clear_tt_entries(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    db.query(TimetableEntry).filter(TimetableEntry.teacher_id == tp.id).delete()
    db.commit()
    return {"message": "All entries cleared"}

# ===== PDF TIMETABLE UPLOAD (auto-parse) =====
@router.post("/timetable-pdf")
async def upload_timetable_pdf(
    file: UploadFile = File(...),
    class_name: str = Form("Class 12"),
    subject: str = Form(""),
    replace: str = Form("false"),
    preview: str = Form("false"),
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    import tt_parser
    raw = await file.read()
    try:
        rows = tt_parser.parse_pdf(raw, force_subject=(subject.strip() or None))
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"PDF parse error: {e}")
    if not rows:
        raise HTTPException(status_code=400, detail="No valid row found in the PDF. A text-based PDF is required.")

    subjects_found = sorted(set(r["subject"] for r in rows))
    # preview mode: sirf parsed rows dikhao, DB me kuch save mat karo
    if preview.lower() == "true":
        try:
            import syllabus_routes
            rows = syllabus_routes.annotate_timetable_rows(db, class_name, rows)
        except Exception:
            pass
        return {"added": 0, "subjects": subjects_found, "preview": rows}
    # replace sirf SAME CLASS ki entries hatao — dusri class ka same-name subject alag timetable hai
    if replace.lower() == "true":
        db.query(TimetableEntry).filter(
            TimetableEntry.teacher_id == tp.id,
            TimetableEntry.subject.in_(subjects_found),
            TimetableEntry.class_name == class_name
        ).delete(synchronize_session=False)

    added = 0
    for r in rows:
        edate = None
        try:
            edate = datetime.strptime(r["date"], "%Y-%m-%d").date()
        except Exception:
            pass
        db.add(TimetableEntry(
            teacher_id=tp.id, subject=r["subject"], class_name=class_name,
            chapter=r["chapter"], part=r["part"], entry_date=edate,
            day=r["day"] or None, time_text=r["time"] or None, entry_type=r["type"]
        ))
        added += 1
    db.commit()
    return {"added": added, "subjects": subjects_found}

# ===== TEACHER: TIMETABLE PDF COMMIT (preview edit ke baad final save) =====
@router.post("/timetable-pdf-commit")
def teacher_timetable_pdf_commit(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    rows = payload.get("rows") or []
    class_name = (payload.get("class_name") or "Class 12").strip()
    replace = str(payload.get("replace") or "false")
    clean = []
    for r in rows:
        sub = (r.get("subject") or "").strip()
        ch = (r.get("chapter") or "").strip()
        if not sub or not ch:
            continue
        clean.append({"subject": sub, "chapter": ch, "part": (r.get("part") or "").strip() or None,
                      "date": r.get("date") or "", "day": (r.get("day") or "").strip(),
                      "time": (r.get("time") or "").strip(), "type": r.get("type") or "chapter"})
    if not clean:
        raise HTTPException(status_code=400, detail="No valid rows left — keep at least 1 chapter.")
    subjects_found = sorted(set(r["subject"] for r in clean))
    if replace.lower() == "true":
        db.query(TimetableEntry).filter(
            TimetableEntry.teacher_id == tp.id,
            TimetableEntry.subject.in_(subjects_found),
            TimetableEntry.class_name == class_name
        ).delete(synchronize_session=False)
    added = 0
    for r in clean:
        edate = None
        try: edate = datetime.strptime(r["date"], "%Y-%m-%d").date()
        except Exception: pass
        db.add(TimetableEntry(
            teacher_id=tp.id, subject=r["subject"], class_name=class_name,
            chapter=r["chapter"], part=r["part"], entry_date=edate,
            day=r["day"] or None, time_text=r["time"] or None, entry_type=r["type"]
        ))
        added += 1
    db.commit()
    return {"added": added, "subjects": subjects_found}

# ===== TEACHER: EDIT TIMETABLE ENTRY TOPIC/PART =====
@router.patch("/timetable-entry/{entry_id}")
def edit_tt_entry(entry_id: int, payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    e = db.query(TimetableEntry).filter(TimetableEntry.id == entry_id).first()
    if not e or (e.subject not in (tp.subjects or []) and e.teacher_id != tp.id):
        raise HTTPException(status_code=404, detail="Entry not found")
    if "part" in payload:
        e.part = (payload.get("part") or "").strip() or None
    if "time" in payload:
        e.time_text = (payload.get("time") or "").strip() or None
    if "chapter" in payload and (payload.get("chapter") or "").strip():
        e.chapter = (payload.get("chapter") or "").strip()
    if "type" in payload and (payload.get("type") or "").strip() in ("chapter", "event"):
        e.entry_type = (payload.get("type") or "").strip()
    if "entry_date" in payload:
        d = (payload.get("entry_date") or "").strip()
        if d:
            try:
                e.entry_date = datetime.strptime(d, "%Y-%m-%d").date()
                e.day = e.entry_date.strftime("%A")
            except Exception:
                raise HTTPException(status_code=400, detail="Date must be in YYYY-MM-DD format")
        else:
            e.entry_date = None
    db.commit()
    return _serialize_tt(e)

# ===== TEACHER: DELETE OWN SUBJECT TIMETABLE (one click, type-to-confirm on frontend) =====
@router.delete("/timetable-subject")
def teacher_delete_tt_subject(subject: str, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher apne kisi subject ka poora timetable ek click me delete kar sakta hai.
    Scope wahi hai jo my-timetable me dikhta hai: subject unke assigned subjects me
    hona chahiye (ya unki khud ki uploaded entries). Doosre subjects delete nahi hote."""
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    from sqlalchemy import or_
    subject = (subject or "").strip()
    if not subject:
        raise HTTPException(status_code=400, detail="Subject is required")
    if subject not in (tp.subjects or []):
        owned = db.query(TimetableEntry).filter(
            TimetableEntry.subject == subject, TimetableEntry.teacher_id == tp.id).first()
        if not owned:
            raise HTTPException(status_code=403, detail="This subject is not assigned to you")
    _sc = _subj_scope_for(db, TimetableEntry, [subject])
    n = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(list(_sc))).delete(synchronize_session=False)
    db.commit()
    return {"deleted": n, "message": f"{n} entries deleted for {subject}"}

# ===== TEACHER: SEND NOTIFICATION TO STUDENTS =====
def _norm_sub(s):
    return re.sub(r"\s+", " ", (s or "").strip().lower())

def _teacher_subjects(tp):
    return [s for s in (tp.subjects or []) if (s or "").strip()]

def _students_for_subject(db, subject):
    """v93: active students jinki profile subjects mein ye subject hai (case-insensitive)."""
    from models import User, StudentProfile
    want = _norm_sub(subject)
    out = []
    rows = (db.query(StudentProfile).join(User, StudentProfile.user_id == User.id)
            .filter(User.is_active == True, User.role == "student").all())
    for sp in rows:
        subs = {_norm_sub(x) for x in (sp.subjects or []) if (x or "").strip()}
        if want in subs:
            out.append(sp)
    return out


def _students_for_subject_class(db, subject, cls):
    """Class-AWARE match: subject NAME + CLASS dono. Isse Mathematics(10) aur
    Mathematics(12), Data Entry Operations(10/12), Economics, Painting etc. kabhi
    merge nahi hote. cls = '10' / '12' (blank ho to sirf naam se — legacy)."""
    from models import User, StudentProfile
    cls = str(cls or "").strip()
    out = []
    rows = (db.query(StudentProfile).join(User, StudentProfile.user_id == User.id)
            .filter(User.is_active == True, User.role == "student").all())
    for sp in rows:
        scls = str(getattr(sp, "class_level", "") or "").strip()
        if cls and scls != cls:
            continue
        if not sp.subjects:
            continue
        if _SR is not None:
            tkey = _SR.canon_key(subject, scls)
            hit = any(_SR.canon_key(x, scls) == tkey for x in sp.subjects)
        else:
            hit = _subj_key(subject) in {_subj_key(x) for x in sp.subjects}
        if hit:
            out.append(sp)
    return out

def _my_students(db, tp):
    """v93: teacher ke kisi bhi subject wale students (union)."""
    mine = {_norm_sub(s) for s in _teacher_subjects(tp)}
    if not mine:
        return None   # fallback: sabhi active students
    from models import User, StudentProfile
    out = []
    rows = (db.query(StudentProfile).join(User, StudentProfile.user_id == User.id)
            .filter(User.is_active == True, User.role == "student").all())
    for sp in rows:
        subs = {_norm_sub(x) for x in (sp.subjects or []) if (x or "").strip()}
        if subs & mine:
            out.append(sp)
    return out

@router.get("/notify-targets")
def teacher_notify_targets(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Notify modal — subject targets SPLIT BY CLASS (student-counts jaisa). Same-name
    subjects (Mathematics 10/12, Data Entry Operations, Economics ...) alag alag rows,
    composite key 'subject|class' se — kabhi merge nahi hote."""
    from models import User, StudentProfile
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    students = (db.query(StudentProfile).join(User, StudentProfile.user_id == User.id)
                .filter(User.is_active == True, User.role == "student").all())
    subjects = []
    seen_sk = set()
    all_ids = set()
    for s in subs:
        sk = _subj_key(s)
        if sk in seen_sk:
            continue
        seen_sk.add(sk)
        per_cls = {}
        for sp in students:
            if not sp.subjects:
                continue
            cls = str(getattr(sp, "class_level", "") or "").strip() or "?"
            if _SR is not None:
                tkey = _SR.canon_key(s, cls)
                hit = any(_SR.canon_key(x, cls) == tkey for x in sp.subjects)
            else:
                hit = sk in {_subj_key(x) for x in sp.subjects}
            if not hit:
                continue
            per_cls.setdefault(cls, []).append(sp.id)
        for cls in sorted(per_cls):
            disp = _SR.canon_display(s, cls) if _SR else s
            ids = per_cls[cls]
            all_ids.update(ids)
            subjects.append({"key": "%s|%s" % (s, cls), "name": disp,
                             "class": cls, "count": len(ids)})
    subjects.sort(key=lambda x: (-x["count"], (x["name"] or ""), x["class"]))
    return {"subjects": subjects, "all_count": len(all_ids), "scoped": bool(subs)}

@router.post("/notify")
def teacher_notify(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import User
    import uuid
    tp = get_teacher_profile(current_user, db)
    title = (payload.get("title") or "").strip()
    message = (payload.get("message") or "").strip()
    subject = (payload.get("subject") or "").strip()
    if not title or not message:
        raise HTTPException(status_code=400, detail="Title and message are required")
    if subject:
        if "|" in subject:
            _nm, _cls = subject.split("|", 1)
            _nm, _cls = _nm.strip(), _cls.strip()
            sps = _students_for_subject_class(db, _nm, _cls)
            label = (_SR.canon_display(_nm, _cls) if _SR is not None else _nm) + \
                    (" (Class %s)" % _cls if _cls and _cls != "?" else "")
        else:
            sps = _students_for_subject(db, subject)   # legacy name-only
            label = subject
    else:
        # All My Students — teacher ke (subject+class) combos ka class-aware union
        seen, sps = set(), []
        for s in (tp.subjects or []):
            for cl in ("10", "12"):
                for sp in _students_for_subject_class(db, s, cl):
                    if sp.id not in seen:
                        seen.add(sp.id); sps.append(sp)
        if sps:
            label = "All My Students"
        else:
            sps = db.query(StudentProfile).join(User, StudentProfile.user_id == User.id).filter(
                User.is_active == True, User.role == "student").all()
            label = "All Students"
    batch = uuid.uuid4().hex[:24]
    sender = "👨‍🏫 " + current_user.name
    sent = 0
    for sp in sps:
        if not sp.user_id:
            continue
        notify(db, sp.user_id, sender + ": " + title, message, "teacher_message",
               sender_id=current_user.id, sender_role="teacher",
               batch_key=batch, batch_label=label)
        sent += 1
    db.commit()
    scope_txt = f"{label} — {sent} student{'s' if sent != 1 else ''}"
    return {"message": f"Sent to {scope_txt}!", "count": sent, "batch_key": batch, "label": label}

@router.get("/notify-log")
def teacher_notify_log(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """v93: meri bheji hui notifications — kitne students ne dekhi (views)."""
    rows = (db.query(Notification)
            .filter(Notification.sender_id == current_user.id,
                    Notification.sender_role == "teacher",
                    Notification.batch_key.isnot(None))
            .order_by(Notification.created_at.desc()).all())
    batches = {}
    for n in rows:
        b = batches.get(n.batch_key)
        if not b:
            b = {"batch_key": n.batch_key, "title": n.title, "message": n.message,
                 "label": n.batch_label or "Students", "sent": 0, "viewed": 0,
                 "created_at": n.created_at.isoformat() if n.created_at else None}
            batches[n.batch_key] = b
        b["sent"] += 1
        if n.is_read:
            b["viewed"] += 1
    return sorted(batches.values(), key=lambda b: b["created_at"] or "", reverse=True)

@router.get("/notify-log/{batch_key}")
def teacher_notify_log_detail(batch_key: str, db: Session = Depends(get_db),
                              current_user=Depends(get_teacher)):
    """v93: ek batch ke recipients — kis student ne dekha, kab dekha."""
    from models import User, StudentProfile
    rows = (db.query(Notification)
            .filter(Notification.sender_id == current_user.id,
                    Notification.sender_role == "teacher",
                    Notification.batch_key == batch_key)
            .order_by(Notification.created_at.asc()).all())
    if not rows:
        raise HTTPException(status_code=404, detail="Batch not found")
    out = []
    for n in rows:
        sp = db.query(StudentProfile).filter(StudentProfile.user_id == n.user_id).first()
        name = "Student"
        if sp and sp.user:
            name = sp.user.name
        else:
            u = db.query(User).filter(User.id == n.user_id).first()
            if u:
                name = u.name
        out.append({"name": name, "read": bool(n.is_read),
                    "read_at": n.read_at.isoformat() if n.read_at else None})
    out.sort(key=lambda r: (r["read"], r["name"].lower()), reverse=True)
    return {"batch_key": batch_key, "title": rows[0].title, "label": rows[0].batch_label or "Students",
            "recipients": out}

# ===== STUDY MATERIAL (PDF upload to DB) =====
@router.post("/material")
async def upload_material(
    file: UploadFile = File(...),
    subject: str = Form(...),
    class_name: str = Form("Class 12"),
    chapter: str = Form(""),
    material_type: str = Form("notes"),   # notes | dpp | test | other
    title: str = Form(""),
    category: str = Form(""),
    duration_min: int = Form(0),
    db: Session = Depends(get_db),
    current_user=Depends(get_teacher)
):
    import base64
    from models import Material
    tp = get_teacher_profile(current_user, db)
    raw = await file.read()
    if len(raw) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File is larger than 20MB. Please use a smaller PDF.")
    b64 = __import__("r2_storage").store_file_value(__import__("r2_storage").new_key("materials", (file.filename or "file.pdf")), raw, (file.content_type or "application/pdf"))
    if _SR is not None:
        subject = _SR.canon_display(subject.strip(), class_name)
    # v123: class hamesha subject ke official class se resolve karo — single-class
    # subject (Social Science=10, Physics=12) pe galat/hardcoded class kabhi save
    # na ho. Dual-class subject (English/Hindi/DEO) me jo class aayi hai use rakho.
    _cls_fixed = _subj_class_digits(db, subject.strip())
    if _cls_fixed:
        class_name = "Class " + _cls_fixed
    elif not class_name.strip():
        class_name = "Class 12"   # legacy default (unknown/dual-class, kuch nahi bheja)
    m = Material(
        teacher_id=tp.id, teacher_name=current_user.name, subject=subject.strip(),
        class_name=class_name.strip(), chapter=chapter.strip(),
        material_type=material_type.strip(), title=(title.strip() or file.filename),
        category=(category.strip() or None),
        filename=file.filename, content_b64=__import__("r2_storage").normalize(b64, "materials", (getattr(file, "content_type", None) or "application/pdf")),
        duration_min=(duration_min or None)
    )
    db.add(m); db.commit(); db.refresh(m)
    # Notify students who have this subject
    try:
        from models import StudentProfile
        label = {"notes": "Class Notes", "dpp": "DPP", "test": "Test"}.get(m.material_type, (m.category or "Material"))
        sps = db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all()
        _nk = _subj_norm(subject.strip())
        for sp in sps:
            if sp.subjects and _nk in {_subj_norm(x) for x in sp.subjects} and sp.user:
                notify(db, sp.user.id, f"📚 New {label}: {subject.strip()}",
                       f"{current_user.name} ne {subject.strip()} ({chapter.strip() or 'General'}) ke liye {label} upload ki hai. Materials section mein dekho!",
                       "new_material")
        db.commit()
    except Exception:
        db.rollback()
    return {"id": m.id, "message": "Uploaded successfully!"}

@router.get("/materials")
def teacher_materials(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material
    tp = get_teacher_profile(current_user, db)
    ms = db.query(Material).options(defer(Material.content_b64)).filter(Material.teacher_id == tp.id,
                                   Material.material_type != "answer").order_by(Material.created_at.desc()).all()
    return [{"id": m.id, "subject": m.subject, "chapter": m.chapter, "type": m.material_type,
             "title": m.title, "filename": m.filename, "duration_min": m.duration_min,
             "date": str(m.created_at)[:10]} for m in ms]

@router.get("/chapter-status")
def chapter_status(subject: str, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """For a subject, list chapters from timetable + whether notes/dpp uploaded."""
    from models import TimetableEntry, Material
    chapters = [r[0] for r in db.query(TimetableEntry.chapter).filter(
        TimetableEntry.subject == subject,
        TimetableEntry.entry_type == "chapter").distinct().all() if r[0]]
    mats = db.query(Material).options(defer(Material.content_b64)).filter(Material.subject == subject).all()
    out = []
    for ch in chapters:
        notes = any(m.chapter == ch and m.material_type == "notes" for m in mats)
        dpp = any(m.chapter == ch and m.material_type == "dpp" for m in mats)
        out.append({"chapter": ch, "notes": notes, "dpp": dpp})
    return out

@router.get("/material/{mid}/download")
def teacher_download(mid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import base64
    from fastapi import Response
    from models import Material
    m = db.query(Material).options(defer(Material.content_b64)).filter(Material.id == mid).first()
    if not m: raise HTTPException(status_code=404, detail="Not found")
    return __import__("r2_storage").proxy_response(m.content_b64, "application/pdf", m.filename or "file.pdf", True, sniff=True)

@router.delete("/material/{mid}")
def delete_material(mid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material
    tp = get_teacher_profile(current_user, db)
    m = db.query(Material).options(defer(Material.content_b64)).filter(Material.id == mid, Material.teacher_id == tp.id).first()
    if not m: raise HTTPException(status_code=404, detail="Not found")
    db.delete(m); db.commit()
    return {"message": "Deleted"}

# ===== TEACHER PROFILE & SUBJECT SELECTION (class-wise) =====
@router.get("/profile")
def teacher_profile(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    sc = tp.subject_classes or []
    return {
        "name": current_user.name,
        "user_id": current_user.user_id,
        "gender": tp.gender,
        "subjects": tp.subjects or [],
        "subject_classes": sc,
        "phone": tp.phone,
        "batch": tp.batch,
        "has_photo": bool(tp.photo_b64),
        "can_see_students": _teacher_sees_students(tp, db),
        "needs_subjects": len(sc) == 0
    }

@router.get("/available-subjects")
def teacher_available_subjects(class_level: str, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import AvailableSubject
    subs = db.query(AvailableSubject).filter(
        AvailableSubject.class_level == class_level, AvailableSubject.is_active == True).all()
    return [{"name": s.name, "code": s.code} for s in subs]

@router.post("/set-subjects")
def teacher_set_subjects(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    selections = payload.get("selections", [])   # [{"subject":..,"class":"10"/"12"}]
    if not selections:
        raise HTTPException(status_code=400, detail="Select at least 1 subject")
    if _SR is not None:
        selections = [dict(x, subject=_SR.canon_display(x.get("subject"), x.get("class") or x.get("class_name")))
                      for x in selections if x.get("subject")]
    tp.subject_classes = selections
    tp.subjects = sorted({s.get("subject") for s in selections if s.get("subject")})
    db.commit()
    return {"message": "Subjects save ho gaye!", "subjects": tp.subjects}

# ===== TEACHER: VIEW TIMETABLE (by their subjects, admin-uploaded) =====
@router.get("/my-timetable")
def my_timetable(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    if not subs:
        return []
    from sqlalchemy import or_
    # Naam me code-suffix variant ho ('X (229)') to bhi match — display canonical naam se
    scope = _subj_scope_for(db, TimetableEntry, subs)
    es = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(list(scope)),
        or_(TimetableEntry.status==None, TimetableEntry.status!='pending')).order_by(
        TimetableEntry.subject, TimetableEntry.entry_date).all()
    _ck, _has = _teacher_classkeys(tp)
    if _has:
        es = [e for e in es if (not str(e.class_name or '').strip()) or _tt_entry_key(e.subject, e.class_name) in _ck]
    return [_serialize_tt(e, scope.get(e.subject)) for e in es]

# ===== TEACHER: TODAY'S CLASSES with material status =====
@router.get("/today-classes")
def today_classes(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, Material
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    if not subs:
        return []
    today = date.today()
    from sqlalchemy import or_
    scope = _subj_scope_for(db, TimetableEntry, subs)
    es = db.query(TimetableEntry).filter(
        TimetableEntry.subject.in_(list(scope)), TimetableEntry.entry_date == today,
        or_(TimetableEntry.status==None, TimetableEntry.status!='pending')).all()
    _ck, _has = _teacher_classkeys(tp)
    if _has:
        es = [e for e in es if (not str(e.class_name or '').strip()) or _tt_entry_key(e.subject, e.class_name) in _ck]
    mats = db.query(Material).options(defer(Material.content_b64)).filter(Material.subject.in_(list(_subj_scope_for(db, Material, subs)))).all()
    out = []
    for e in es:
        notes = any(m.chapter == e.chapter and _subj_eq(m.subject, e.subject) and m.material_type == "notes" for m in mats)
        dpp = any(m.chapter == e.chapter and _subj_eq(m.subject, e.subject) and m.material_type == "dpp" for m in mats)
        d = _serialize_tt(e, scope.get(e.subject)); d["notes"] = notes; d["dpp"] = dpp
        out.append(d)
    out.sort(key=lambda x: x.get("time") or "")
    return out

# ===== TEACHER: REQUEST EXTRA CLASS (needs admin approval) =====
@router.post("/request-class")
def request_class(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, User, UserRole, Notification
    tp = get_teacher_profile(current_user, db)
    subject = (payload.get("subject") or "").strip()
    if subject not in (tp.subjects or []):
        raise HTTPException(status_code=400, detail="This is not your subject")
    edate = None
    if payload.get("date"):
        try:
            from datetime import datetime as _dt
            edate = _dt.strptime(payload["date"], "%Y-%m-%d").date()
        except Exception:
            pass
    day = edate.strftime("%a") if edate else None
    e = TimetableEntry(
        teacher_id=tp.id, subject=subject, class_name=payload.get("class_name", "Class 12"),
        chapter=(payload.get("chapter") or payload.get("topic") or "Extra Class").strip(),
        part=((payload.get("topic") or "").strip() or None),
        entry_date=edate, day=day, time_text=(payload.get("time") or "").strip() or None,
        entry_type="chapter", status="pending"
    )
    db.add(e); db.flush()
    # notify admins
    for adm in db.query(User).filter(User.role == UserRole.admin).all():
        db.add(Notification(user_id=adm.id, title="New Extra Class Request",
                            message=f"{current_user.name} requested an extra class for {subject} ({payload.get('date','')} {payload.get('time','')}). Please approve it.",
                            notif_type="class_request"))
    db.commit(); db.refresh(e)
    return {"id": e.id, "message": "Request sent to the admin! It will appear in the timetable once approved."}

# ===== TEACHER: SUBJECT-WISE STUDENT COUNTS =====

def _exam_ranking_rows(db, exam_id):
    """Graded attempts -> ranked rows (top 3 podium + list). Shared by all roles."""
    from models import Exam, ExamAttempt, StudentProfile
    exam = db.query(Exam).filter(Exam.id == exam_id).first()
    if not exam:
        return None
    atts = (db.query(ExamAttempt)
            .filter(ExamAttempt.exam_id == exam_id, ExamAttempt.status == "graded")
            .all())
    rows = []
    for a in atts:
        sp = db.query(StudentProfile).filter(StudentProfile.id == a.student_id).first()
        nm = a.student_name or ((sp.user.name if sp and sp.user else "") or "Student")
        att_n = None
        if a.attempted:
            att_n = len(a.attempted)
        elif a.mcq_answers:
            att_n = len([v for v in (a.mcq_answers or {}).values() if v not in (None, "", [])])
        rows.append({"student_id": a.student_id, "name": nm,
                     "marks": round(float(a.total_awarded or 0), 1),
                     "attempted": att_n,
                     "batch": (sp.batch_name if sp else None),
                     "has_photo": bool(sp and sp.photo_b64)})
    rows.sort(key=lambda r: (-r["marks"], r["name"].lower()))
    for i, r in enumerate(rows, 1):
        r["rank"] = i
    return {"exam": {"id": exam.id, "title": exam.title, "subject": exam.subject,
                     "chapter": exam.chapter, "total_marks": exam.total_marks,
                     "test_type": exam.test_type},
            "graded": len(rows), "rows": rows}

@router.get("/exam/{exam_id}/ranking")
def teacher_exam_ranking(exam_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    data = _exam_ranking_rows(db, exam_id)
    if not data:
        raise HTTPException(status_code=404, detail="Test not found")
    data["me_id"] = None
    return data


# ================== DPP PACKS (Create / Upload / Results) ==================
def _dpp_pack_out(db, pk, with_counts=True):
    out = {"id": pk.id, "subject": pk.subject, "class_name": pk.class_name,
           "chapter": pk.chapter, "part": pk.part, "title": pk.title,
           "medium": pk.medium, "source": pk.source,
           "created_at": pk.created_at.strftime("%d %b %Y") if pk.created_at else None}
    if with_counts:
        from models import DppAnswer, DppEvent
        subs = (db.query(DppAnswer).filter(DppAnswer.pack_id == pk.id,
                                           DppAnswer.status != "staged").all())
        out["submitted"] = len(subs)
        out["checked"] = sum(1 for a in subs if a.status == "checked")
        # tracker: kitne DISTINCT students ne view / download kiya
        out["views"] = (db.query(DppEvent.student_id)
                        .filter(DppEvent.pack_id == pk.id, DppEvent.event == "view")
                        .distinct().count())
        out["downloads"] = (db.query(DppEvent.student_id)
                            .filter(DppEvent.pack_id == pk.id, DppEvent.event == "download")
                            .distinct().count())
    return out


@router.get("/dpp-packs")
def teacher_dpp_packs(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import DppPack
    tp = get_teacher_profile(current_user, db)
    packs = (db.query(DppPack).filter(DppPack.teacher_id == tp.id)
             .order_by(DppPack.created_at.desc()).all())
    return {"packs": [_dpp_pack_out(db, pk) for pk in packs]}


# ===== DPP RANKING — sabse pehle submit karne wala rank #1 (teacher = full details) =====
@router.get("/dpp-packs/{pack_id}/ranking")
def teacher_dpp_ranking(pack_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import DppPack, DppAnswer, StudentProfile, User
    tp = get_teacher_profile(current_user, db)
    _see = _teacher_sees_students(tp, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(404, "DPP pack not found")
    rows = (db.query(DppAnswer)
            .filter(DppAnswer.pack_id == pack_id, DppAnswer.status != "staged")
            .order_by(DppAnswer.submitted_at.asc(), DppAnswer.id.asc()).all())
    out = []
    for i, a in enumerate(rows):
        srow = db.query(StudentProfile).filter(StudentProfile.id == a.student_id).first()
        nm = ""
        if srow:
            u = db.query(User).filter(User.id == srow.user_id).first()
            nm = (u.name if u else "") or ""
        nm = nm or f"Student #{a.student_id}"
        out.append({"rank": i + 1, "name": nm, "student_id": a.student_id,
                    "submitted_at": a.submitted_at.strftime("%d %b %Y, %I:%M %p") if a.submitted_at else "",
                    "status": a.status or "submitted",
                    "checked_at": a.checked_at.strftime("%d %b %Y, %I:%M %p") if getattr(a, "checked_at", None) else "",
                    "checked_by": getattr(a, "checked_by", None) or "",
                    "remarks": a.remarks or "",
                    "filename": a.filename or "",
                    "class_name": (srow.class_name if srow else "") or "",
                    "phone": (((srow.phone if srow else "") or "") if _see else "")})
    return {"pack": {"id": pk.id, "title": pk.title or "DPP", "subject": pk.subject or "",
                     "chapter": pk.chapter or "", "part": pk.part or "",
                     "class_name": getattr(pk, "class_name", "") or ""},
            "count": len(out), "rows": out}


# ===== DPP TRACK-LIST — kis kis student ne view/download kiya (clickable chips ke liye) =====
@router.get("/dpp-packs/{pack_id}/track-list")
def teacher_dpp_track_list(pack_id: int, event: str = "view",
                           db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import DppPack, DppEvent, StudentProfile, User
    tp = get_teacher_profile(current_user, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(404, "DPP pack not found")
    ev = (event or "view").lower()
    if ev not in ("view", "download"):
        raise HTTPException(400, "event 'view' ya 'download' hona chahiye")
    rows = (db.query(DppEvent)
            .filter(DppEvent.pack_id == pack_id, DppEvent.event == ev)
            .order_by(DppEvent.created_at.desc(), DppEvent.id.desc()).all())
    seen = set()
    out = []
    for e in rows:
        if e.student_id in seen:
            continue
        seen.add(e.student_id)
        nm = ""
        srow = db.query(StudentProfile).filter(StudentProfile.id == e.student_id).first()
        if srow:
            u = db.query(User).filter(User.id == srow.user_id).first()
            nm = (u.name if u else "") or ""
        out.append({"name": nm or f"Student #{e.student_id}",
                    "at": e.created_at.strftime("%d %b %Y, %I:%M %p") if e.created_at else ""})
    return {"pack": {"id": pk.id, "title": pk.title or "DPP"},
            "event": ev, "count": len(out), "rows": out}


def _dpp_build_pdf(db, pk, kind="q", med=None):
    """Created pack se on-demand premium PDF (language select ke saath).
    kind=q -> questions paper (no answers), kind=s -> solutions paper."""
    from types import SimpleNamespace as _NS
    import base64 as _b64
    import exam_pdf
    tname = ""
    tphoto = None
    try:
        from models import TeacherProfile as _TP
        _tp = db.query(_TP).filter(_TP.id == pk.teacher_id).first()
        tname = _tp.user.name if _tp and _tp.user else "Teacher"
        tphoto = _tp.photo_b64 if _tp else None
    except Exception:
        tname = "Teacher"
    med = (med or pk.medium or "english").lower()
    med = "both" if med.startswith("bo") else ("hindi" if med.startswith("hin") else "english")
    ex = _NS(teacher_name=tname, teacher_photo_b64=tphoto, title=pk.title, subject=pk.subject,
             chapter=pk.chapter, part=pk.part, class_name=(pk.class_name or ""), test_type="DPP",
             duration_mins=None, total_marks=None)
    qobjs = []
    for i, q in enumerate(pk.questions or [], 1):
        base = dict(q_no=i, question_text=q.get("q"), max_marks=q.get("marks"),
                    options=None, correct_option=None, image_b64=_r2img(q.get("image")),
                    alt_image_b64=_r2img(q.get("alt_image")), question_text_hi=q.get("q_hi"),
                    options_hi=None, explanation=None, explanation_hi=None)
        if kind == "s":
            qobjs.append(_NS(**base, model_answer=q.get("model"),
                             model_answer_hi=q.get("model_hi"),
                             model_answer_image=_r2img(q.get("model_image"))))
        else:
            qobjs.append(_NS(**base, model_answer=None, model_answer_hi=None,
                             model_answer_image=None))
    # professional gold-band DPP layout (question paper / solutions)
    return _b64.b64encode(exam_pdf.build_dpp_pdf(ex, qobjs, med, kind)).decode()


@router.get("/dpp-packs/{pack_id}/pdf")
def teacher_dpp_pdf(pack_id: int, kind: str = "q", medium: str = "",
                    db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """View/Download ke liye: created DPP -> language ke saath fresh premium PDF;
    uploaded DPP -> stored file (language ignore)."""
    from models import DppPack
    tp = get_teacher_profile(current_user, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(status_code=404, detail="DPP not found")
    blob = None
    if pk.source == "created" and pk.questions:
        try:
            blob = _dpp_build_pdf(db, pk, kind, medium or pk.medium)
            # Cache: agli baar instant (default medium wala hi store karo)
            try:
                med_l = (medium or pk.medium or "").lower()
                if not med_l.startswith("hin") or (pk.medium or "").lower().startswith("hin"):
                    if kind == "s":
                        pk.s_pdf = pk.s_pdf or blob
                    else:
                        pk.q_pdf = pk.q_pdf or blob
                    db.commit()
            except Exception:
                db.rollback()
        except Exception:
            blob = (pk.q_pdf if kind == "q" else pk.s_pdf)
    else:
        blob = (pk.q_pdf if kind == "q" else pk.s_pdf)
    blob = _dpp_shrink(db, pk, kind, blob)
    if not blob:
        raise HTTPException(status_code=404, detail="File not available")
    med = ("both" if (medium or "").lower().startswith("bo") else ("hindi" if (medium or "").lower().startswith("hin") else "english"))
    fname = (pk.title or "DPP").replace("/", "-") + \
        ("-solutions" if kind == "s" else "-questions") + "-" + med + ".pdf"
    # R2 URL ho to redirect, base64 ho to decode (dono safe)
    return __import__("r2_storage").file_response(blob, "application/pdf", fname, False)


@router.post("/dpp-packs/create")
def teacher_dpp_create(data: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Editor se bana DPP — har question ka model answer MANDATORY;
    questions paper (no answers) + solutions paper (with answers) dono PDFs generate."""
    from models import DppPack
    tp = get_teacher_profile(current_user, db)
    subject = (data.get("subject") or "").strip()
    chapter = (data.get("chapter") or "").strip()
    part = (data.get("part") or "").strip()
    title = (data.get("title") or "").strip() or ("DPP - " + (part or chapter or subject))
    medium = (data.get("medium") or "English").strip()
    questions = data.get("questions") or []
    if not subject:
        raise HTTPException(status_code=400, detail="Subject is required")
    if not questions:
        raise HTTPException(status_code=400, detail="Add at least one question")
    for i, q in enumerate(questions, 1):
        if not (q.get("q") or "").strip() and not q.get("image"):
            raise HTTPException(status_code=400, detail=f"Question {i} is empty")
        if not (q.get("model") or "").strip() and not q.get("model_image"):
            raise HTTPException(status_code=400,
                                detail=f"Answer mandatory: Question {i} ka model answer bharo")
    if _SR is not None:
        subject = _SR.canon_display(subject, data.get("class_name"))
    # Idempotency: network error ke baad retry pe DUPLICATE DPP na bane
    # (same teacher + same title, 90 sec ke andar -> purana pack hi return)
    ck = (data.get("client_key") or "").strip()
    if ck:
        try:
            dup = (db.query(DppPack)
                   .filter(DppPack.teacher_id == tp.id, DppPack.title == title,
                           DppPack.created_at >= datetime.utcnow() - timedelta(seconds=90))
                   .first())
            if dup:
                return {"ok": True, "pack": _dpp_pack_out(db, dup, False), "duplicate": True}
        except Exception:
            db.rollback()
    pk = DppPack(teacher_id=tp.id, subject=subject, class_name=(data.get("class_name") or ""),
                 chapter=chapter, part=part, title=title, medium=medium,
                 source="created", questions=questions)
    # PDFs LAZY banenge (view/download pe on-demand) — create instant + crash-proof.
    # Bade images ke saath sync PDF build se server OOM/crash ho sakta tha.
    db.add(pk); db.commit(); db.refresh(pk)
    return {"ok": True, "pack": _dpp_pack_out(db, pk, False)}


@router.patch("/dpp-packs/{pack_id}")
def teacher_dpp_update(pack_id: int, data: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Editor se bana DPP (source='created') edit karo — title/subject/chapter/part/
    medium/class update + questions poora replace. Uploaded (PDF) DPP builder me edit
    nahi hote. Har question ka model answer yahan bhi MANDATORY (create jaisa)."""
    from models import DppPack
    tp = get_teacher_profile(current_user, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(status_code=404, detail="DPP not found")
    if (pk.source or "created") != "created":
        raise HTTPException(status_code=400, detail="Uploaded DPPs can't be edited here — delete and upload again.")
    questions = data.get("questions") or []
    if not questions:
        raise HTTPException(status_code=400, detail="Add at least one question")
    for i, q in enumerate(questions, 1):
        if not (q.get("q") or "").strip() and not q.get("image"):
            raise HTTPException(status_code=400, detail=f"Question {i} is empty")
        if not (q.get("model") or "").strip() and not q.get("model_image"):
            raise HTTPException(status_code=400,
                                detail=f"Answer mandatory: Question {i} ka model answer bharo")
    subject = (data.get("subject") or pk.subject or "").strip()
    if _SR is not None and subject:
        subject = _SR.canon_display(subject, data.get("class_name") or pk.class_name)
    if subject:
        pk.subject = subject
    if data.get("class_name") is not None:
        pk.class_name = (data.get("class_name") or "").strip()
    if data.get("chapter") is not None:
        pk.chapter = (data.get("chapter") or "").strip()
    if data.get("part") is not None:
        pk.part = (data.get("part") or "").strip()
    if (data.get("title") or "").strip():
        pk.title = (data.get("title") or "").strip()
    if (data.get("medium") or "").strip():
        pk.medium = (data.get("medium") or "").strip()
    pk.questions = questions
    # Stored PDFs are stale after an edit — clear so they rebuild lazily on next view/download.
    for _attr in ("q_pdf", "s_pdf", "q_pdf_hi", "s_pdf_hi"):
        if hasattr(pk, _attr):
            try:
                setattr(pk, _attr, None)
            except Exception:
                pass
    db.add(pk); db.commit(); db.refresh(pk)
    return {"ok": True, "pack": _dpp_pack_out(db, pk, False)}


def _compress_pdf(pdf_bytes):
    """Uploaded PDF ko chhota karo (lossless optimize) — download fast + server load kam,
    quality bilkul same. fitz (PyMuPDF) na ho ya fail ho to file as-is (kabhi crash nahi)."""
    try:
        if not pdf_bytes or len(pdf_bytes) < 60000:   # chhoti file ko chhedne ki zaroorat nahi
            return pdf_bytes
        import fitz
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        out = doc.tobytes(garbage=4, deflate=True, deflate_images=True,
                          deflate_fonts=True, clean=True)
        doc.close()
        if out and 0 < len(out) < len(pdf_bytes):
            return out
    except Exception:
        pass
    return pdf_bytes


def _dpp_shrink(db, pk, kind, blob):
    """Uploaded pack ka BADA PDF ek baar compress karke DB me chhota save + wahi serve.
    Guarded: chhoti (<500KB) file ko chhodo; >8% shrink pe hi save. Fail ho to as-is —
    download kabhi nahi tootega. Purane bade DPPs isse next open pe hi chhote ho jaate hain."""
    try:
        if not blob or (isinstance(blob, str) and blob.startswith("http")) \
                or (getattr(pk, "source", "") == "created"):
            return blob
        raw = base64.b64decode(blob)
        if len(raw) < 500000:
            return blob
        smaller = _compress_pdf(raw)
        if smaller and len(smaller) < len(raw) * 0.92:
            nb = base64.b64encode(smaller).decode()
            if kind == "s":
                pk.s_pdf = nb
            else:
                pk.q_pdf = nb
            db.commit()
            return nb
    except Exception:
        try:
            db.rollback()
        except Exception:
            pass
    return blob


@router.post("/dpp-packs/upload")
async def teacher_dpp_upload(subject: str = Form(...), chapter: str = Form(""), part: str = Form(""),
                             title: str = Form(""), medium: str = Form("English"),
                             class_name: str = Form(""),
                             q_pdf: UploadFile = File(...), s_pdf: UploadFile = File(...),
                             db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Ready-made DPP upload — questions PDF + solutions PDF dono MANDATORY."""
    from models import DppPack
    tp = get_teacher_profile(current_user, db)
    qd = await q_pdf.read(); sd = await s_pdf.read()
    if not qd or not sd:
        raise HTTPException(status_code=400, detail="Questions aur Solutions dono PDF upload karna zaroori hai")
    qd = _compress_pdf(qd); sd = _compress_pdf(sd)   # upload pe hi compress -> download fast
    if _SR is not None:
        subject = _SR.canon_display(subject.strip(), class_name)
    pk = DppPack(teacher_id=tp.id, subject=subject.strip(), class_name=class_name.strip(),
                 chapter=chapter.strip(), part=part.strip(),
                 title=(title.strip() or "DPP - " + (part.strip() or chapter.strip() or subject.strip())),
                 medium=medium, source="uploaded", questions=[],
                 q_pdf=__import__("r2_storage").store_file_value(__import__("r2_storage").new_key("dpp-pdf", "q.pdf"), qd, "application/pdf"), s_pdf=__import__("r2_storage").store_file_value(__import__("r2_storage").new_key("dpp-pdf", "s.pdf"), sd, "application/pdf"))
    db.add(pk); db.commit(); db.refresh(pk)
    return {"ok": True, "pack": _dpp_pack_out(db, pk, False)}


@router.get("/dpp-packs/{pack_id}/questions")
def teacher_dpp_questions(pack_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Panel view ke liye pack ke questions — model answers SAHIT (teacher ka apna pack)."""
    from models import DppPack
    tp = get_teacher_profile(current_user, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(status_code=404, detail="DPP not found")
    qs = []
    for i, q in enumerate(pk.questions or [], 1):
        qs.append({"qno": i, "q": q.get("q") or "", "q_hi": q.get("q_hi") or "",
                   "image": q.get("image"), "alt_image": q.get("alt_image"),
                   "model": q.get("model") or "", "model_hi": q.get("model_hi") or "",
                   "model_image": q.get("model_image")})
    tname = ""
    tph = False
    try:
        tname = tp.user.name if tp and tp.user else ""
        tph = bool(tp.photo_b64)
    except Exception:
        pass
    return {"id": pk.id, "title": pk.title, "subject": pk.subject, "chapter": pk.chapter,
            "part": pk.part, "medium": pk.medium, "source": pk.source, "teacher": tname,
            "teacher_id": tp.id, "has_teacher_photo": tph,
            "class_name": pk.class_name, "has_solution": bool(pk.s_pdf) or any(
                (q.get("model") or q.get("model_image")) for q in (pk.questions or [])),
            "questions": qs}


@router.get("/dpp-packs/{pack_id}/answers")
def teacher_dpp_answers(pack_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import DppPack, DppAnswer, StudentProfile, User
    tp = get_teacher_profile(current_user, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(status_code=404, detail="DPP not found")
    rows = (db.query(DppAnswer).filter(DppAnswer.pack_id == pack_id,
                                       DppAnswer.status != "staged")
            .order_by(DppAnswer.submitted_at.desc()).all())
    out = []
    for a in rows:
        sp = db.query(StudentProfile).filter(StudentProfile.id == a.student_id).first()
        nm = ""
        if sp and sp.user_id:
            u = db.query(User).filter(User.id == sp.user_id).first()
            nm = u.name if u else ""
        out.append({"id": a.id, "student": nm or f"Student #{a.student_id}",
                    "class_level": (sp.class_level if sp else None),
                    "filename": a.filename, "status": a.status, "remarks": a.remarks,
                    "allow_resubmit": bool(getattr(a, "allow_resubmit", False)),
                    "submitted_at": a.submitted_at.strftime("%d %b %Y, %I:%M %p") if a.submitted_at else None,
                    "checked_at": a.checked_at.strftime("%d %b %Y") if a.checked_at else None})
    return {"pack": _dpp_pack_out(db, pk, False), "answers": out}


@router.post("/dpp-packs/{pack_id}/answers/{answer_id}/allow-resubmit")
def teacher_dpp_allow_resubmit(pack_id: int, answer_id: int, data: dict = Body(...),
                               db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher kisi student ka re-submit ON/OFF kare — ON pe student ko dobara
    'Re-submit' option dikhta hai; wo ek baar use hote hi apne aap OFF ho jata hai."""
    from models import DppPack, DppAnswer
    tp = get_teacher_profile(current_user, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(status_code=404, detail="DPP not found")
    a = db.query(DppAnswer).filter(DppAnswer.id == answer_id, DppAnswer.pack_id == pack_id,
                                   DppAnswer.status != "staged").first()
    if not a:
        raise HTTPException(status_code=404, detail="Submission not found")
    a.allow_resubmit = bool((data or {}).get("allow", True))
    db.commit()
    return {"ok": True, "allow_resubmit": a.allow_resubmit}


@router.get("/dpp-answers/{answer_id}/file")
def teacher_dpp_answer_file(answer_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import DppAnswer
    get_teacher_profile(current_user, db)
    a = db.query(DppAnswer).filter(DppAnswer.id == answer_id).first()
    if not a or not a.answer_b64:
        raise HTTPException(status_code=404, detail="File not found")
    return __import__("r2_storage").file_response(a.answer_b64, "application/pdf", a.filename or "dpp-answer.pdf", True)


@router.delete("/dpp-packs/{pack_id}")
def delete_dpp_pack(pack_id: int, db: Session = Depends(get_db), user=Depends(get_teacher)):
    """Teacher apni banaayi/upload ki gayi DPP delete kar sake — saari
    submissions bhi saath hat jati hain."""
    from models import DppPack, DppAnswer, DppEvent, DppChunk
    tp = db.query(TeacherProfile).filter(TeacherProfile.user_id == user.id).first()
    pk = db.query(DppPack).get(pack_id)
    if not pk or not tp or pk.teacher_id != tp.id:
        raise HTTPException(status_code=404, detail="DPP not found")
    # v155: saare child rows pehle hatao warna FK constraint (dpp_events/dpp_chunks
    # /dpp_answers -> dpp_packs) delete ko block karti hai.
    db.query(DppAnswer).filter(DppAnswer.pack_id == pack_id).delete(synchronize_session=False)
    db.query(DppEvent).filter(DppEvent.pack_id == pack_id).delete(synchronize_session=False)
    db.query(DppChunk).filter(DppChunk.pack_id == pack_id).delete(synchronize_session=False)
    db.delete(pk)
    db.commit()
    return {"message": "DPP deleted"}


@router.get("/dpp-packs/{pack_id}/file")
def teacher_dpp_pack_file(pack_id: int, kind: str = "q", db: Session = Depends(get_db),
                          current_user=Depends(get_teacher)):
    from models import DppPack
    tp = get_teacher_profile(current_user, db)
    pk = db.query(DppPack).filter(DppPack.id == pack_id, DppPack.teacher_id == tp.id).first()
    if not pk:
        raise HTTPException(status_code=404, detail="DPP not found")
    blob = None
    if pk.source == "created" and pk.questions:
        # created pack: FRESH premium PDF (naya format + teacher photo) — cache sirf fallback
        try:
            blob = _dpp_build_pdf(db, pk, kind, pk.medium)
        except Exception:
            blob = None
    if not blob:
        blob = pk.s_pdf if kind == "s" else pk.q_pdf
    if not blob:
        raise HTTPException(status_code=404, detail="File not generated")
    fname = (pk.title or "DPP").replace("/", "-") + ("-solutions.pdf" if kind == "s" else "-questions.pdf")
    # R2 URL ho to server-side stream (viewer same-origin fetch chalega, crash nahi);
    # base64 ho to decode. Dono safe.
    return __import__("r2_storage").proxy_response(blob, "application/pdf", fname, True)


@router.post("/dpp-answers/{answer_id}/check")
def teacher_dpp_check(answer_id: int, data: dict, db: Session = Depends(get_db),
                      current_user=Depends(get_teacher)):
    """Marks nahi — sirf Checked + Remarks. Student ko notification teacher ke naam se."""
    import datetime as _dt
    from models import DppAnswer, DppPack, Notification, StudentProfile
    tp = get_teacher_profile(current_user, db)
    a = db.query(DppAnswer).filter(DppAnswer.id == answer_id).first()
    if not a:
        raise HTTPException(status_code=404, detail="Submission not found")
    pk = db.query(DppPack).filter(DppPack.id == a.pack_id).first()
    tname = tp.user.name if tp.user else "Teacher"
    a.status = "checked"
    a.remarks = (data.get("remarks") or "").strip()
    a.checked_by = tname
    a.checked_at = _dt.datetime.utcnow()
    sp = db.query(StudentProfile).filter(StudentProfile.id == a.student_id).first()
    if sp and sp.user_id:
        msg = tname + " checked your DPP \"" + (pk.title if pk else "") + "\"."
        if a.remarks:
            msg += " Remarks: " + a.remarks
        db.add(Notification(user_id=sp.user_id, title="DPP Checked - " + tname, message=msg))
    db.commit()
    return {"ok": True, "status": "checked"}


@router.get("/student-counts")
def teacher_student_counts(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Per-subject counts SPLIT BY CLASS LEVEL — English (Class 10, code 202) aur
    English (Class 12, code 302) alag alag cards. Pehle same-name subjects merge
    ho jaate the aur Class 10/12 ke bache ek hi list me dikhte the."""
    from models import StudentProfile, AvailableSubject
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    students = db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all()
    # (subject_key, class_level) -> NIOS subject code
    code_map = {}
    for a in db.query(AvailableSubject).all():
        code_map[(_subj_key(a.name), str(a.class_level or "").strip())] = a.code
    out = []
    seen_ids = set()
    seen_sk = set()  # same subject do baar list me ho to duplicate card na bane
    for s in subs:
        sk = _subj_key(s)
        if sk in seen_sk:
            continue
        seen_sk.add(sk)
        per_cls = {}
        for sp in students:
            if not sp.subjects:
                continue
            cls = str(sp.class_level or "").strip() or "?"
            if _SR is not None:
                tkey = _SR.canon_key(s, cls)
                hit = any(_SR.canon_key(x, cls) == tkey for x in sp.subjects)
            else:
                hit = sk in {_subj_key(x) for x in sp.subjects}
            if not hit:
                continue
            seen_ids.add(sp.id)
            per_cls[cls] = per_cls.get(cls, 0) + 1
        for cls in sorted(per_cls):
            disp = _SR.canon_display(s, cls) if _SR else s
            code = code_map.get((sk, cls))
            if _SR is not None:
                code = (_SR.canon_subject(s, cls) or {}).get("code") or code
            out.append({"subject": disp, "class": cls,
                        "code": code, "count": per_cls[cls]})
    out.sort(key=lambda x: (-x["count"], (x["subject"] or ""), x["class"]))
    return {"total": len(seen_ids), "subjects": out}

# ===== TEACHER: VIEW SUBMISSIONS + GIVE MARKS =====
@router.get("/dpp-results")
def teacher_dpp_results(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Every DPP the teacher uploaded, with submission counts (submitted /
    checked / pending) so the DPP Result page can show progress at a glance."""
    from models import Material, MaterialView, StudentProfile
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    dpps = db.query(Material).options(defer(Material.content_b64)).filter(
        Material.material_type == "dpp",
        Material.subject.in_(subs)).order_by(Material.created_at.desc()).all() if subs else []
    ids = [m.id for m in dpps]
    answers = db.query(Material).options(defer(Material.content_b64)).filter(
        Material.material_type == "answer",
        Material.parent_id.in_(ids)).all() if ids else []
    # how many students should be doing each DPP (same subject)
    roster = {}
    for sp in db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all():
        for s in (sp.subjects or []):
            roster[s] = roster.get(s, 0) + 1
    views, downloads = {}, {}
    if ids:
        for v in db.query(MaterialView).filter(MaterialView.material_id.in_(ids)).all():
            (downloads if v.action == "download" else views).setdefault(v.material_id, set()).add(v.student_id)
    out = []
    for m in dpps:
        mine = [a for a in answers if a.parent_id == m.id]
        checked = sum(1 for a in mine if (a.marks or "").strip())
        total_students = roster.get(m.subject, 0)
        out.append({
            "id": m.id, "subject": m.subject, "chapter": m.chapter, "part": m.part,
            "title": m.title, "filename": m.filename,
            "date": str(m.created_at)[:10] if m.created_at else "",
            "views": len(views.get(m.id, ())), "downloads": len(downloads.get(m.id, ())),
            "submitted": len(mine), "checked": checked, "pending": len(mine) - checked,
            "total_students": total_students,
        })
    totals = {
        "dpps": len(out),
        "submitted": sum(o["submitted"] for o in out),
        "checked": sum(o["checked"] for o in out),
        "pending": sum(o["pending"] for o in out),
    }
    return {"totals": totals, "rows": out}


@router.get("/submissions")
def teacher_submissions(parent_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material
    subs = db.query(Material).options(defer(Material.content_b64)).filter(Material.material_type == "answer",
                                     Material.parent_id == parent_id).order_by(Material.created_at.desc()).all()
    return [{"id": m.id, "student_name": m.student_name, "marks": m.marks,
             "date": str(m.created_at)[:16]} for m in subs]

@router.post("/submission/{sid}/marks")
def set_marks(sid: int, payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material, StudentProfile, Notification
    m = db.query(Material).options(defer(Material.content_b64)).filter(Material.id == sid, Material.material_type == "answer").first()
    if not m:
        raise HTTPException(status_code=404, detail="Submission not found")
    m.marks = str(payload.get("marks", "")).strip()
    # notify student
    if m.student_id:
        sp = db.query(StudentProfile).filter(StudentProfile.id == m.student_id).first()
        if sp and sp.user:
            db.add(Notification(user_id=sp.user.id, title="DPP Checked!",
                                message=f"{current_user.name} checked your {m.subject} DPP. Marks: {m.marks}",
                                notif_type="marks"))
    db.commit()
    return {"message": "Marks save ho gaye!"}

# ===== TEACHER: OWN PHOTO + MY STUDENTS LIST =====
@router.get("/my-photo")
def teacher_my_photo(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import base64
    from fastapi import Response
    tp = get_teacher_profile(current_user, db)
    if not tp.photo_b64:
        raise HTTPException(status_code=404, detail="No photo")
    return __import__("r2_storage").photo_response(tp.photo_b64)

@router.get("/teacher/{tid}/photo")
def teacher_peer_photo(tid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """v112: doubt thread avatar — kisi bhi teacher ki photo (fallback initials)."""
    import base64
    from fastapi import Response
    tp = db.query(TeacherProfile).filter(TeacherProfile.id == tid).first()
    if not tp or not tp.photo_b64:
        raise HTTPException(status_code=404, detail="No photo")
    return __import__("r2_storage").photo_response(tp.photo_b64)

@router.get("/student/{sid}/photo")
def teacher_student_photo(sid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import base64
    from fastapi import Response
    from models import StudentProfile
    sp = db.query(StudentProfile).filter(StudentProfile.id == sid).first()
    if not sp or not sp.photo_b64:
        raise HTTPException(status_code=404, detail="No photo")
    return __import__("r2_storage").photo_response(sp.photo_b64)

def _subj_key(name):
    """Subject naam ko compare karne layak banata hai: extra space, case aur
    "(Class 12)" jaisa suffix hata deta hai. Excel upload se aksar "Physics "
    jaise trailing space aa jaate the aur student teacher ki list me hi
    nahi dikhta tha."""
    t = str(name or "")
    t = re.sub(r"\((?:class\s*)?\d+(?:th)?\)", " ", t, flags=re.I)
    t = re.sub(r"[^a-z0-9]+", " ", t.lower())
    return " ".join(t.split()).strip()


@router.get("/my-students-list")
def teacher_my_students_list(q: str = "", subject: str = "", cls: str = "", db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import StudentProfile
    tp = get_teacher_profile(current_user, db)
    if not _teacher_sees_students(tp, db):
        return []   # is teacher ke liye "My Students" access admin ne band kiya hai
    subs = tp.subjects or []
    # Class-aware canonical keys: 'PHYSICS'(Cl-12 student) aur teacher ka 'Physics'
    # dono ka key 'c312' — koi bhi case/alias/class-variant miss nahi hoga.
    def _tkeys(cls):
        if _SR is not None:
            return {_SR.canon_key(x, cls) for x in subs}
        return {_subj_key(x) for x in subs if _subj_key(x)}
    def _skeys(ssubs, cls):
        if _SR is not None:
            return {_SR.canon_key(x, cls) for x in ssubs}
        return {_subj_key(x) for x in ssubs if _subj_key(x)}
    want_cls = (cls or "").strip()
    ql = " ".join((q or "").split()).strip().lower()
    q_tokens = [t for t in ql.split(" ") if t]
    rows = db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all()
    out = []
    for sp in rows:
        ssubs = sp.subjects or []
        scls = "".join(ch for ch in str(sp.class_level or "") if ch.isdigit())[:2] or None
        tkeys = _tkeys(scls)
        matched = [x for x in ssubs if (_SR.canon_key(x, scls) if _SR else _subj_key(x)) in tkeys]
        if not matched:
            continue
        if subject:
            wk = (_SR.canon_key(subject, scls) if _SR else _subj_key(subject))
            if wk not in _skeys(ssubs, scls):
                continue
        if want_cls and str(sp.class_level or "").strip() != want_cls:
            continue
        nm = (sp.user.name if sp.user else "") or ""
        if q_tokens:
            hay = " ".join([
                " ".join(nm.split()).lower(),
                (sp.phone or ""),
                (sp.user.user_id if sp.user else "") or "",
                (sp.batch_name or ""),
            ]).lower()
            # har shabd alag se match ho - "tanu sharma" "TANU  SHARMA" se bhi mile
            if not all(t in hay for t in q_tokens):
                continue
        out.append({"id": sp.id, "name": nm, "phone": sp.phone, "class": sp.class_level,
                    "user_id": (sp.user.user_id if sp.user else None),
                    "batch": sp.batch_name, "medium": sp.medium,
                    "email": sp.email,
                    "class_name": sp.class_name, "nios_ref": sp.nios_ref,
                    "exam_session": sp.exam_session, "exam_stream": sp.exam_stream,
                    "goal": (sp.goal_custom if sp.goal == "other" else sp.goal),
                    "last_seen": sp.last_seen.strftime("%d %b %Y, %I:%M %p") if sp.last_seen else None,
                    "is_verified": bool(sp.is_verified),
                    "all_subjects": (_SR.canon_list(ssubs, scls) if _SR else ssubs),
                    "subjects": (_SR.canon_list(matched, scls) if _SR else matched),
                    "has_photo": bool(sp.photo_b64)})
    out.sort(key=lambda x: (x["name"] or "").lower())
    return {"total": len(out), "students": out}

# ===== TEACHER -> ADMIN MESSAGE =====
@router.post("/message-admin")
def teacher_message_admin(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    title = (payload.get("title") or "").strip()
    message = (payload.get("message") or "").strip()
    if not title or not message:
        raise HTTPException(status_code=400, detail="Title and message are required")
    admins = db.query(User).filter(User.role == "admin").all()
    sender = current_user.name
    for a in admins:
        notify(db, a.id, f"\u2709\ufe0f {sender}: {title}", message, "teacher_to_admin")
    db.commit()
    return {"message": "Message sent to the admin"}

# ===== TEACHER ACCOUNTABILITY: classes with status, mark-complete, compliance =====
def _class_status(e):
    """Upcoming | Live | Completed | Missed based on date/time + completed flag."""
    if getattr(e, "completed", False):
        return "Completed"
    today = date.today()
    if e.entry_date is None:
        return "Upcoming"
    if e.entry_date < today:
        return "Missed"
    if e.entry_date > today:
        return "Upcoming"
    return "Pending"  # today, not yet completed

@router.get("/my-classes")
def teacher_my_classes(scope: str = "all", db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry
    from sqlalchemy import or_
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    if not subs:
        return []
    q = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(subs),
        or_(TimetableEntry.status == None, TimetableEntry.status != "pending"))
    if scope == "today":
        q = q.filter(TimetableEntry.entry_date == date.today())
    es = q.order_by(TimetableEntry.entry_date, TimetableEntry.time_text).all()
    out = []
    for e in es:
        d = _serialize_tt(e); d["live_status"] = _class_status(e)
        out.append(d)
    return out

@router.post("/class/{entry_id}/complete")
def teacher_complete_class(entry_id: int, payload: dict, background_tasks: BackgroundTasks = None,
                           db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry
    tp = get_teacher_profile(current_user, db)
    e = db.query(TimetableEntry).filter(TimetableEntry.id == entry_id).first()
    if not e or e.subject not in (tp.subjects or []):
        raise HTTPException(status_code=404, detail="Class not found")
    e.completed = True
    e.completed_at = datetime.now()
    e.topic_covered = (payload.get("topic_covered") or e.chapter or "").strip() or None
    e.start_time = (payload.get("start_time") or "").strip() or None
    e.end_time = (payload.get("end_time") or "").strip() or None
    e.homework = (payload.get("homework") or "").strip() or None
    e.dpp_given = bool(payload.get("dpp_given"))
    e.remarks = (payload.get("remarks") or "").strip() or None
    db.commit()
    _maybe_warn_late(db, tp, e)
    if background_tasks is not None:
        background_tasks.add_task(_notify_class_done, e.subject, e.chapter or "",
                                  e.part or "", current_user.name or "Your teacher",
                                  bool(e.dpp_given))
    return {"message": "Class marked as completed."}


_LATE_WARN_THRESHOLD = 2      # more than this many late starts in a month -> remind


def _maybe_warn_late(db, tp, entry):
    """If this teacher has now started more than _LATE_WARN_THRESHOLD classes late
    this month, remind them once. Repeat delays hurt the institute's reputation and
    unsettle students, so the reminder is sent automatically."""
    try:
        from models import TimetableEntry
        d = _delay_of(entry)
        if _delay_band(d) not in ("minor", "late"):
            return
        today = date.today()
        month_start = date(today.year, today.month, 1)
        rows = db.query(TimetableEntry).filter(
            TimetableEntry.teacher_id == tp.id,
            TimetableEntry.completed == True,
            TimetableEntry.entry_date >= month_start).all()
        late = sum(1 for r in rows if _delay_band(_delay_of(r)) in ("minor", "late"))
        if late <= _LATE_WARN_THRESHOLD or not tp.user:
            return
        title = "\u26a0\ufe0f Class Punctuality Reminder"
        # only remind once a month
        seen = db.query(Notification).filter(
            Notification.user_id == tp.user.id, Notification.title == title,
            Notification.created_at >= datetime(today.year, today.month, 1)).first()
        if seen:
            return
        msg = ("Your classes started late %d times this month.\n\n"
               "This affects MVS Foundation's reputation and makes the children anxious. "
               "It will also reflect in your monthly report.\n\n"
               "Please start your classes on time." % late)
        notify(db, tp.user.id, title, msg, "warning")
        db.commit()
    except Exception:
        db.rollback()


def _notify_class_done(subject, chapter, part, teacher_name, dpp_given):
    """Tell the subject's students the class report is up, and guide them to the
    class notes / DPP / verification so they know exactly what to do next."""
    from database import SessionLocal
    from models import StudentProfile
    db = SessionLocal()
    try:
        topic = " \u00b7 ".join([x for x in (chapter, part) if x])
        msg = ("%s completed the %s class%s.\n"
               "\u2022 Class notes: Materials \u2192 %s\n"
               "%s"
               "\u2022 Open the Time Table and tap 'Mark Done' to verify the lecture (earn XP)."
               % (teacher_name, subject, (" (" + topic + ")") if topic else "", subject,
                  ("\u2022 DPP: download it from the DPP Submit page, solve it and upload it back\n"
                   if dpp_given else "")))
        for sp in db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all():
            if subject in (sp.subjects or []) and sp.user:
                notify(db, sp.user.id, "\U0001F4DA %s class complete" % subject, msg, "class")
        db.commit()
    except Exception:
        db.rollback()
    finally:
        db.close()

@router.get("/compliance")
def teacher_compliance(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, Material
    from sqlalchemy import or_
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    subj_count = max(1, len(subs))
    today = date.today()
    month_start = date(today.year, today.month, 1)
    classes = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(subs),
        or_(TimetableEntry.status == None, TimetableEntry.status != "pending"),
        TimetableEntry.entry_type == "chapter").all() if subs else []
    due = [c for c in classes if c.entry_date and c.entry_date <= today]
    completed = [c for c in due if getattr(c, "completed", False)]
    mats = db.query(Material).options(defer(Material.content_b64)).filter(Material.subject.in_(subs)).all() if subs else []
    dpp_count = sum(1 for m in mats if m.material_type == "dpp")
    notes_count = sum(1 for m in mats if m.material_type == "notes")
    test_count = sum(1 for m in mats if m.material_type == "test")
    # component scores (0..1)
    cc = (len(completed) / len(due)) if due else 1.0
    dpp_s = min(1.0, dpp_count / subj_count)
    mat_s = min(1.0, notes_count / subj_count)
    test_s = min(1.0, test_count / subj_count)
    score = round(cc * 40 + dpp_s * 25 + mat_s * 20 + test_s * 15)
    band = "green" if score >= 81 else ("yellow" if score >= 61 else "red")
    return {
        "score": score, "band": band,
        "breakdown": {
            "class_completion": {"weight": 40, "pct": round(cc * 100), "got": round(cc * 40)},
            "dpp_upload": {"weight": 25, "pct": round(dpp_s * 100), "got": round(dpp_s * 25)},
            "study_material": {"weight": 20, "pct": round(mat_s * 100), "got": round(mat_s * 20)},
            "test_creation": {"weight": 15, "pct": round(test_s * 100), "got": round(test_s * 15)},
        },
        "stats": {
            "classes_due": len(due), "classes_completed": len(completed),
            "dpp_count": dpp_count, "notes_count": notes_count, "test_count": test_count,
            "classes_today": sum(1 for c in classes if c.entry_date == today),
            "completed_today": sum(1 for c in completed if c.entry_date == today),
            "pending_today": sum(1 for c in classes if c.entry_date == today and not getattr(c, "completed", False)),
            "missed": sum(1 for c in due if not getattr(c, "completed", False)),
            "subject_count": len(subs),
        }
    }

# ===== TEACHER: DOUBT STATS (pending, resolved, avg response time) =====
@router.get("/doubt-stats")
def teacher_doubt_stats(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Doubt, DoubtStatus
    tp = get_teacher_profile(current_user, db)
    ds = db.query(Doubt).filter(Doubt.teacher_id == tp.id).all()
    # v93: admin ko assigned doubts meri pending/responsibility se bahar
    ds = [d for d in ds if not getattr(d, "assigned_to_admin", False)]
    # v112: resolved doubt pe naya student follow-up bhi pending attention hai
    pending = sum(1 for d in ds if (d.status.value if hasattr(d.status, "value") else d.status) == "pending")
    pending += sum(1 for d in ds
                   if (d.status.value if hasattr(d.status, "value") else d.status) == "resolved"
                   and _doubt_needs_attention(db, d.id))
    resolved_list = [d for d in ds if (d.status.value if hasattr(d.status, "value") else d.status) == "resolved" and d.resolved_at and d.created_at]
    resolved = sum(1 for d in ds if (d.status.value if hasattr(d.status, "value") else d.status) == "resolved")
    avg_min = None
    if resolved_list:
        total = sum((d.resolved_at - d.created_at).total_seconds() for d in resolved_list)
        avg_min = round(total / len(resolved_list) / 60)
    return {"pending": pending, "resolved": resolved, "total": len(ds), "avg_response_minutes": avg_min}

# ===== TEACHER: PERFORMANCE (aggregates + recent activity + monthly) =====
@router.get("/performance")
def teacher_performance(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import TimetableEntry, Material
    from sqlalchemy import or_
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    classes = db.query(TimetableEntry).filter(TimetableEntry.subject.in_(subs),
        or_(TimetableEntry.status == None, TimetableEntry.status != "pending"),
        TimetableEntry.entry_type == "chapter").all() if subs else []
    completed = [c for c in classes if getattr(c, "completed", False)]
    mats = db.query(Material).options(defer(Material.content_b64)).filter(Material.subject.in_(subs)).all() if subs else []
    dpp_count = sum(1 for m in mats if m.material_type == "dpp")
    notes_count = sum(1 for m in mats if m.material_type in ("notes", "other"))
    test_count = sum(1 for m in mats if m.material_type == "test")
    # monthly completed (last 6 months)
    from collections import OrderedDict
    today = date.today()
    months = OrderedDict()
    for i in range(5, -1, -1):
        y = today.year; mo = today.month - i
        while mo <= 0:
            mo += 12; y -= 1
        months[f"{y}-{mo:02d}"] = 0
    for c in completed:
        if c.completed_at:
            key = f"{c.completed_at.year}-{c.completed_at.month:02d}"
            if key in months:
                months[key] += 1
    monthly = [{"month": k, "count": v} for k, v in months.items()]
    # recent activity (completions + uploads)
    acts = []
    for c in completed:
        if c.completed_at:
            acts.append({"type": "class", "text": f"Completed {c.subject} — {c.topic_covered or c.chapter or ''}", "at": c.completed_at})
    for m in mats:
        if m.created_at:
            acts.append({"type": m.material_type, "text": f"Uploaded {m.material_type.upper()}: {m.title or m.chapter or m.subject}", "at": m.created_at})
    acts.sort(key=lambda x: x["at"], reverse=True)
    recent = [{"type": a["type"], "text": a["text"], "at": str(a["at"])[:16]} for a in acts[:12]]
    return {
        "classes_assigned": len(classes), "classes_completed": len(completed),
        "dpp_uploaded": dpp_count, "materials_uploaded": notes_count, "tests_created": test_count,
        "monthly": monthly, "recent": recent
    }

# ===== TEACHER: MATERIAL ANALYTICS (views/downloads per material) =====
@router.get("/material-analytics")
def teacher_material_analytics(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material, MaterialView
    from sqlalchemy import func as _f
    tp = get_teacher_profile(current_user, db)
    subs = tp.subjects or []
    mats = db.query(Material).options(defer(Material.content_b64)).filter(Material.subject.in_(subs),
        Material.material_type.in_(["notes", "dpp", "test", "other"])).order_by(Material.created_at.desc()).all() if subs else []
    out = []
    for m in mats:
        viewed = db.query(_f.count(_f.distinct(MaterialView.student_id))).filter(MaterialView.material_id == m.id).scalar() or 0
        downloads = db.query(_f.count(MaterialView.id)).filter(MaterialView.material_id == m.id, MaterialView.action == "download").scalar() or 0
        out.append({
            "id": m.id, "type": m.material_type, "category": m.category,
            "title": m.title or m.chapter or m.subject, "subject": m.subject,
            "upload_date": str(m.created_at)[:10] if m.created_at else None,
            "students_viewed": viewed, "downloads": downloads,
            "approval_status": getattr(m, "approval_status", "approved") or "approved",
        })
    return out

# ===== TEACHER: STUDENT ENGAGEMENT =====
@router.get("/student-engagement")
def teacher_student_engagement(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import Material, MaterialView, StudentProfile
    from sqlalchemy import func as _f, or_
    tp = get_teacher_profile(current_user, db)
    _see = _teacher_sees_students(tp, db)
    subs = tp.subjects or []
    if not subs:
        return []
    # students who have any of the teacher's subjects
    students = []
    for sp in db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all():
        if set(sp.subjects or []) & set(subs):
            students.append(sp)
    # teacher material ids
    mat_ids = [m.id for m in db.query(Material).options(defer(Material.content_b64)).filter(Material.subject.in_(subs)).all()]
    out = []
    for sp in students:
        answers = db.query(Material).options(defer(Material.content_b64)).filter(Material.material_type == "answer", Material.student_id == sp.id).all()
        pids = [a.parent_id for a in answers if a.parent_id]
        ptypes = {}
        if pids:
            for pm in db.query(Material).options(defer(Material.content_b64)).filter(Material.id.in_(pids)).all():
                ptypes[pm.id] = pm.material_type
        dpp_done = sum(1 for a in answers if ptypes.get(a.parent_id) == "dpp")
        test_done = sum(1 for a in answers if ptypes.get(a.parent_id) == "test")
        downloads = db.query(_f.count(MaterialView.id)).filter(
            MaterialView.student_id == sp.id, MaterialView.action == "download",
            MaterialView.material_id.in_(mat_ids) if mat_ids else False).scalar() or 0
        last_act = db.query(MaterialView).filter(MaterialView.student_id == sp.id).order_by(MaterialView.created_at.desc()).first()
        out.append({
            "name": (sp.user.name if sp.user else "Student"),
            "phone": (sp.phone if _see else ""), "subjects": sp.subjects or [],
            "class_level": sp.class_level,
            "dpp_completed": dpp_done, "tests_completed": test_done,
            "material_downloads": downloads,
            "last_active": str(last_act.created_at)[:16] if last_act else None,
        })
    out.sort(key=lambda x: (x["material_downloads"] + x["dpp_completed"] + x["tests_completed"]), reverse=True)
    return out


# ===================== EXAM / TEST ENGINE (teacher) =====================

# ---- exam engine: lazy column migration + helpers (safe to call anywhere) ----
_EXAM_COLS_READY = False

def _ensure_exam_columns(db):
    """Add scheduled_at / attempted / skipped columns on first use (MySQL/Postgres/SQLite).
    Runs once per process; every ALTER is best-effort so existing databases upgrade themselves."""
    global _EXAM_COLS_READY
    if _EXAM_COLS_READY:
        return
    from sqlalchemy import text as _text
    stmts = [
        ("ALTER TABLE exams ADD COLUMN scheduled_at DATETIME NULL",
         "ALTER TABLE exams ADD COLUMN scheduled_at TIMESTAMP NULL"),
        ("ALTER TABLE exam_attempts ADD COLUMN attempted JSON NULL",
         "ALTER TABLE exam_attempts ADD COLUMN attempted TEXT NULL"),
        ("ALTER TABLE exam_attempts ADD COLUMN skipped JSON NULL",
         "ALTER TABLE exam_attempts ADD COLUMN skipped TEXT NULL"),
        ("ALTER TABLE exam_questions ADD COLUMN alt_image_b64 LONGTEXT NULL",
         "ALTER TABLE exam_questions ADD COLUMN alt_image_b64 TEXT NULL"),
        ("ALTER TABLE exams ADD COLUMN class_name VARCHAR(50) NULL",
         "ALTER TABLE exams ADD COLUMN class_name TEXT NULL"),
    ]
    for group in stmts:
        for s in group:
            try:
                db.execute(_text(s))
                db.commit()
                break
            except Exception:
                db.rollback()
    _EXAM_COLS_READY = True


def _exam_parse_dt(v):
    """Parse an ISO-ish datetime from the portal; returns None on failure."""
    if not v:
        return None
    if isinstance(v, datetime):
        return v
    s = str(v).strip()
    try:
        dt = datetime.fromisoformat(s.replace("Z", "+00:00"))
        if dt.tzinfo is not None:
            dt = dt.astimezone().replace(tzinfo=None)
        return dt
    except Exception:
        pass
    for fmt in ("%Y-%m-%dT%H:%M", "%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
        try:
            return datetime.strptime(s[:19], fmt)
        except Exception:
            continue
    return None

def _parse_dur(v):
    """Test duration ko normalize: 0 / blank / invalid -> None (no time limit);
    positive int -> wahi. None matlab student jab chahe kar sakta hai, koi countdown nahi."""
    try:
        n = int(v)
    except Exception:
        return None
    return n if n > 0 else None


@router.post("/exam")
def create_exam(payload: dict = Body(...), background_tasks: BackgroundTasks = None, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    qs = payload.get("questions") or []
    if not payload.get("title") or not qs:
        raise HTTPException(400, "Title and at least one question are required")
    ttype = payload.get("test_type", "subjective")
    total = sum(int(q.get("max_marks", 1) or 1) for q in qs)
    _subj_in = payload.get("subject", "")
    if _SR is not None and _subj_in:
        _subj_in = _SR.canon_display(_subj_in, payload.get("class_name"))
    ex = Exam(teacher_id=tp.id, teacher_name=current_user.name,
              subject=_subj_in, title=payload["title"],
              chapter=payload.get("chapter"), test_type=ttype,
              class_name=(payload.get("class_name") or "").strip(),
              medium=payload.get("medium", "English"),
              total_marks=total, duration_min=_parse_dur(payload.get("duration_min")),
              scheduled_at=_exam_parse_dt(payload.get("scheduled_at")))
    db.add(ex); db.flush()
    for i, q in enumerate(qs, start=1):
        co = q.get("correct_option")
        opts_hi = q.get("options_hi") if ttype == "mcq" else None
        db.add(ExamQuestion(exam_id=ex.id, q_no=i,
               question_text=q.get("question_text", ""),
               max_marks=int(q.get("max_marks", 1) or 1),
               model_answer=q.get("model_answer"),
               options=q.get("options") if ttype == "mcq" else None,
               correct_option=(str(co) if co not in (None, "") else None),
               image_b64=_r2img(q.get("image_b64")),
               question_text_hi=(q.get("question_text_hi") or None),
               model_answer_hi=(q.get("model_answer_hi") or None),
               options_hi=(opts_hi if opts_hi else None),
               model_answer_image=_r2img(q.get("model_answer_image")),
               alt_image_b64=_r2img(q.get("alt_image_b64")),
               explanation=(q.get("explanation") or None),
               explanation_hi=(q.get("explanation_hi") or None)))
    db.commit()
    # Bilingual Hindi is now filled on-demand by the portal (free). Paid Gemini
    # auto-translation is disabled to avoid API costs. (Function kept for manual use.)
    # if (ex.medium or "").lower().startswith("bi") and background_tasks is not None:
    #     background_tasks.add_task(_bg_translate_exam, ex.id)
    return {"id": ex.id, "total_marks": total, "questions": len(qs),
            "test_type": ttype, "medium": ex.medium,
            "scheduled_at": ex.scheduled_at.isoformat() if getattr(ex, "scheduled_at", None) else None}


def _bg_translate_exam(exam_id):
    """Fill in any missing Hindi fields for a bilingual test using Gemini.
    Runs after the response so test creation stays fast. Only fills blanks."""
    from database import SessionLocal
    db = SessionLocal()
    try:
        ex = db.query(Exam).filter(Exam.id == exam_id).first()
        if not ex:
            return
        qs = db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).order_by(ExamQuestion.q_no).all()
        for q in qs:
            need_q = not (q.question_text_hi or "").strip()
            need_a = (ex.test_type != "mcq") and not (q.model_answer_hi or "").strip()
            need_o = (ex.test_type == "mcq") and not q.options_hi
            if not (need_q or need_a or need_o):
                continue
            tr = grading.translate_question_to_hindi(
                q.question_text or "", q.model_answer or "",
                (q.options or []) if ex.test_type == "mcq" else None, ex.subject or "")
            if not tr:
                continue
            if need_q and tr.get("question"):
                q.question_text_hi = tr["question"]
            if need_a and tr.get("answer"):
                q.model_answer_hi = tr["answer"]
            if need_o and tr.get("options") and len(tr["options"]) == len(q.options or []):
                q.options_hi = tr["options"]
        db.commit()
    except Exception:
        db.rollback()
    finally:
        db.close()


@router.get("/exam/{exam_id}/pdf")
def teacher_exam_pdf(exam_id: int, medium: str = "english",
                     db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Download the full question+answer paper as a PDF in English or Hindi medium."""
    _ensure_exam_columns(db)
    import exam_pdf
    from fastapi import Response
    tp = get_teacher_profile(current_user, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Test not found")
    qs = db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).order_by(ExamQuestion.q_no).all()
    med = ("both" if str(medium).lower().startswith("bo") else ("hindi" if str(medium).lower().startswith("hi") else "english"))
    try:
        data = exam_pdf.build_exam_pdf(ex, qs, med)
    except Exception as e:
        raise HTTPException(500, "Could not generate the PDF. The server needs fpdf2, "
                                 "uharfbuzz and the Devanagari font. (%s)" % e)
    safe = (ex.title or "test").replace('"', "").replace("/", "-").strip()[:60] or "test"
    return Response(content=data, media_type="application/pdf",
                    headers={"Content-Disposition": 'attachment; filename="%s-%s.pdf"' % (safe, med)})

@router.get("/exams")
def list_exams(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    _ensure_exam_columns(db)
    from models import ExamView
    tp = get_teacher_profile(current_user, db)
    rows = db.query(Exam).filter(Exam.teacher_id == tp.id, Exam.is_active == True).order_by(Exam.created_at.desc()).all()
    ids = [e.id for e in rows]
    views, downloads = {}, {}
    if ids:
        for v in db.query(ExamView).filter(ExamView.exam_id.in_(ids)).all():
            (downloads if v.action == "download" else views).setdefault(v.exam_id, set()).add(v.student_id)
        # a student who attempted the test has obviously seen it - count them as a
        # viewer too, so tests taken before view-tracking existed still show up
        for a in db.query(ExamAttempt).filter(ExamAttempt.exam_id.in_(ids)).all():
            views.setdefault(a.exam_id, set()).add(a.student_id)
    out = []
    for e in rows:
        nq = db.query(ExamQuestion).filter(ExamQuestion.exam_id == e.id).count()
        na = db.query(ExamAttempt).filter(ExamAttempt.exam_id == e.id).count()
        ng = db.query(ExamAttempt).filter(ExamAttempt.exam_id == e.id, ExamAttempt.status == "graded").count()
        out.append({"id": e.id, "title": e.title, "subject": e.subject, "chapter": e.chapter,
                    "class_name": getattr(e, "class_name", "") or "",
                    "test_type": e.test_type, "total_marks": e.total_marks, "duration_min": e.duration_min,
                    "medium": e.medium, "questions": nq, "attempts": na, "graded": ng,
                    "views": len(views.get(e.id, ())), "downloads": len(downloads.get(e.id, ())),
                    "scheduled_at": e.scheduled_at.isoformat() if getattr(e, "scheduled_at", None) else None,
                    "created_at": e.created_at.isoformat() if e.created_at else None})
    return out


@router.get("/exam/{exam_id}/audience")
def exam_audience(exam_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Which students opened / downloaded this test."""
    _ensure_exam_columns(db)
    from models import ExamView, StudentProfile
    tp = get_teacher_profile(current_user, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Test not found")
    rows = db.query(ExamView).filter(ExamView.exam_id == exam_id).all()
    attempts = db.query(ExamAttempt).filter(ExamAttempt.exam_id == exam_id).all()
    sids = list({r.student_id for r in rows} | {a.student_id for a in attempts})
    smap = {}
    if sids:
        for sp in db.query(StudentProfile).filter(StudentProfile.id.in_(sids)).all():
            smap[sp.id] = (sp.user.name if sp.user else ("Student #%d" % sp.id))
    viewers, downloaders, seen = [], [], set()
    for r in rows:
        entry = {"student_id": r.student_id, "name": smap.get(r.student_id, "Student"),
                 "at": str(r.created_at)[:16]}
        if r.action == "download":
            downloaders.append(entry)
        else:
            seen.add(r.student_id)
            viewers.append(entry)
    # attempted => viewed (covers tests taken before view-tracking existed)
    for a in attempts:
        if a.student_id in seen:
            continue
        seen.add(a.student_id)
        viewers.append({"student_id": a.student_id, "name": smap.get(a.student_id, "Student"),
                        "at": str(a.submitted_at)[:16] if getattr(a, "submitted_at", None) else ""})
    return {"material": {"id": ex.id, "title": ex.title, "type": "test",
                         "subject": ex.subject, "chapter": ex.chapter, "part": None},
            "viewers": viewers, "downloaders": downloaders}

@router.post("/exam/{exam_id}/reset-attempts")
def exam_reset_attempts(exam_id: int, payload: dict = Body(default={}),
                        db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Is test ke attempts clear karo taaki students dobara de sakein. `only_broken=True`
    ho to sirf wo MCQ attempts hatao jinke answers empty the (purane bug ke shikaar)."""
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    _see = _teacher_sees_students(tp, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Exam not found")
    only_broken = bool((payload or {}).get("only_broken"))
    atts = db.query(ExamAttempt).filter(ExamAttempt.exam_id == exam_id).all()
    removed = 0
    for a in atts:
        if only_broken:
            ma = a.mcq_answers or {}
            has_ans = any(v not in (None, "", []) for v in ma.values())
            if has_ans:
                continue
        db.query(ExamResult).filter(ExamResult.attempt_id == a.id).delete(synchronize_session=False)
        db.delete(a)
        removed += 1
    db.commit()
    return {"ok": True, "removed": removed}


@router.get("/exam/{exam_id}/attempts")
def exam_attempts(exam_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import StudentProfile
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    _see = _teacher_sees_students(tp, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Exam not found")
    atts = db.query(ExamAttempt).filter(ExamAttempt.exam_id == exam_id).order_by(ExamAttempt.submitted_at.desc()).all()
    out = []
    for a in atts:
        _atl = getattr(a, "attempted", None)
        _skl = getattr(a, "skipped", None)
        _sp = db.query(StudentProfile).filter(StudentProfile.id == a.student_id).first()
        _su = _sp.user if _sp else None
        out.append({"attempt_id": a.id, "student_id": a.student_id, "student_name": a.student_name,
            "status": a.status, "total_awarded": a.total_awarded, "verdict": a.verdict,
            "has_answer": bool(a.answer_image_b64),
            "feedback": a.overall_feedback,
            "results": [{"q_no": rr.q_no, "marks": rr.marks_awarded,
                         "max": rr.max_marks, "remark": rr.remark or ""}
                        for rr in db.query(ExamResult)
                                    .filter(ExamResult.attempt_id == a.id)
                                    .order_by(ExamResult.q_no).all()],
            "phone": ((_sp.phone if _sp else None) if _see else None),
            "student_code": (_su.user_id if _su else None),
            "batch": (_sp.batch_name if _sp else None),
            "class_level": (_sp.class_level if _sp else None),
            "medium": (_sp.medium if _sp else None),
            "email": (_sp.email if _sp else None),
            "subjects": ((_sp.subjects or []) if _sp else []),
            "attempted_count": (len(_atl) if isinstance(_atl, list) else None),
            "skipped_count": (len(_skl) if isinstance(_skl, list) else None),
            "submitted_at": a.submitted_at.isoformat() if a.submitted_at else None})
    qrows = db.query(ExamQuestion).filter(ExamQuestion.exam_id == exam_id).order_by(ExamQuestion.q_no).all()
    questions = [{"q_no": q.q_no, "question_text": q.question_text, "max_marks": q.max_marks,
                  "model_answer": q.model_answer, "options": q.options,
                  "correct_option": q.correct_option, "image_b64": q.image_b64,
                  "question_text_hi": q.question_text_hi, "model_answer_hi": q.model_answer_hi,
                  "options_hi": q.options_hi, "model_answer_image": q.model_answer_image,
                  "alt_image_b64": getattr(q, "alt_image_b64", None),
                  "explanation": q.explanation, "explanation_hi": q.explanation_hi} for q in qrows]
    return {"exam": {"id": ex.id, "title": ex.title, "total_marks": ex.total_marks,
                     "test_type": ex.test_type, "subject": ex.subject, "chapter": ex.chapter,
                     "medium": ex.medium, "duration_min": ex.duration_min,
                     "scheduled_at": ex.scheduled_at.isoformat() if getattr(ex, "scheduled_at", None) else None},
            "questions": questions, "attempts": out}


def _exam_verdict_t(aw, tot):
    if not tot:
        return "Good"
    p = aw / tot * 100
    return "Excellent" if p >= 80 else ("Good" if p >= 50 else "Needs Improvement")

def _notify_exam_result_t(db, att, ex):
    """Notify the student that their test result is ready."""
    try:
        from models import StudentProfile
        sp = db.query(StudentProfile).filter(StudentProfile.id == att.student_id).first()
        if sp and sp.user_id:
            try:
                sc = "%g" % float(att.total_awarded)
            except Exception:
                sc = str(att.total_awarded)
            db.add(Notification(
                user_id=sp.user_id,
                title="Result ready: %s" % (ex.title or "Test"),
                message="Your test has been checked. You scored %s/%s. Tap to view your result and download your answer sheet." % (sc, ex.total_marks),
                notif_type="exam_result"))
    except Exception:
        pass


@router.post("/attempt/{attempt_id}/grade")
def grade_attempt_now(attempt_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    att = db.query(ExamAttempt).filter(ExamAttempt.id == attempt_id).first()
    if not att:
        raise HTTPException(404, "Attempt not found")
    ex = db.query(Exam).filter(Exam.id == att.exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(403, "Not your test")
    if att.status == "graded":
        return {"status": "graded", "message": "Already graded"}
    qs = db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).order_by(ExamQuestion.q_no).all()
    teacher = ex.teacher_name or "your teacher"
    if ex.test_type == "mcq":
        results, total = grading.grade_mcq(qs, att.mcq_answers or {})
        feedback, verdict = "", _exam_verdict_t(total, ex.total_marks)
    else:
        results, total, feedback, verdict = grading.grade_subjective(qs, att.answer_image_b64 or "", "image/jpeg")
        if results is None:
            raise HTTPException(400, "AI grading failed: " + (feedback or "unknown error") + " -- you can use Grade Manually instead.")
        verdict = verdict or _exam_verdict_t(total, ex.total_marks)
    db.query(ExamResult).filter(ExamResult.attempt_id == att.id).delete()
    for r in results:
        db.add(ExamResult(attempt_id=att.id, q_no=r["q_no"], marks_awarded=r["marks"],
               max_marks=r["max"], remark=r.get("remark", "")))
    att.total_awarded = total
    att.status = "graded"
    att.graded_at = datetime.utcnow()
    att.verdict = verdict
    att.overall_feedback = feedback or ("Graded by teacher. \u2014 %s" % teacher)
    _notify_exam_result_t(db, att, ex)
    db.commit()
    return {"status": "graded", "total_awarded": total, "verdict": verdict}


# ===================== AI AUTO-MAGIC ENDPOINTS (Phase 2) =====================
@router.get("/ai-format-config")
def teacher_ai_format_config(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """v122: teacher exam-builder ke liye AI formatting config (read-only).
    Admin Subjects screen se set hota hai — {'all': bool, 'subjects': [names]}."""
    from models import AppSetting
    import json as _json
    cfg = {"all": False, "subjects": []}
    row = db.query(AppSetting).filter(AppSetting.key == "ai_format_cfg").first()
    if row and row.value:
        try:
            data = _json.loads(row.value)
            cfg["all"] = bool(data.get("all"))
            subs = data.get("subjects") or []
            if isinstance(subs, list):
                cfg["subjects"] = sorted({str(s).strip()[:80] for s in subs if str(s).strip()})
        except Exception:
            pass
    cfg["ai_available"] = bool((_os.environ.get("GEMINI_API_KEY") or "").strip())
    return cfg

@router.post("/ocr-question")
def ocr_question(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    img = payload.get("image_b64") or ""
    if not img:
        raise HTTPException(400, "No image provided")
    res = grading.ocr_extract_question(img, payload.get("test_type", "subjective"),
                                       payload.get("mime_type", "image/jpeg"))
    if res is None:
        raise HTTPException(503, "AI could not read the image. Check GEMINI_API_KEY or try a clearer screenshot.")
    return res

@router.post("/format-text")
def format_text(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    out = grading.format_text_latex(payload.get("text", ""))
    if out is None:
        raise HTTPException(503, "AI formatting is unavailable. Check GEMINI_API_KEY.")
    return {"text": out}

@router.post("/format-text-ai")
def format_text_ai(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """v117: direct Gemini formatting endpoint with current model names.

    The legacy /format-text route goes through grading.py, whose pinned model
    name may be retired by Google. This endpoint calls the Gemini REST API
    directly with a fallback model chain and the same GEMINI_API_KEY, so AI
    formatting keeps working without touching grading.py.
    """
    import os
    import urllib.request
    import urllib.error

    text = str((payload or {}).get("text") or "").strip()
    if not text:
        raise HTTPException(400, "No text provided")
    if len(text) > 8000:
        raise HTTPException(400, "Text too long (max 8000 characters)")

    api_key = (os.environ.get("GEMINI_API_KEY") or "").strip()
    if not api_key:
        raise HTTPException(503, "GEMINI_API_KEY is not set on the server")

    prompt = (
        "You are a formatting engine for a school exam portal covering ALL subjects "
        "(Physics, Chemistry, Biology, Maths, Social Science, English, Accountancy, etc.). "
        "Rewrite the teacher's raw text as clean, display-ready content.\n"
        "RULES:\n"
        "1. Convert every mathematical expression into LaTeX wrapped in single $...$ (use $$...$$ only for a full-line display equation).\n"
        "2. Fractions use \\frac{num}{den}; roots use \\sqrt{...}; powers use ^{...}; subscripts use _{...}; integrals use \\int.\n"
        "3. Matrices and determinants use \\begin{bmatrix} ... \\end{bmatrix} (pmatrix/vmatrix if the source implies them).\n"
        "4. Units are upright, e.g. $\\text{cm}^{3}$ or $\\text{m}^{2}$.\n"
        "5. If the pieces of one formula are scattered across separate lines (integral sign, denominator, numerator), reassemble them into ONE correct LaTeX expression in the intended order.\n"
        "6. Chemistry: use \\\\ce{...} (mhchem) for formulae/equations, e.g. $\\\\ce{2H2 + O2 -> 2H2O}$; keep states (s)(l)(g)(aq).\n"
        "7. TABLES (Match the following / Differentiate between / Compare / any 2-3 column table): output ONE block as $\\\\begin{array}{l|l}\\\\hline \\\\colorbox{#dce8f8}{\\\\text{\\\\textbf{Header-1}}} & \\\\colorbox{#dce8f8}{\\\\text{\\\\textbf{Header-2}}} \\\\\\\\ \\\\hline cell & cell \\\\\\\\ \\\\hline ... \\\\end{array}$ — one column per group, one row per item. Do NOT add an empty spacer column.\n"
        "8. STRUCTURED ANSWERS (several named items/types/features, e.g. the five kingdoms, differences, points): put EACH item on its OWN line as a bullet starting with '- ' and bold the key term, e.g. '- **Monera:** Prokaryotic, unicellular ...'.\n"
        "9. Put 'Step 1:', 'Step 2:', 'Formula:', 'Substitute:', 'Answer:', 'Therefore', '(i)', '(ii)' etc. on their own lines.\n"
        "10. NEVER change the meaning, numbers, language (keep Hindi as Hindi, English as English), or option labels (A)(B)(C)(D). Only fix structure and formatting.\n"
        "11. Preserve existing markdown (**bold**, __underline__) and standalone OR lines.\n"
        "12. Output ONLY the rewritten text. No preamble, no code fences, no explanation.\n\n"
        "RAW TEXT:\n"
        + text
    )
    body = {
        "contents": [{"parts": [{"text": prompt}]}],
        # v118: temperature/top_p/top_k are deprecated on Gemini 3.x models — omit them
        "generationConfig": {"maxOutputTokens": 4096},
    }
    # v118: current model chain (gemini-2.0-flash was shut down by Google in Jun 2026).
    # Busy/quota (429) or retired (404) -> automatically tries the next model.
    # Ops override: set GEMINI_MODELS="model-a,model-b" to pin your own chain.
    models_env = (os.environ.get("GEMINI_MODELS") or "").strip()
    if models_env:
        models = [m.strip() for m in models_env.split(",") if m.strip()]
    else:
        models = ["gemini-3.6-flash", "gemini-3.5-flash", "gemini-2.5-flash", "gemini-3.5-flash-lite"]
    last_err = "no response"
    for model in models:
        url = "https://generativelanguage.googleapis.com/v1beta/models/{}:generateContent?key={}".format(model, api_key)
        req = urllib.request.Request(url, data=json.dumps(body).encode("utf-8"),
                                     headers={"Content-Type": "application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = json.loads(resp.read().decode("utf-8", "ignore"))
            parts = ((data.get("candidates") or [{}])[0].get("content") or {}).get("parts") or []
            out = "".join(str(p.get("text") or "") for p in parts).strip()
            out = re.sub(r"^```[a-zA-Z]*\s*|\s*```$", "", out).strip()
            if not out:
                last_err = "empty response"
                continue
            return {"text": out, "model": model}
        except urllib.error.HTTPError as e:
            detail = ""
            try:
                detail = e.read().decode("utf-8", "ignore")[:400]
            except Exception:
                pass
            if e.code == 403 or (e.code == 400 and "API key" in detail):
                raise HTTPException(503, "GEMINI_API_KEY is invalid or blocked by Google")
            last_err = "{} HTTP {}".format(model, e.code)
            continue  # 404/429 etc -> try the next model
        except Exception as e:
            last_err = "{} {}".format(model, type(e).__name__)
            continue
    raise HTTPException(503, "AI formatting failed ({})".format(last_err))


@router.post("/ocr-question-ai")
def ocr_question_ai(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """v163: screenshot -> Gemini VISION se EXACT question text (clean LaTeX). PYQ
    integral/fraction jaise cases me copy-paste tootta hai; screenshot bulletproof hai.
    Direct Gemini call (grading.py bypass) + model fallback chain."""
    import os, urllib.request, urllib.error
    img = str((payload or {}).get("image_b64") or "")
    mime = str((payload or {}).get("mime_type") or "image/jpeg")
    if not img:
        raise HTTPException(400, "No image provided")
    if img.startswith("data:"):
        c = img.find(",")
        if c >= 0:
            head = img[:c]
            img = img[c + 1:]
            m = re.search(r"data:([^;]+)", head)
            if m:
                mime = m.group(1)
    api_key = (os.environ.get("GEMINI_API_KEY") or "").strip()
    if not api_key:
        raise HTTPException(503, "GEMINI_API_KEY is not set on the server")
    prompt = (
        "You are reading a question from a screenshot for a school exam portal. "
        "Output the EXACT question shown, with EVERY mathematical expression in clean LaTeX wrapped in $...$.\n"
        "RULES:\n"
        "1. Reproduce fractions and roots EXACTLY as displayed. A fraction stays a fraction: e.g. the integral of 1 over root(16 - x^2) is $\\\\int \\\\frac{dx}{\\\\sqrt{16-x^{2}}}$ (NEVER flatten it to $\\\\int(16-x^2)dx$).\n"
        "2. \\\\frac{num}{den} for fractions, \\\\sqrt{...} / \\\\sqrt[n]{...} for roots, ^{...} powers, _{...} subscripts, \\\\int / \\\\sum / \\\\lim, \\\\ce{...} for chemistry.\n"
        "3. Keep the exact numbers, symbols, language (Hindi stays Hindi), and option labels (A)(B)(C)(D) — put each option on its own line.\n"
        "4. If it is a match-the-following / difference table, output it as $\\\\begin{array}{l|l}\\\\hline Header-1 & Header-2 \\\\\\\\ \\\\hline a & b \\\\\\\\ \\\\hline \\\\end{array}$.\n"
        "5. Output ONLY the question text — no preamble, no explanation, no code fences."
    )
    body = {
        "contents": [{"parts": [{"text": prompt},
                                 {"inline_data": {"mime_type": mime, "data": img}}]}],
        "generationConfig": {"maxOutputTokens": 4096},
    }
    models_env = (os.environ.get("GEMINI_MODELS") or "").strip()
    models = [m.strip() for m in models_env.split(",") if m.strip()] if models_env else \
        ["gemini-3.6-flash", "gemini-3.5-flash", "gemini-2.5-flash", "gemini-3.5-flash-lite"]
    last_err = "no response"
    for model in models:
        url = "https://generativelanguage.googleapis.com/v1beta/models/{}:generateContent?key={}".format(model, api_key)
        req = urllib.request.Request(url, data=json.dumps(body).encode("utf-8"),
                                     headers={"Content-Type": "application/json"}, method="POST")
        try:
            with urllib.request.urlopen(req, timeout=40) as resp:
                data = json.loads(resp.read().decode("utf-8", "ignore"))
            parts = ((data.get("candidates") or [{}])[0].get("content") or {}).get("parts") or []
            out = "".join(str(p.get("text") or "") for p in parts).strip()
            out = re.sub(r"^```[a-zA-Z]*\s*|\s*```$", "", out).strip()
            if not out:
                last_err = "empty response"
                continue
            return {"text": out, "model": model}
        except urllib.error.HTTPError as e:
            detail = ""
            try:
                detail = e.read().decode("utf-8", "ignore")[:400]
            except Exception:
                pass
            if e.code == 403 or (e.code == 400 and "API key" in detail):
                raise HTTPException(503, "GEMINI_API_KEY is invalid or blocked by Google")
            last_err = "{} HTTP {}".format(model, e.code)
            continue
        except Exception as e:
            last_err = "{} {}".format(model, type(e).__name__)
            continue
    raise HTTPException(503, "AI could not read the screenshot ({})".format(last_err))

@router.post("/parse-exam-docx")
async def parse_exam_docx(file: UploadFile = File(...), test_type: str = Form("subjective"),
                          db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    import io
    try:
        from docx import Document
    except Exception:
        raise HTTPException(503, "Word parsing is not enabled on the server (add python-docx to requirements.txt).")
    data = await file.read()
    try:
        doc = Document(io.BytesIO(data))
    except Exception:
        raise HTTPException(400, "Could not open the Word file. Please upload a valid .docx file.")
    full = "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
    # also pull text from tables
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                full += "\n" + " | ".join(cells)
    if not full.strip():
        raise HTTPException(400, "The Word file appears to be empty.")
    qs = grading.local_structure_questions(full, test_type)
    if qs is None:
        raise HTTPException(400, grading.LAST_ERROR or "Could not read questions from the document.")
    return {"questions": qs, "count": len(qs)}


@router.post("/parse-exam-pdf")
async def parse_exam_pdf(file: UploadFile = File(...), test_type: str = Form("subjective"),
                         db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Extract questions + answers from an uploaded PDF question paper."""
    import io
    data = await file.read()
    full = ""
    # try PyMuPDF first, then pdfplumber as a fallback
    try:
        import fitz
        doc = fitz.open(stream=data, filetype="pdf")
        full = "\n".join(page.get_text() for page in doc)
        doc.close()
    except Exception:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(data)) as pdf:
                full = "\n".join((pg.extract_text() or "") for pg in pdf.pages)
        except Exception:
            raise HTTPException(503, "PDF parsing is not enabled on the server (add pymupdf to requirements.txt).")
    if not full.strip():
        raise HTTPException(400, "Could not read any text from this PDF. If it is a scanned image, please use a text PDF or the screenshot auto-fill.")
    qs = grading.local_structure_questions(full, test_type)
    if qs is None:
        raise HTTPException(400, grading.LAST_ERROR or "Could not read questions from the PDF.")
    return {"questions": qs, "count": len(qs), "note": grading.LAST_ERROR or None}


@router.get("/attempt/{attempt_id}/answer")
def attempt_answer_image(attempt_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    _ensure_exam_columns(db)
    import base64          # ROOT CAUSE: ye import missing tha -> NameError ->
                           # unhandled 500 (CORS headers ke bina) -> portal par
                           # "Failed to fetch". Isi wajah se sheet kabhi nahi khulti thi.
    from fastapi import Response
    tp = get_teacher_profile(current_user, db)
    att = db.query(ExamAttempt).filter(ExamAttempt.id == attempt_id).first()
    if not att:
        raise HTTPException(404, "Attempt not found")
    ex = db.query(Exam).filter(Exam.id == att.exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(403, "Not your test")
    if (att.status or "") == "grading":
        att.status = "marking"   # teacher opened the sheet -> "being checked"
        db.commit()
    if not att.answer_image_b64:
        raise HTTPException(404, "No answer sheet uploaded")
    # R2 URL ho to redirect (naye uploads normalize se URL hote hain; migration ke baad
    # purane bhi URL) — warna neeche base64 decode toot jaata.
    if str(att.answer_image_b64).startswith("http"):
        # migration ne sabko .jpg/image/jpeg bana diya tha, chahe student ne PDF upload kiya ho —
        # isi wajah se PDF sheets "damaged/broken" aati thi. sniff=True -> file ke ASAL magic
        # bytes se sahi content-type (PDF/JPEG/PNG) se serve karo (same-origin stream, no CORS).
        return __import__("r2_storage").proxy_response(att.answer_image_b64, "application/octet-stream", None, False, sniff=True)
    # Students upload a photo OR a PDF. Pehle hamesha image/jpeg bheja jaata tha
    # aur decode fail hone par unhandled 500 aata tha - browser use CORS ke bina
    # block kar deta tha, isliye portal par "Failed to fetch" dikhta tha.
    raw = att.answer_image_b64 or ""
    mime = "image/jpeg"
    if raw.startswith("data:") and "," in raw:
        header, raw = raw.split(",", 1)
        try:
            mime = header.split(":", 1)[1].split(";", 1)[0] or "image/jpeg"
        except Exception:
            mime = "image/jpeg"
    raw = "".join(raw.split())          # stray whitespace/newlines hatao
    raw += "=" * (-len(raw) % 4)        # padding theek karo
    try:
        data = base64.b64decode(raw)
    except Exception:
        raise HTTPException(400, "The uploaded answer sheet could not be read. Ask the student to upload it again.")
    if not data:
        raise HTTPException(404, "No answer sheet uploaded")
    ext = "pdf" if "pdf" in mime else ("png" if "png" in mime else "jpg")
    safe = "".join(c for c in (att.student_name or "student") if c.isalnum() or c in " -_").strip() or "student"
    return Response(content=data, media_type=mime,
                    headers={"Content-Disposition": 'inline; filename="answer-%s.%s"' % (safe, ext),
                             "Content-Length": str(len(data))})

@router.post("/attempt/{attempt_id}/grade-manual")
def grade_attempt_manual(attempt_id: int, payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    att = db.query(ExamAttempt).filter(ExamAttempt.id == attempt_id).first()
    if not att:
        raise HTTPException(404, "Attempt not found")
    ex = db.query(Exam).filter(Exam.id == att.exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(403, "Not your test")
    qmap = {q.q_no: q for q in db.query(ExamQuestion).filter(ExamQuestion.exam_id == ex.id).all()}
    results = payload.get("results") or []
    db.query(ExamResult).filter(ExamResult.attempt_id == att.id).delete()
    total = 0.0
    for r in results:
        try:
            qn = int(r.get("q_no"))
        except Exception:
            continue
        mx = qmap[qn].max_marks if qn in qmap else int(r.get("max", 1) or 1)
        try:
            mk = float(r.get("marks", 0) or 0)
        except Exception:
            mk = 0.0
        mk = max(0.0, min(mk, float(mx)))
        total += mk
        db.add(ExamResult(attempt_id=att.id, q_no=qn, marks_awarded=mk, max_marks=mx, remark=r.get("remark", "")))
    att.total_awarded = total
    att.status = "graded"
    att.graded_at = datetime.utcnow()
    att.verdict = payload.get("verdict") or _exam_verdict_t(total, ex.total_marks)
    fb = payload.get("feedback") or ""
    att.overall_feedback = fb if fb else ("Checked by %s." % (ex.teacher_name or "your teacher"))
    _notify_exam_result_t(db, att, ex)
    db.commit()
    return {"status": "graded", "total_awarded": total, "verdict": att.verdict}


@router.get("/ai-status")
def ai_status(current_user=Depends(get_teacher)):
    return grading.ai_status()


# ============================================================
#  SMART LECTURE VERIFICATION — TEACHER SIDE
# ============================================================
from models import Lecture, LectureQuestion, StudentProfile

_LQ_TYPES = {"mcq", "image_mcq", "numerical", "fill_blank", "true_false"}


@router.get("/timetable-entries-lite")
def teacher_tt_entries_lite(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Timetable entries a lecture report can optionally be linked to. Matches by
    the teacher's own entries AND by their subjects, so linking works even when
    an entry was created without a teacher_id."""
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry
    from sqlalchemy import or_
    subs = tp.subjects or []
    q = db.query(TimetableEntry)
    conds = [TimetableEntry.teacher_id == tp.id]
    if subs:
        conds.append(TimetableEntry.subject.in_(subs))
    es = q.filter(or_(*conds)).order_by(TimetableEntry.entry_date.desc()).limit(120).all()
    # skip test/event rows - only teachable chapter parts make sense to verify
    out = []
    for e in es:
        et = (getattr(e, "entry_type", "") or "").lower()
        if et in ("test", "exam", "event"):
            continue
        out.append({"id": e.id, "subject": e.subject, "chapter": e.chapter, "part": e.part,
                    "date": str(e.entry_date) if e.entry_date else None,
                    "class_name": e.class_name})
    return out


def _lec_cls5(x):
    """lectures.class_level is VARCHAR(5) (meant for '10' / '12'). Timetable
    class_name can be 'Class 10' (8 chars) — inserting it raw makes MySQL strict
    mode reject the whole class report (error 1406, "Data too long"). Normalize:
    first digit run wins ('Class 10' -> '10'); no digits -> raw, truncated;
    empty -> None. Result can never exceed 5 chars."""
    import re as _re
    raw = str(x or "").strip()
    if not raw:
        return None
    m = _re.search(r"\d+", raw)
    return (m.group(0) if m else raw)[:5]


@router.post("/lecture")
async def create_lecture(payload: dict = Body(...), background_tasks: BackgroundTasks = None,
                         db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Publish a lecture report (summary + class notes). v98: verification
    questions removed — students mark the lecture done directly, so an empty
    question set is fine."""
    tp = get_teacher_profile(current_user, db)
    subject = (payload.get("subject") or "").strip()
    if not subject:
        raise HTTPException(400, "Subject is required")
    qs = payload.get("questions") or []
    valid_qs = [q for q in qs if (q.get("question") or "").strip() and (q.get("qtype") in _LQ_TYPES)]

    # v123: class resolution — single-class subject pe hamesha official class;
    # warna payload ka class; warna linked timetable entry ka class_name.
    _cls_fixed = _subj_class_digits(db, subject)
    if _cls_fixed:
        _eff_cls = "Class " + _cls_fixed
    else:
        _eff_cls = (payload.get("class_level") or "").strip()
        if not _eff_cls and payload.get("timetable_entry_id"):
            try:
                from models import TimetableEntry as _TTE
                _te = db.query(_TTE).filter(_TTE.id == int(payload.get("timetable_entry_id"))).first()
                if _te:
                    _eff_cls = (_te.class_name or "").strip()
            except Exception:
                pass

    lec = Lecture(
        teacher_id=tp.id, teacher_name=(current_user.name or ""),
        subject=subject, class_level=_lec_cls5(_eff_cls),
        chapter=(payload.get("chapter") or None), part=(payload.get("part") or None),
        title=(payload.get("title") or (subject + " Lecture")),
        timetable_entry_id=(payload.get("timetable_entry_id") or None),
        summary=(payload.get("summary") or None), homework=(payload.get("homework") or None),
        pdf_b64=(payload.get("pdf_b64") or None), pdf_filename=(payload.get("pdf_filename") or None),
        dpp_b64=(payload.get("dpp_b64") or None), dpp_filename=(payload.get("dpp_filename") or None),
        is_active=True,
    )
    from datetime import date as _date
    ld = payload.get("lecture_date")
    if ld:
        try:
            lec.lecture_date = _date.fromisoformat(ld)
        except Exception:
            pass
    db.add(lec); db.flush()

    # Mirror the uploads into Materials so students find them under Study Material
    # and view/download analytics work exactly like any other material.
    from models import Material
    def _mk(kind, b64, fname):
        if not b64:
            return
        stored = __import__("r2_storage").normalize(b64, "materials", "application/pdf")
        db.add(Material(
            teacher_id=tp.id, teacher_name=(current_user.name or ""),
            subject=subject, class_name=(_eff_cls or None),
            chapter=(payload.get("chapter") or None), part=(payload.get("part") or None),
            material_type=kind, title=(lec.title or subject),
            filename=(fname or ("%s.pdf" % kind)), content_b64=stored))
    _mk("notes", payload.get("pdf_b64"), payload.get("pdf_filename"))
    _mk("dpp", payload.get("dpp_b64"), payload.get("dpp_filename"))
    db.flush()

    for q in valid_qs:
        db.add(LectureQuestion(
            lecture_id=lec.id, qtype=q.get("qtype"),
            question=(q.get("question") or ""), question_hi=(q.get("question_hi") or None),
            image_b64=_r2img(q.get("image_b64")),
            options=(q.get("options") or None), options_hi=(q.get("options_hi") or None),
            option_images=(q.get("option_images") or None),
            correct=str(q.get("correct") if q.get("correct") is not None else ""),
            tolerance=(float(q["tolerance"]) if q.get("tolerance") not in (None, "") else None),
        ))
    db.commit(); db.refresh(lec)

    # notify students of this subject in the background (fast response)
    if background_tasks is not None:
        background_tasks.add_task(_notify_lecture_students, subject, lec.title, current_user.name)
    return {"id": lec.id, "message": "Lecture report published"}


def _notify_lecture_students(subject, title, teacher_name):
    from database import SessionLocal
    db = SessionLocal()
    try:
        studs = db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all()
        for sp in studs:
            if subject in (sp.subjects or []) and sp.user:
                notify(db, sp.user.id, "\U0001F4DA New Lecture: %s" % subject,
                       "%s ne '%s' ka lecture report daala hai. Mark it done to verify." % (teacher_name, title),
                       "lecture")
        db.commit()
    except Exception:
        db.rollback()
    finally:
        db.close()


@router.get("/lectures")
def teacher_lectures(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import LectureVerification
    lecs = db.query(Lecture).options(defer(Lecture.pdf_b64), defer(Lecture.dpp_b64)).filter(Lecture.teacher_id == tp.id).order_by(Lecture.created_at.desc()).all()
    out = []
    for l in lecs:
        nq = db.query(LectureQuestion).filter(LectureQuestion.lecture_id == l.id).count()
        verified = db.query(LectureVerification).filter(
            LectureVerification.lecture_id == l.id, LectureVerification.status == "verified").count()
        attempted = db.query(LectureVerification).filter(LectureVerification.lecture_id == l.id).count()
        out.append({"id": l.id, "title": l.title, "subject": l.subject, "chapter": l.chapter,
                    "date": str(l.lecture_date) if l.lecture_date else str(l.created_at)[:10],
                    "questions": nq, "verified": verified, "attempted": attempted,
                    "has_pdf": bool(l.pdf_b64), "has_dpp": bool(l.dpp_b64), "is_active": l.is_active})
    return out


@router.delete("/lecture/{lecture_id}")
def delete_lecture(lecture_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    lec = db.query(Lecture).options(defer(Lecture.pdf_b64), defer(Lecture.dpp_b64)).filter(Lecture.id == lecture_id, Lecture.teacher_id == tp.id).first()
    if not lec:
        raise HTTPException(404, "Lecture not found")
    lec.is_active = False
    db.commit()
    return {"message": "Lecture removed"}


# ===================================================================== CLASS REPORTS
# Delay tracking + teaching-hours analytics, computed from the timetable's
# scheduled slot (time_text) vs what the teacher actually reported (start_time).

def _parse_hhmm(s):
    """'6:30 pm' / '18:30' / '6.30pm' -> minutes since midnight, or None."""
    if not s:
        return None
    m = re.search(r"(\d{1,2})\s*[:.]\s*(\d{2})\s*(am|pm)?", str(s), re.I)
    if not m:
        m2 = re.search(r"(\d{1,2})\s*(am|pm)", str(s), re.I)
        if not m2:
            return None
        h = int(m2.group(1)) % 12
        if m2.group(2).lower() == "pm":
            h += 12
        return h * 60
    h, mi = int(m.group(1)), int(m.group(2))
    ap = (m.group(3) or "").lower()
    if ap == "pm" and h < 12:
        h += 12
    if ap == "am" and h == 12:
        h = 0
    return h * 60 + mi


def _slot_start(time_text):
    """The scheduled slot may be a range ('6:30 pm - 7:30 pm'); take its start."""
    if not time_text:
        return None
    first = re.split(r"[-\u2013\u2014to]+", str(time_text))[0]
    return _parse_hhmm(first)


def _delay_of(e):
    """Minutes the class started late (negative = early). None when unknown."""
    sched = _slot_start(getattr(e, "time_text", None))
    actual = _parse_hhmm(getattr(e, "start_time", None))
    if sched is None or actual is None:
        return None
    return actual - sched


def _duration_of(e):
    a = _parse_hhmm(getattr(e, "start_time", None))
    b = _parse_hhmm(getattr(e, "end_time", None))
    if a is None or b is None:
        return 0
    d = b - a
    if d < 0:
        d += 24 * 60          # crossed midnight
    return d if 0 < d <= 6 * 60 else 0


def _delay_band(d):
    if d is None:
        return "unknown"
    if d <= 5:
        return "ontime"       # up to 5 min = on time
    if d <= 15:
        return "minor"
    return "late"


def _report_rows(db, subjects, teacher_map=None):
    from models import TimetableEntry
    q = db.query(TimetableEntry).filter(TimetableEntry.completed == True,
                                        TimetableEntry.entry_type == "chapter")
    if subjects is not None:
        if not subjects:
            return []
        q = q.filter(TimetableEntry.subject.in_(subjects))
    es = q.order_by(TimetableEntry.entry_date.desc()).limit(300).all()
    rows = []
    for e in es:
        d = _delay_of(e)
        rows.append({
            "id": e.id, "subject": e.subject, "chapter": e.chapter, "part": e.part,
            "date": str(e.entry_date) if e.entry_date else None,
            "scheduled": e.time_text, "start_time": e.start_time, "end_time": e.end_time,
            "delay_min": d, "delay_band": _delay_band(d),
            "duration_min": _duration_of(e),
            "topic_covered": e.topic_covered, "homework": e.homework,
            "dpp_given": bool(e.dpp_given), "remarks": e.remarks,
            "teacher_id": e.teacher_id,
            "teacher_name": (teacher_map or {}).get(e.teacher_id, ""),
        })
    return rows


def _report_summary(rows):
    today = date.today()
    week_start = today - timedelta(days=today.weekday())
    month_start = date(today.year, today.month, 1)
    by_subject = {}
    wk_min = mo_min = 0
    delays = [r["delay_min"] for r in rows if r["delay_min"] is not None]
    bands = {"ontime": 0, "minor": 0, "late": 0}
    for r in rows:
        bands[r["delay_band"]] = bands.get(r["delay_band"], 0) + (1 if r["delay_band"] in bands else 0)
        if not r["date"]:
            continue
        try:
            d = date.fromisoformat(r["date"])
        except Exception:
            continue
        s = by_subject.setdefault(r["subject"], {"subject": r["subject"], "classes": 0,
                                                 "week_min": 0, "month_min": 0, "total_min": 0})
        s["classes"] += 1
        s["total_min"] += r["duration_min"]
        if d >= week_start:
            s["week_min"] += r["duration_min"]; wk_min += r["duration_min"]
        if d >= month_start:
            s["month_min"] += r["duration_min"]; mo_min += r["duration_min"]
    return {
        "week_hours": round(wk_min / 60.0, 1),
        "month_hours": round(mo_min / 60.0, 1),
        "classes_done": len(rows),
        "avg_delay": (round(sum(delays) / len(delays)) if delays else None),
        "on_time_pct": (round(bands["ontime"] * 100 / len(delays)) if delays else None),
        "bands": bands,
        "by_subject": sorted(by_subject.values(), key=lambda x: -x["total_min"]),
    }


@router.get("/class-reports")
def teacher_class_reports(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """The teacher's own submitted class reports + delay/hours analytics."""
    tp = get_teacher_profile(current_user, db)
    rows = _report_rows(db, tp.subjects or [])
    return {"summary": _report_summary(rows), "rows": rows[:60]}


# ============================================================ MATERIAL ANALYTICS
# Subject -> chapter -> part view of everything uploaded, with view/download
# counts and the actual student lists behind those counts. Shared shape so the
# teacher portal and the admin portal render from the same renderer.

def _material_tree(db, subjects=None):
    from models import Material, MaterialView, StudentProfile
    q = db.query(Material).options(defer(Material.content_b64)).filter(Material.material_type != "answer")
    if subjects is not None:
        if not subjects:
            return []
        q = q.filter(Material.subject.in_(list(_subj_scope_for(db, Material, subjects))))
    mats = q.order_by(Material.created_at.desc()).all()
    ids = [m.id for m in mats]
    views, downloads = {}, {}
    if ids:
        for v in db.query(MaterialView).filter(MaterialView.material_id.in_(ids)).all():
            d = downloads if v.action == "download" else views
            d.setdefault(v.material_id, set()).add(v.student_id)
    import re as _re
    def _cls_d(x):
        m2 = _re.search(r"\d+", str(x or ""))
        return m2.group(0) if m2 else ""
    # Stored variant naam ('X (229)') ko assigned canonical naam pe lao (teacher scope me)
    _canon = {}
    if subjects:
        try:
            _sc = _subj_scope_for(db, Material, subjects)
            _canon = {_subj_norm(k): v for k, v in _sc.items()}
        except Exception:
            _canon = {}
    def _cn(s):
        return _canon.get(_subj_norm(s), s or "General")
    # v123: single-class subject (Social Science=10, Physics=12) ke purane
    # galat/blank-tagged materials bhi EK hi card me merge ho jaate hain —
    # effective class hamesha official class. Class-split sirf dual-class
    # subjects (English/Hindi/DEO — dono classes me hote hain) ke liye rehta hai.
    _fixed_cls = {}
    for m in mats:
        nm = _cn(m.subject)
        if nm not in _fixed_cls:
            try:
                _fixed_cls[nm] = _subj_class_digits(db, m.subject)
            except Exception:
                _fixed_cls[nm] = None
    def _ecd(m):
        fx = _fixed_cls.get(_cn(m.subject))
        return fx if fx else _cls_d(getattr(m, "class_name", ""))
    # Same-naam subjects (e.g. English / Data Entry Operations) Class 10 & 12 dono
    # me ho to tree me alag-alag node banao: "Subject (Class 10)" / "Subject (Class 12)".
    raw = {}
    for m in mats:
        cd = _ecd(m)
        raw.setdefault(_cn(m.subject), set()).add(cd)
    def _disp(sub, cd):
        cls_set = raw.get(sub) or set()
        real = {c for c in cls_set if c}
        if len(real) > 1 and cd:
            return "%s (Class %s)" % (sub, cd)
        return sub
    tree = {}
    for m in mats:
        cd = _ecd(m)
        label = _disp(_cn(m.subject), cd)
        sub = tree.setdefault(label, {"subject": label, "chapters": {}})
        ch = sub["chapters"].setdefault(m.chapter or "General", {"chapter": m.chapter or "General", "items": []})
        ch["items"].append({
            "id": m.id, "part": m.part or "", "type": m.material_type,
            "category": m.category or "", "title": m.title or "",
            "filename": m.filename or "", "teacher_name": m.teacher_name or "",
            "class_name": getattr(m, "class_name", "") or "",
            "date": str(m.created_at)[:10] if m.created_at else "",
            "views": len(views.get(m.id, ())), "downloads": len(downloads.get(m.id, ())),
        })
    # v130: editor se bane DPP (DppPack, source='created') bhi Classes Material tree
    # me dikhein taaki wahin se download ho. Poora block guarded — koi error aaye to
    # normal material tree waisa hi rehta hai.
    try:
        from models import DppPack
        _dpacks = db.query(DppPack).filter(DppPack.source == "created").order_by(
            DppPack.created_at.desc()).all()
        if subjects is not None:
            _allowed = set()
            try:
                _allowed |= {_subj_norm(k) for k in _subj_scope_for(db, Material, subjects).keys()}
            except Exception:
                pass
            _allowed |= {_subj_norm(x) for x in (subjects or [])}
            _dpacks = [pk for pk in _dpacks
                       if _subj_norm(pk.subject) in _allowed or _subj_norm(_cn(pk.subject)) in _allowed]
        for pk in _dpacks:
            nm = _cn(pk.subject)
            cd = _fixed_cls.get(nm) or _cls_d(getattr(pk, "class_name", ""))
            label = _disp(nm, cd)
            sub = tree.setdefault(label, {"subject": label, "chapters": {}})
            ch = sub["chapters"].setdefault(pk.chapter or "General",
                                            {"chapter": pk.chapter or "General", "items": []})
            ch["items"].append({
                "id": None, "dpp_pack_id": pk.id, "part": pk.part or "",
                "type": "dpp", "category": "", "title": pk.title or "",
                "filename": (pk.title or "DPP") + ".pdf", "teacher_name": "",
                "class_name": getattr(pk, "class_name", "") or "",
                "date": str(pk.created_at)[:10] if pk.created_at else "",
                "views": int(getattr(pk, "views", 0) or 0),
                "downloads": int(getattr(pk, "downloads", 0) or 0),
            })
    except Exception:
        pass
    out = []
    for s in tree.values():
        chapters = []
        for c in s["chapters"].values():
            c["items"].sort(key=lambda x: (x["part"], x["type"]))
            chapters.append(c)
        chapters.sort(key=lambda c: c["chapter"])
        out.append({"subject": s["subject"], "chapters": chapters})
    out.sort(key=lambda s: s["subject"])
    return out


def _material_audience(db, material_id):
    """Who viewed / downloaded this material."""
    from models import Material, MaterialView, StudentProfile, User
    m = db.query(Material).options(defer(Material.content_b64)).filter(Material.id == material_id).first()
    if not m:
        raise HTTPException(404, "Material not found")
    rows = db.query(MaterialView).filter(MaterialView.material_id == material_id).all()
    sids = list({r.student_id for r in rows})
    smap = {}
    if sids:
        for sp in db.query(StudentProfile).filter(StudentProfile.id.in_(sids)).all():
            smap[sp.id] = (sp.user.name if sp.user else ("Student #%d" % sp.id))
    seen, viewers, downloaders = {}, [], []
    for r in rows:
        nm = smap.get(r.student_id, "Student")
        key = (r.student_id, r.action)
        if key in seen:
            continue
        seen[key] = True
        entry = {"student_id": r.student_id, "name": nm, "at": str(r.created_at)[:16]}
        (downloaders if r.action == "download" else viewers).append(entry)
    return {"material": {"id": m.id, "title": m.title, "type": m.material_type,
                         "subject": m.subject, "chapter": m.chapter, "part": m.part},
            "viewers": viewers, "downloaders": downloaders}


@router.get("/materials-tree")
def teacher_materials_tree(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    return {"subjects": _material_tree(db, tp.subjects or [])}


@router.get("/material/{mid}/audience")
def teacher_material_audience(mid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    return _material_audience(db, mid)

# ==================================================================
#  SMART EXTRA CLASS — auto-shift ke saath
#  Teacher extra class daalta hai -> uske baad ki us subject ki saari
#  classes apne aap aage khisak jaati hain (sirf un weekdays par jinpe
#  us subject ki class hoti hai). Session end (default 10 Sept) ke baad
#  jaane par warning + extra weekdays ka suggestion.
# ==================================================================
import os as _os
from datetime import date as _date, timedelta as _td, datetime as _dt2

WEEK = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]


def _default_end():
    raw = (_os.getenv("SESSION_END_DATE") or "").strip()
    if raw:
        try:
            return _dt2.strptime(raw, "%Y-%m-%d").date()
        except Exception:
            pass
    today = _date.today()
    end = _date(today.year, 9, 10)          # fallback: 10 September
    if end < today:
        end = _date(today.year + 1, 9, 10)
    return end


def session_end_for(db, subject=None, class_level=None, batch=None):
    """Deadline priority: subject+class > subject > batch > global > default.
    Admin ise Settings -> Session Deadlines me set karta hai."""
    from models import SessionDeadline as _SD
    rows = db.query(_SD).all()
    if not rows:
        return _default_end(), "default"

    def find(scope, key):
        k = (key or "").strip().lower()
        for r in rows:
            if r.scope == scope and (r.key or "").strip().lower() == k:
                return r
        return None

    if subject and class_level:
        r = find("subject", f"{subject}|{class_level}")
        if r:
            return r.end_date, f"{subject} (Class {class_level})"
    if subject:
        r = find("subject", subject)
        if r:
            return r.end_date, subject
    if batch:
        r = find("batch", batch)
        if r:
            return r.end_date, batch
    r = find("global", "")
    if r:
        return r.end_date, "All batches"
    return _default_end(), "default"


def _subject_weekdays(entries):
    """Us subject ki classes kin weekdays par hoti hain (0=Mon)."""
    days = sorted({e.entry_date.weekday() for e in entries if e.entry_date})
    return days or [0, 2, 4]               # fallback: Mon/Wed/Fri


def _next_slots(start_after, weekdays, count, busy_dates):
    """Agli `count` free dates jo `weekdays` par aati hain (busy dates skip)."""
    out, d = [], start_after + _td(days=1)
    guard = 0
    while len(out) < count and guard < 800:
        guard += 1
        if d.weekday() in weekdays and d not in busy_dates:
            out.append(d)
        d += _td(days=1)
    return out


def _cls_norm(s):
    """'Class 12' == '12' == 'XII' na bhi ho to kam se kam number/base match."""
    return re.sub(r"[^0-9a-z]", "", (s or "").lower().replace("class", "").strip())


def _plan_shift(db, tp, subject, new_date, class_name=None):
    """Preview: extra class ke baad kya-kya shift hoga.
    Scope wahi jo teacher ko timetable me DIKHTA hai — subject-scope (teacher_id
    se nahi) + subject name variants ('English (229)') + optional class filter.
    Warna 'No classes after this date' ka jhootha message aata tha."""
    from models import TimetableEntry
    canon = _subj_canon(subject)
    scope = _subj_scope_for(db, TimetableEntry, tp.subjects or [])
    variants = [s for s, c in scope.items() if c == canon] or [subject]
    rows = db.query(TimetableEntry).filter(
        TimetableEntry.subject.in_(variants),
        TimetableEntry.status == "approved",
    ).all()
    # class filter: 10th ki extra class 12th ki timeline na khiskaye (aur vice-versa)
    if class_name and any((e.class_name or "").strip() for e in rows):
        cn = _cls_norm(class_name)
        rows = [e for e in rows if _cls_norm(e.class_name) == cn]
    dated = [e for e in rows if e.entry_date]
    weekdays = _subject_weekdays(dated)

    # jo classes new_date ke din ya uske baad hain, wo ek slot aage khiskengi
    affected = sorted([e for e in dated if e.entry_date >= new_date],
                      key=lambda e: (e.entry_date, e.id))
    # TeacherProfile.batch plain String hai (StudentProfile ka enum nahi) —
    # dono case sambhal lo, warna batch wale teacher ka extra-class 500 deta hai.
    _b = getattr(tp, "batch", None)
    _b = (getattr(_b, "value", None) or _b) or None
    end, end_src = session_end_for(db, subject=subject,
                                   class_level=(dated[0].class_name if dated else None),
                                   batch=_b)
    if not affected:
        return {"shifted": [], "weekdays": [WEEK[i] for i in weekdays], "overflow": False,
                "last_date": None, "session_end": str(end), "deadline_for": end_src}

    busy = {new_date}
    slots = _next_slots(new_date, weekdays, len(affected), busy)
    shifted, last = [], None
    for e, nd in zip(affected, slots):
        shifted.append({"id": e.id, "chapter": e.chapter, "part": e.part or "",
                        "from": str(e.entry_date), "to": str(nd),
                        "day": WEEK[nd.weekday()]})
        last = nd
    overflow = bool(last and last > end)
    # overflow -> baaki weekdays suggest karo (jinpe abhi class nahi hoti)
    free_days = [WEEK[i] for i in range(7) if i not in weekdays and i != 6]
    return {"shifted": shifted, "weekdays": [WEEK[i] for i in weekdays],
            "overflow": overflow, "last_date": str(last) if last else None,
            "session_end": str(end), "deadline_for": end_src, "suggest_days": free_days,
            "over_by": (last - end).days if overflow and last else 0}


@router.post("/extra-class/preview")
def extra_class_preview(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Extra class daalne se pehle dikhao ki kitni classes shift hongi."""
    tp = get_teacher_profile(current_user, db)
    subject = (payload.get("subject") or "").strip()
    if subject not in (tp.subjects or []):
        raise HTTPException(status_code=400, detail="This is not your subject")
    try:
        nd = _dt2.strptime(payload["date"], "%Y-%m-%d").date()
    except Exception:
        raise HTTPException(status_code=400, detail="Please choose a valid date")
    if not bool(payload.get("shift", True)):
        end, src = session_end_for(db, subject=subject)
        return {"shifted": [], "overflow": False, "session_end": str(end), "deadline_for": src}
    return _plan_shift(db, tp, subject, nd, class_name=(payload.get("class_name") or None))


@router.post("/extra-class")
def create_extra_class(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """v86: Extra class request — KOI AUTO-SHIFT NAHI. Missed/absent class ki
       compensation teacher khud alag slot pe karta hai; baaki classes apni
       jagah rehti hain. Sab kuch admin approval par hi live hota hai."""
    from models import TimetableEntry, User, UserRole, Notification
    tp = get_teacher_profile(current_user, db)
    _ensure_v86(db)
    subject = (payload.get("subject") or "").strip()
    if subject not in (tp.subjects or []):
        raise HTTPException(status_code=400, detail="This is not your subject")
    try:
        nd = _dt2.strptime(payload["date"], "%Y-%m-%d").date()
    except Exception:
        raise HTTPException(status_code=400, detail="Please choose a valid date")
    time_text = (payload.get("time") or "").strip()
    chapter = (payload.get("chapter") or payload.get("topic") or "Extra Class").strip()
    part = (payload.get("topic") or "").strip() or None

    # Same slot pe pehle se class hai to rok do — extra class alag timing pe hi honi chahiye
    def _tmin(s):
        m = _parse_hhmm(s)
        return m if m is not None else -1
    clash = None
    for x in db.query(TimetableEntry).filter(
            TimetableEntry.teacher_id == tp.id,
            TimetableEntry.entry_date == nd,
            TimetableEntry.status == "approved").all():
        if time_text and _tmin(x.time_text) == _tmin(time_text) and _tmin(time_text) >= 0:
            clash = x
            break
    if clash:
        raise HTTPException(status_code=400,
            detail=f"Is date/time pe aapki {clash.subject} class already hai "
                   f"({clash.time_text or ''}). Extra class ke liye alag timing choose karein.")

    e = TimetableEntry(
        teacher_id=tp.id, subject=subject,
        class_name=payload.get("class_name", "Class 12"),
        chapter=chapter, part=part, entry_date=nd, day=WEEK[nd.weekday()],
        time_text=time_text or None, entry_type="chapter", status="pending",
    )
    db.add(e); db.flush()

    for adm in db.query(User).filter(User.role == UserRole.admin).all():
        msg = (f"{current_user.name} requested an extra class for {subject} "
               f"({nd} {time_text}). Baaki classes shift nahi hongi — ye alag slot pe hogi.")
        db.add(Notification(user_id=adm.id, title="New Extra Class Request",
                            message=msg, notif_type="class_request"))
    db.commit(); db.refresh(e)
    return {"id": e.id, "shift_count": 0,
            "message": "Request sent to the admin. Once approved, this extra class will be added at your chosen time — other classes stay as they are."}

# =====================================================================
# v86: CLASS MOVE HELPERS (leave auto-move / admin move) + TEACHER RESCHEDULE
# Sirf TEACHER ki khud ki request reschedule-count me jati hai —
# admin ya approved-leave ki wajah se hua move count/salary ko touch nahi karta.
# =====================================================================
def _students_of_subject(db, subject):
    """Us subject ke enrolled students (canon match) — notification targeting."""
    from models import StudentProfile
    nk = (_SR.canon_norm(subject) if _SR else subject)
    out = []
    for sp in db.query(StudentProfile).options(defer(StudentProfile.photo_b64)).all():
        try:
            subs = sp.subjects or []
        except Exception:
            subs = []
        if subs and nk in {(_SR.canon_norm(x) if _SR else x) for x in subs} and sp.user:
            out.append(sp)
    return out

def _notify_class_moved(db, tp, moved, headline, note):
    """moved = [(entry, old_date, old_time), ...] — students ko teacher ki PHOTO ke
    saath notification (image_url -> student photo endpoint)."""
    from models import Notification
    img = f"/api/student/teacher/{tp.id}/photo" if getattr(tp, "photo_b64", None) else None
    by_sub = {}
    for e, od, ot in moved:
        by_sub.setdefault(e.subject or "", []).append((e, od, ot))
    for subj, rows in by_sub.items():
        parts = []
        for e, od, ot in rows[:6]:
            cls = f" ({e.class_name})" if e.class_name else ""
            ods = od.strftime("%d %b") if od else "—"
            nds = e.entry_date.strftime("%d %b") if e.entry_date else "—"
            parts.append(f"{subj}{cls}: {ods} {ot or ''} → {nds} {e.time_text or ''}".replace("  ", " ").strip())
        if len(rows) > 6:
            parts.append(f"+{len(rows) - 6} more")
        msg = (note + " " if note else "") + "; ".join(parts) + ". Check your time table."
        for sp in _students_of_subject(db, subj):
            db.add(Notification(user_id=sp.user.id, title=f"{headline}{subj}",
                                message=msg, notif_type="class_rescheduled",
                                image_url=img))

def _auto_move_leave_classes(db, tp, lv):
    """Approved FULL leave ke dinon ki pending classes aage move karta hai —
    leave ke span ke barabar days aage; same-time collision pe aur aage (max 14).
    resched_by='leave' => teacher ke reschedule count/salary pe koi asar NAHI."""
    from models import TimetableEntry
    _ensure_v86(db)
    if (lv.leave_type or "full") != "full":
        return []
    span = max(1, (lv.end_date - lv.start_date).days + 1)
    rows = db.query(TimetableEntry).filter(
        TimetableEntry.teacher_id == tp.id,
        TimetableEntry.entry_date >= lv.start_date,
        TimetableEntry.entry_date <= lv.end_date,
        TimetableEntry.status == "approved",
        TimetableEntry.completed == False).all()
    moved = []
    for e in rows:
        od, ot = e.entry_date, e.time_text
        nd = e.entry_date + timedelta(days=span)
        em = _parse_hhmm(e.time_text)
        for _try in range(14):
            conf = db.query(TimetableEntry).filter(
                TimetableEntry.teacher_id == tp.id,
                TimetableEntry.entry_date == nd,
                TimetableEntry.id != e.id,
                TimetableEntry.status == "approved").all()
            if em is None or not any(_parse_hhmm(x.time_text) == em for x in conf):
                break
            nd += timedelta(days=1)
        e.entry_date = nd
        e.day = nd.strftime("%A")
        e.resched_by = "leave"
        moved.append((e, od, ot))
    return moved

@router.post("/tt-reschedule")
def request_tt_reschedule(payload: dict = Body(...), db: Session = Depends(get_db),
                          current_user=Depends(get_teacher)):
    """Topic-edit modal ka 'Reschedule' button — timetable entry ki nayi date/time
    request. Admin approval ke baad apply hoti hai aur TEACHER ke monthly
    reschedule count me jati hai."""
    from models import (TimetableEntry, RescheduleRequest, RescheduleStatus,
                        User, UserRole, Notification)
    from datetime import time as _time
    tp = get_teacher_profile(current_user, db)
    _ensure_v86(db)
    try:
        eid = int(payload.get("entry_id"))
    except Exception:
        raise HTTPException(status_code=400, detail="entry_id is required")
    e = db.query(TimetableEntry).filter(TimetableEntry.id == eid).first()
    if not e or (e.subject not in (tp.subjects or []) and e.teacher_id != tp.id):
        raise HTTPException(status_code=404, detail="Entry not found")
    if e.completed:
        raise HTTPException(status_code=400, detail="Completed class reschedule nahi ho sakti")
    try:
        nd = _dt2.strptime((payload.get("new_date") or "").strip(), "%Y-%m-%d").date()
    except Exception:
        raise HTTPException(status_code=400, detail="Please choose a valid new date")
    nt_txt = (payload.get("new_time") or "").strip()
    mins = _parse_hhmm(nt_txt)
    if not nt_txt or mins is None:
        raise HTTPException(status_code=400, detail="Please enter a valid new time (e.g. 5:00 PM)")
    reason = (payload.get("reason") or "").strip()
    if not reason:
        raise HTTPException(status_code=400, detail="Please write a short reason")
    dup = db.query(RescheduleRequest).filter(
        RescheduleRequest.tt_entry_id == e.id,
        RescheduleRequest.status == RescheduleStatus.pending).first()
    if dup:
        raise HTTPException(status_code=400,
                            detail="Is class ki reschedule request pehle se pending hai")
    om = _parse_hhmm(e.time_text)
    h, m = divmod(mins, 60)
    rs = RescheduleRequest(
        class_entry_id=None, tt_entry_id=e.id, teacher_id=tp.id,
        original_date=e.entry_date,
        original_time=(_time(*divmod(om, 60)) if om is not None else None),
        new_date=nd, new_time=_time(h, m), reason=reason,
        status=RescheduleStatus.pending)
    db.add(rs)
    for adm in db.query(User).filter(User.role == UserRole.admin).all():
        db.add(Notification(
            user_id=adm.id, title="Reschedule Request",
            message=(f"{current_user.name} wants to move {e.subject} ({e.class_name or ''}) "
                     f"of {e.entry_date} {e.time_text or ''} to {nd} {nt_txt}. Reason: {reason}"),
            notif_type="reschedule_request"))
    db.commit()
    return {"message": "Request sent to the admin. Approval ke baad class move hogi — ye aapke monthly reschedule count me judegi."}

# =====================================================================
# TEACHER ATTENDANCE (PUNCH IN / PUNCH OUT) + CONTRACT + PAYOUT
# =====================================================================
def _ist_now():
    """Railway server UTC pe chalta hai — IST me convert karke store/show karte hain."""
    return datetime.utcnow() + timedelta(hours=5, minutes=30)

def _fmt_t(dt):
    return dt.strftime("%I:%M %p") if dt else None

def _att_hours(a):
    if a and a.punch_in and a.punch_out:
        return round((a.punch_out - a.punch_in).total_seconds() / 3600, 1)
    return None

# Present tabhi count hoga jab punch-in aur punch-out ke beech kam se kam itne ghante hon.
MIN_PRESENT_HOURS = 1.0

def _is_present(a):
    """Present = punch-in + punch-out dono, aur gap >= MIN_PRESENT_HOURS."""
    if not a or not a.punch_in or not a.punch_out:
        return False
    return (a.punch_out - a.punch_in).total_seconds() >= MIN_PRESENT_HOURS * 3600

def _is_short(a):
    """Short day = punch complete but gap < MIN_PRESENT_HOURS (present nahi, absent bhi nahi)."""
    if not a or not a.punch_in or not a.punch_out:
        return False
    return (a.punch_out - a.punch_in).total_seconds() < MIN_PRESENT_HOURS * 3600

def _leave_days_map(db, teacher_id, start, end, unpaid_only=False):
    """Approved leaves -> {date: 1.0|0.5} month range ke andar.
    v86: unpaid_only=True -> sirf UNPAID leaves (salary deduction ke liye);
    paid leaves attendance me 'Leave' dikhti hain par pay se kati nahi."""
    from models import TeacherLeave
    out = {}
    q = db.query(TeacherLeave).filter(
        TeacherLeave.teacher_id == teacher_id,
        TeacherLeave.status == "approved",
        TeacherLeave.start_date < end, TeacherLeave.end_date >= start)
    for lv in q.all():
        if unpaid_only and bool(getattr(lv, "paid", False)):
            continue
        d = max(lv.start_date, start)
        e = min(lv.end_date, end - timedelta(days=1))
        val = 0.5 if lv.leave_type == "half" else 1.0
        while d <= e:
            out[d] = val
            d += timedelta(days=1)
    return out

def _elapsed_days(start, end):
    """Month range me aaj tak kitne din guzar chuke (future month = 0)."""
    today = _ist_now().date()
    if today < start:
        return 0
    last = min(today, end - timedelta(days=1))
    return (last - start).days + 1

def _month_range(month: str):
    """'2026-07' -> (date(2026,7,1), date(2026,8,1)). Galat format pe current month."""
    try:
        y, m = month.split("-"); y = int(y); m = int(m)
        assert 1 <= m <= 12
    except Exception:
        n = _ist_now(); y, m = n.year, n.month
    start = date(y, m, 1)
    end = date(y + 1, 1, 1) if m == 12 else date(y, m + 1, 1)
    return start, end

# =============================================
# SMART WORK POLICY (per-teacher timing — admin set karta hai)
# =============================================
WORK_TYPES = ("full_time", "part_time")
POLICY_MODES = ("fixed", "hours", "flexible")
DEFAULT_REQUIRED_HOURS = 8.0
MAX_WORK_HOURS = 16.0

def _default_policy(teacher_id=None):
    return {"teacher_id": teacher_id, "configured": False,
            "work_type": "full_time", "mode": "hours",
            "required_hours": DEFAULT_REQUIRED_HOURS,
            "entry_time": "", "exit_time": "", "break_minutes": 0,
            "disabled": False}

def _parse_hhmm24(s):
    """'09:30' (24h policy time) -> minutes since midnight. Invalid -> None.
    NOTE: naam alag rakha hai — 2487 wala flexible _parse_hhmm (am/pm samajhta hai)
    _delay_of/_duration_of ke liye chahiye; yeh usko shadow nahi karna chahiye."""
    try:
        h, m = str(s or "").strip().split(":")
        h, m = int(h), int(m)
        if 0 <= h <= 23 and 0 <= m <= 59:
            return h * 60 + m
    except Exception:
        pass
    return None

def _policy_from_row(p, teacher_id):
    """Saved row + defaults merge -> normalized policy dict."""
    d = _default_policy(teacher_id)
    if p:
        d.update({"configured": True,
                  "work_type": p.work_type if p.work_type in WORK_TYPES else "full_time",
                  "mode": p.mode if p.mode in POLICY_MODES else "hours",
                  "required_hours": float(p.required_hours or DEFAULT_REQUIRED_HOURS),
                  "entry_time": p.entry_time or "", "exit_time": p.exit_time or "",
                  "break_minutes": int(p.break_minutes or 0),
                  "disabled": bool(getattr(p, "disabled", False))})
    if d["work_type"] != "full_time":
        d["break_minutes"] = 0          # lunch break sirf full-time me count hota hai
    return d

def _policy_dict(db, teacher_id):
    """Teacher ki saved policy + defaults merge. Punch/reports/payout sab yahi use karte hain."""
    from models import TeacherWorkPolicy
    p = db.query(TeacherWorkPolicy).filter(TeacherWorkPolicy.teacher_id == teacher_id).first()
    return _policy_from_row(p, teacher_id)

def _policy_map(db, teacher_ids):
    """Bulk: {teacher_id: policy dict} — admin reports me N+1 queries se bachne ke liye."""
    from models import TeacherWorkPolicy
    rows = db.query(TeacherWorkPolicy).filter(TeacherWorkPolicy.teacher_id.in_(teacher_ids or [0])).all()
    by_id = {r.teacher_id: r for r in rows}
    return {tid: _policy_from_row(by_id.get(tid), tid) for tid in (teacher_ids or [])}

def _policy_required(pol):
    """Ek din ke required NET hours.
    flexible = sirf minimum 1h; fixed = entry-exit span - break; hours = saved value."""
    if pol.get("mode") == "flexible":
        return MIN_PRESENT_HOURS
    if pol.get("mode") == "fixed":
        en, ex = _parse_hhmm24(pol.get("entry_time")), _parse_hhmm24(pol.get("exit_time"))
        if en is not None and ex is not None and ex > en:
            span = (ex - en) / 60.0 - (pol.get("break_minutes") or 0) / 60.0
            return max(MIN_PRESENT_HOURS, round(span, 2))
    return max(MIN_PRESENT_HOURS, float(pol.get("required_hours") or DEFAULT_REQUIRED_HOURS))

def _net_hours(a, pol):
    """Punch gap - lunch break (sirf full_time) = counted working hours."""
    gross = _att_hours(a)
    if gross is None:
        return None
    brk = (pol.get("break_minutes") or 0) / 60.0 if pol.get("work_type") == "full_time" else 0.0
    return round(max(0.0, gross - brk), 2)

def _day_status(a, pol):
    """Policy-aware day status: working | present | short | none.
    short = punch complete par required hours se kam => 'Present (Short)' count hota hai."""
    if not a or not a.punch_in:
        return "none"
    if not a.punch_out:
        return "working"
    net = _net_hours(a, pol)
    return "present" if (net or 0) >= _policy_required(pol) else "short"

def _extra_hours(a, pol):
    """Required se zyada net hours — SIRF display ke liye, extra payout NAHI."""
    if pol.get("mode") == "flexible":
        return 0.0                       # flexible me assigned hours hi nahi — extra concept nahi
    net = _net_hours(a, pol)
    if net is None:
        return 0.0
    return round(max(0.0, net - _policy_required(pol)), 2)

def _policy_label(pol):
    """UI chip text, e.g. 'Full Time · 8h/day' / 'Part Time · 09:30-18:30'."""
    wt = "Full Time" if pol.get("work_type") == "full_time" else "Part Time"
    mode = pol.get("mode")
    if mode == "flexible":
        return wt + " · Flexible (min 1h)"
    if mode == "fixed" and pol.get("entry_time") and pol.get("exit_time"):
        return "%s · %s-%s" % (wt, pol["entry_time"], pol["exit_time"])
    return "%s · %sh/day" % (wt, _policy_required(pol))

# =====================================================================
# PERFORMANCE PAYOUT ENGINE (monthly task-based, 1 Aug 2026 se effective)
# =====================================================================
PAYOUT_PERF_START = "2026-08"   # is month se performance system apply hota hai

# default template: (key, label, source, weight%, target) - total weight 100
PERF_DEFAULT_TEMPLATE = [
    ("live_class", "Live Classes",            "auto",   40, 0),   # target auto = timetable count
    ("dpp",        "DPP (1 per chapter)",     "auto",   15, 0),   # target auto = us month padhe gaye chapters
    ("test",       "Weekly Tests",            "auto",   15, 4),
    ("doubt",      "Doubt Resolution",        "auto",   5,  8),
    ("content",    "Notes / Free Content",    "auto",   5,  4),
    ("oneshot",    "One Shot Videos",         "manual", 8,  8),
    ("rapid",      "Rapid Revision Videos",   "manual", 4,  4),
    ("ytlive",     "YouTube Live Sessions",   "manual", 4,  4),
    ("shorts",     "Shorts",                  "manual", 4,  8),
]

_PAYOUT_TABLES_READY = False
def _ensure_payout_tables(db):
    """payout_templates / payout_tasks / payout_months tables pehli use me bana do
    (server pe Base.metadata.create_all na ho to bhi upgrade ho jaye). MySQL/Postgres
    dono dialects handle; har statement best-effort."""
    global _PAYOUT_TABLES_READY
    if _PAYOUT_TABLES_READY:
        return
    from sqlalchemy import text as _text
    stmts = [
        """CREATE TABLE IF NOT EXISTS payout_templates (
             id INTEGER PRIMARY KEY, teacher_id INTEGER, key VARCHAR(30),
             label VARCHAR(80), target INTEGER DEFAULT 0,
             weight_pct FLOAT DEFAULT 0, source VARCHAR(10) DEFAULT 'manual',
             sort INTEGER DEFAULT 0)""",
        "CREATE INDEX IF NOT EXISTS ix_payout_templates_teacher ON payout_templates (teacher_id)",
        """CREATE TABLE IF NOT EXISTS payout_tasks (
             id INTEGER PRIMARY KEY, teacher_id INTEGER, month VARCHAR(7),
             key VARCHAR(30), title VARCHAR(200), status VARCHAR(20) DEFAULT 'pending',
             ref_id INTEGER, done_date DATE, note VARCHAR(300),
             approved_by VARCHAR(120), approved_at DATETIME, created_at DATETIME)""",
        "CREATE INDEX IF NOT EXISTS ix_payout_tasks_teacher ON payout_tasks (teacher_id)",
        "CREATE INDEX IF NOT EXISTS ix_payout_tasks_month ON payout_tasks (month)",
        """CREATE TABLE IF NOT EXISTS payout_months (
             id INTEGER PRIMARY KEY, teacher_id INTEGER, month VARCHAR(7),
             status VARCHAR(20) DEFAULT 'finalized', snapshot TEXT,
             finalized_at DATETIME, paid_at DATETIME, created_at DATETIME)""",
        "CREATE INDEX IF NOT EXISTS ix_payout_months_teacher ON payout_months (teacher_id)",
    ]
    for st in stmts:
        try:
            db.execute(_text(st))
            db.commit()
        except Exception:
            db.rollback()
    _PAYOUT_TABLES_READY = True

def _chapter_key(name):
    """Chapter naam ko compare-able key me badalta hai: 'Chapter 3: Motion' -> 'n3',
    'Lesson 3 - X (Part 2)' -> 'n3'. Parts ignore hote hain (kisi bhi part ka DPP chalega).
    Number na mile to normalized string."""
    import re as _re
    s2 = _re.sub(r"\bpart\s*\d+\b", "", (name or "").lower())
    m = _re.search(r"(?:chapter|lesson|ch\.?|path|paath)\s*[-\u2013:.]?\s*(\d+)", s2)
    if m:
        return "n" + m.group(1)
    m2 = _re.search(r"(\d+)", s2)
    if m2:
        return "n" + m2.group(1)
    return _re.sub(r"[^a-z0-9\u0900-\u097F]+", " ", s2).strip()[:40]

def _is_grammar_entry(name):
    """Grammar/writing entries ke liye DPP compulsory nahi (book chapters ke liye hi)."""
    import re as _re
    return bool(_re.search(r"grammar|tense|narration|direct|indirect|voice|clause|modal|preposition|conjunction|writing|essay|letter|notice|punctuation|comprehension", (name or ""), _re.I))

def _perf_seed(db, tp_id):
    """Teacher ka template nahi hai to default bana do (idempotent)."""
    from models import PayoutTemplate
    _ensure_payout_tables(db)
    if db.query(PayoutTemplate).filter(PayoutTemplate.teacher_id == tp_id).count():
        return
    for i, (k, lbl, src, w, tg) in enumerate(PERF_DEFAULT_TEMPLATE):
        db.add(PayoutTemplate(teacher_id=tp_id, key=k, label=lbl, source=src,
                              weight_pct=w, target=tg, sort=i))
    db.commit()

def compute_performance(db, teacher_id: int, month: str):
    """Month ka performance calculation - policy ke 5 rules ke saath:
    1. sab kuch same month me complete -> 100% payout
    2. same month me postpone karke complete -> completed (no deduction)
    3. next month me complete -> previous month ke liye count NAHI (delayed)
    4. delayed sirf record ke liye; next month ka assigned ho tabhi wahan count
    5. assigned/completed/pending/delayed/completion%/payout% sab auto.
    Category weight% ke hisaab se proportional deduction."""
    from models import (PayoutTemplate, PayoutTask, PayoutMonth, TimetableEntry,
                        DPP, Doubt, Material)
    start, end = _month_range(month)
    mk = start.strftime("%Y-%m")
    today = _ist_now().date()
    period_end = min(today, end - timedelta(days=1))
    is_current = (start.year == today.year and start.month == today.month)
    started = mk >= PAYOUT_PERF_START

    _perf_seed(db, teacher_id)
    tpl = db.query(PayoutTemplate).filter(
        PayoutTemplate.teacher_id == teacher_id).order_by(PayoutTemplate.sort, PayoutTemplate.id).all()

    tasks = db.query(PayoutTask).filter(
        PayoutTask.teacher_id == teacher_id, PayoutTask.month == mk).all()
    def _task_out(t):
        delayed = bool(t.done_date and not (start <= t.done_date < end))
        return {"id": t.id, "key": t.key, "title": t.title, "status": t.status,
                "done_date": str(t.done_date) if t.done_date else None,
                "note": t.note or "", "delayed": delayed,
                "approved_by": t.approved_by, "ref_id": t.ref_id,
                "created_at": t.created_at.strftime("%d %b") if t.created_at else ""}

    cats = []
    for t in tpl:
        target = t.target or 0
        done = pending = delayed = 0
        missing = []
        if not started:
            cats.append({"key": t.key, "label": t.label, "source": t.source,
                         "target": 0, "done": 0, "pending": 0, "delayed": 0,
                         "weight": t.weight_pct, "completion": 0})
            continue
        if t.key == "live_class":
            entries = db.query(TimetableEntry).filter(
                TimetableEntry.teacher_id == teacher_id,
                TimetableEntry.entry_date >= start, TimetableEntry.entry_date < end,
                TimetableEntry.entry_type == "chapter",
                TimetableEntry.status == "approved").all()
            missed = {x.ref_id for x in tasks if x.key == "live_class" and x.status == "missed"}
            target = len(entries)
            due = [e for e in entries if e.entry_date and e.entry_date <= period_end and e.id not in missed]
            done = len(due)
            pending = max(0, target - done)
        elif t.source == "auto":
            if t.key == "dpp":
                # RULE: 1 chapter = 1 DPP compulsory. Chapter ke 4 parts ho to kisi
                # bhi 1 part me DPP aa jaye -> chapter complete (baaki parts optional).
                # Kisi bhi part me DPP nahi -> wo chapter miss -> proportional deduction.
                entries = db.query(TimetableEntry).filter(
                    TimetableEntry.teacher_id == teacher_id,
                    TimetableEntry.entry_date >= start, TimetableEntry.entry_date < end,
                    TimetableEntry.entry_type == "chapter",
                    TimetableEntry.status == "approved").all()
                chapters = {}
                for e in entries:
                    if _is_grammar_entry(e.chapter):
                        continue
                    chapters.setdefault(((e.subject or "").strip().lower(), _chapter_key(e.chapter)),
                                        (e.chapter or "").strip())
                target = len(chapters)
                dpps = db.query(DPP).filter(
                    DPP.teacher_id == teacher_id, DPP.is_active == True,
                    DPP.created_at >= start, DPP.created_at < end).all()
                have = {((d.subject or "").strip().lower(), _chapter_key(d.reference)) for d in dpps}
                missing = [nm for k, nm in chapters.items() if k not in have]
                done = target - len(missing)
            elif t.key == "test":
                from models import Exam
                done = db.query(Exam).filter(Exam.teacher_id == teacher_id,
                        Exam.created_at >= start, Exam.created_at < end).count()
            elif t.key == "doubt":
                done = db.query(Doubt).filter(Doubt.teacher_id == teacher_id,
                        Doubt.resolved_at >= start, Doubt.resolved_at < end).count()
            elif t.key == "content":
                done = db.query(Material).options(defer(Material.content_b64)).filter(Material.teacher_id == teacher_id,
                        Material.created_at >= start, Material.created_at < end).count()
            pending = max(0, target - done) if target else 0
        else:  # manual - approved + done_date is month me ho tabhi count (rule 2/3)
            mine = [x for x in tasks if x.key == t.key and x.status == "approved"]
            in_month = [x for x in mine if x.done_date and start <= x.done_date < end]
            delayed = len([x for x in mine if x.done_date and not (start <= x.done_date < end)])
            done = len(in_month)
            pending = max(0, target - done) if target else 0
        completion = (min(done, target) / target) if target else 0
        row = {"key": t.key, "label": t.label, "source": t.source,
               "target": target, "done": done, "pending": pending,
               "delayed": delayed, "weight": t.weight_pct,
               "completion": round(completion, 4)}
        if t.key == "dpp" and started:
            row["missing"] = sorted(missing)
        cats.append(row)

    # weight renormalize: jin categories ka target 0 hai wo calculation se baahar
    active = [c for c in cats if c["target"] > 0]
    wsum = sum(c["weight"] for c in active) or 1
    perf_pct = sum(c["weight"] * c["completion"] for c in active) / wsum
    totals = {
        "target": sum(c["target"] for c in active),
        "done": sum(min(c["done"], c["target"]) for c in active),
        "pending": sum(c["pending"] for c in active),
        "delayed": sum(c["delayed"] for c in active),
        "completion_pct": round(perf_pct * 100, 1),
    }
    fin = db.query(PayoutMonth).filter(
        PayoutMonth.teacher_id == teacher_id, PayoutMonth.month == mk).first()
    return {
        "month": mk, "started": started, "is_current_month": is_current,
        "perf_start": PAYOUT_PERF_START,
        "perf_ratio": round(perf_pct, 6),
        "categories": cats, "totals": totals,
        "perf_pct": round(perf_pct * 100, 1),
        "tasks": [_task_out(x) for x in sorted(tasks, key=lambda z: (z.key, z.id))],
        "awaiting_approval": len([x for x in tasks if x.status == "pending"]),
        "finalized": bool(fin), "paid": bool(fin and fin.status == "paid"),
        "finalized_at": fin.finalized_at.strftime("%d %b %Y") if fin and fin.finalized_at else None,
    }

def compute_payout(db, teacher_id: int, month: str):
    """Transparent payout breakdown — teacher aur admin dono yahi dekhte hain.
    Net = Base + Allowances + Extras + Bonus - Manual Deductions - Attendance Deduction.
    Attendance Deduction = (absent + UNPAID approved leave) x per-day rate.
    v86: admin leave approve karte waqt PAID/UNPAID choose karta hai — PAID leave pe
    koi deduction nahi; UNPAID leave (default) per-day rate se katti hai.
    Bina approval ki chhutti (absent) ka penalty rule baad me add hoga."""
    from models import TeacherContract, TeacherAttendance, PayoutAdjustment
    _ensure_geofence(db); _ensure_v86(db)
    c = db.query(TeacherContract).filter(TeacherContract.teacher_id == teacher_id).first()
    if not c:
        return None
    start, end = _month_range(month)
    month_key = start.strftime("%Y-%m")
    att_rows = db.query(TeacherAttendance).filter(
        TeacherAttendance.teacher_id == teacher_id,
        TeacherAttendance.att_date >= start, TeacherAttendance.att_date < end).all()
    # Smart policy: short day (required hours se kam) bhi PRESENT/paid hai — deduction nahi.
    # Isliye present + short ka total pehle jaisa hi rehta hai => payout figure bilkul same.
    pol = _policy_dict(db, teacher_id)
    full_days = sum(1 for a in att_rows if _day_status(a, pol) == "present")
    short = sum(1 for a in att_rows if _day_status(a, pol) == "short")
    present = full_days + short
    extra_hours = round(sum(_extra_hours(a, pol) for a in att_rows), 1)
    leave_days = sum(_leave_days_map(db, teacher_id, start, end).values())
    unpaid_leave = sum(_leave_days_map(db, teacher_id, start, end, unpaid_only=True).values())
    paid_leave = round(leave_days - unpaid_leave, 1)
    wd = c.working_days or 26
    base = c.base_salary or 0
    per_day = round(base / wd) if wd else 0
    # Approved UNPAID leave = per-day rate ki deduction (penalty nahi);
    # PAID leave = deduction NAHI; short day (assigned hours se kam) PRESENT hai;
    # bina approval ki chhutti (absent) = deduction
    # (absent par alag PENALTY rule baad me add hoga).
    absent = max(0, round(wd - present - leave_days))
    leave_ded = round(per_day * unpaid_leave)
    att_deduction = per_day * absent + leave_ded
    adjs = db.query(PayoutAdjustment).filter(
        PayoutAdjustment.teacher_id == teacher_id,
        PayoutAdjustment.month == month_key).order_by(PayoutAdjustment.created_at).all()
    extras = sum(a.amount or 0 for a in adjs if a.kind == "extra")
    bonus = sum(a.amount or 0 for a in adjs if a.kind == "bonus")
    manual_ded = sum(a.amount or 0 for a in adjs if a.kind == "deduction")
    perf = compute_performance(db, teacher_id, month)
    gross = base + (c.allowances or 0)
    perf_pct = (perf.get("perf_ratio", perf["perf_pct"] / 100.0)) if (perf and perf["started"]) else 1.0
    perf_pay = round(gross * perf_pct)
    perf_ded = gross - perf_pay
    net = perf_pay + extras + bonus - manual_ded - att_deduction
    now = _ist_now()
    return {
        "month": month_key,
        "is_current_month": (start.year == now.year and start.month == now.month),
        "base_salary": base, "allowances": c.allowances or 0,
        "working_days": wd, "present_days": present, "absent_days": absent,
        "short_days": short, "full_days": full_days, "leave_days": leave_days,
        "paid_leave_days": paid_leave, "unpaid_leave_days": unpaid_leave,
        "extra_hours": extra_hours,
        "min_hours": MIN_PRESENT_HOURS, "required_hours": _policy_required(pol),
        "per_day_rate": per_day, "attendance_deduction": att_deduction,
        "leave_deduction": leave_ded, "unpaid_days": absent + unpaid_leave,
        "extras": extras, "bonus": bonus, "manual_deductions": manual_ded,
        "gross_salary": gross, "performance": perf,
        "perf_pct": perf["perf_pct"] if perf["started"] else None,
        "perf_pay": perf_pay, "perf_deduction": perf_ded,
        "net_payout": net,
        "rules": [r.strip() for r in (c.rules_text or "").splitlines() if r.strip()],
        "adjustments": [{"id": a.id, "kind": a.kind, "amount": a.amount, "note": a.note or ""} for a in adjs],
        "designation": c.designation, "accepted": bool(c.accepted)
    }

# =====================================================================
# EARNINGS-BASED PAYOUT (v80) — appointment-letter model
# pay structure: retainer (60% of max) + quality + notes/DPP + doubts + project
# =====================================================================
EARNINGS_DEFAULTS = {
    "class_retainer": 15000, "class_quality": 1000, "notes_dpp": 2000,
    "doubt_resolution": 1000, "project_delivery": 6000,
    "tests_target": 4, "videos_target": 8, "live_target": 4, "shorts_target": 8,
}
EARNINGS_PAY_FIELDS = ("class_retainer", "class_quality", "notes_dpp",
                       "doubt_resolution", "project_delivery")
EARNINGS_TARGET_FIELDS = ("tests_target", "videos_target", "live_target", "shorts_target")


def _jr(x):
    """JS Math.round ke barabar (half-up) — Python round() banker's rounding karta hai."""
    import math as _m
    return int(_m.floor(float(x) + 0.5))


def get_pay_config(db, teacher_id):
    """Saved config ho to wo, warna defaults ke saath transient (unsaved) row."""
    from models import TeacherPayConfig
    cfg = db.query(TeacherPayConfig).filter(TeacherPayConfig.teacher_id == teacher_id).first()
    if cfg:
        return cfg
    return TeacherPayConfig(teacher_id=teacher_id, **EARNINGS_DEFAULTS)


def calc_earnings(a, pay):
    """Appointment-letter earnings rule ka exact port.
    a = month activity dict, pay = 5 amounts + 4 targets.
    STRICT (v82): koi free credit nahi — jis component ki activity/denominator
    0 hai uska amount bhi 0. Sirf actual recorded activity se earning banti hai."""
    max_potential = sum(int(pay[k]) for k in EARNINGS_PAY_FIELDS)
    if not any([a["classes_scheduled"], a["classes_conducted"], a["notes_uploaded"],
                a["dpp_uploaded"], a["tests_created"], a["videos_made"], a["live_sessions"],
                a["shorts_made"], a["doubts_assigned"], a["doubts_resolved"],
                a["tasks_assigned"], a["tasks_on_time"]]):
        return {
            "class_earned": 0, "class_quality_earned": 0, "notes_earned": 0,
            "dpp_earned": 0, "tests_earned": 0, "doubt_earned": 0, "task_earned": 0,
            "gross_earned": 0, "max_potential": max_potential, "perf_score": 0,
            "tds": 0, "other_deduct": 0, "net_payable": 0,
            "pcts": {"class": 0, "quality": 0, "notes": 0, "dpp": 0, "tests": 0,
                     "doubt": 0, "task": 0, "content": 0},
        }
    sched = a["classes_scheduled"]
    cond = a["classes_conducted"]

    def _cap(x, t):
        # v82: target 0 hai aur actual bhi 0 -> 0 (free full credit nahi);
        # target 0 par actual > 0 -> full credit (target exceed kar diya)
        if t > 0:
            return min(x / t, 1)
        return 1 if x > 0 else 0

    class_pct = cond / sched if sched > 0 else 0
    class_earned = _jr(pay["class_retainer"] * class_pct)

    # Figma parity: late scheduled ke against measure hota hai (1 late of 20 -> 950),
    # lekin v82: ek bhi class conduct nahi hui to quality bhi 0 (free credit nahi)
    quality_pct = (max(0, 1 - a["late_classes"] / sched)
                   if (sched > 0 and cond > 0) else 0)
    class_quality_earned = _jr(pay["class_quality"] * quality_pct)

    notes_pct = a["notes_uploaded"] / cond if cond > 0 else 0
    # v83: DPP pct chapter-coverage se (1 chapter = 1 DPP). Purane payloads me
    # dpp_covered nahi hai to raw upload count par fallback.
    dpp_pct = a.get("dpp_covered", a["dpp_uploaded"]) / cond if cond > 0 else 0
    test_pct = _cap(a["tests_created"], pay["tests_target"])
    notes_earned = _jr(pay["notes_dpp"] * 0.40 * notes_pct)
    dpp_earned = _jr(pay["notes_dpp"] * 0.40 * dpp_pct)
    tests_earned = _jr(pay["notes_dpp"] * 0.20 * test_pct)

    doubt_pct = a["doubts_resolved"] / a["doubts_assigned"] if a["doubts_assigned"] > 0 else 0
    doubt_earned = _jr(pay["doubt_resolution"] * doubt_pct)

    task_pct = a["tasks_on_time"] / a["tasks_assigned"] if a["tasks_assigned"] > 0 else 0
    content_pct = (_cap(a["tests_created"], pay["tests_target"]) * 0.25 +
                   _cap(a["videos_made"], pay["videos_target"]) * 0.40 +
                   _cap(a["live_sessions"], pay["live_target"]) * 0.24 +
                   _cap(a["shorts_made"], pay["shorts_target"]) * 0.11)
    task_earned = _jr(pay["project_delivery"] * (task_pct * 0.5 + content_pct * 0.5))

    gross = (class_earned + class_quality_earned + notes_earned + dpp_earned +
             tests_earned + doubt_earned + task_earned)
    max_potential = sum(int(pay[k]) for k in EARNINGS_PAY_FIELDS)
    perf = _jr(gross / max_potential * 100) if max_potential else 0
    return {
        "class_earned": class_earned, "class_quality_earned": class_quality_earned,
        "notes_earned": notes_earned, "dpp_earned": dpp_earned, "tests_earned": tests_earned,
        "doubt_earned": doubt_earned, "task_earned": task_earned,
        "gross_earned": gross, "max_potential": max_potential, "perf_score": perf,
        "tds": 0, "other_deduct": 0, "net_payable": gross,
        "pcts": {"class": round(class_pct, 4), "quality": round(quality_pct, 4),
                 "notes": round(notes_pct, 4), "dpp": round(dpp_pct, 4), "tests": round(test_pct, 4),
                 "doubt": round(doubt_pct, 4), "task": round(task_pct, 4),
                 "content": round(content_pct, 4)},
    }


def _month_activity(db, tp, month):
    """Portal activity logs se us month ke salary-input stats."""
    from models import (TimetableEntry, Material, Test, DPP, Doubt, VideoTask,
                        RescheduleRequest, RescheduleStatus)
    y, m = int(month[:4]), int(month[5:7])
    start = date(y, m, 1)
    end = date(y + 1, 1, 1) if m == 12 else date(y, m + 1, 1)
    dt0, dt1 = datetime.combine(start, datetime.min.time()), datetime.combine(end, datetime.min.time())

    entries = db.query(TimetableEntry).filter(
        TimetableEntry.teacher_id == tp.id,
        TimetableEntry.entry_type == "chapter",
        TimetableEntry.status == "approved",
        TimetableEntry.entry_date >= start, TimetableEntry.entry_date < end).all()
    scheduled = len(entries)
    done = [e for e in entries if e.completed]
    conducted = len(done)
    late = sum(1 for e in done if _delay_band(_delay_of(e)) == "late")

    mats = db.query(Material).options(defer(Material.content_b64)).filter(
        Material.teacher_id == tp.id,
        Material.created_at >= dt0, Material.created_at < dt1).all()
    notes = sum(1 for x in mats if (x.material_type or "") == "notes")
    mat_dpp = sum(1 for x in mats if (x.material_type or "") == "dpp")
    mat_test = sum(1 for x in mats if (x.material_type or "") == "test")
    dpps = db.query(DPP).filter(
        DPP.teacher_id == tp.id, DPP.created_at >= dt0, DPP.created_at < dt1).count()
    tests = db.query(Test).filter(
        Test.teacher_id == tp.id, Test.created_at >= dt0, Test.created_at < dt1).count()

    vids = db.query(VideoTask).filter(
        VideoTask.teacher_id == tp.id,
        VideoTask.submitted_at != None,
        VideoTask.submitted_at >= dt0, VideoTask.submitted_at < dt1,
        VideoTask.status != "rejected").all()
    videos = live = shorts = 0
    for t in vids:
        vt = (t.video_type or "").lower()
        if "short" in vt:
            shorts += 1
        elif "live" in vt:
            live += 1
        else:
            videos += 1

    doubts_assigned = db.query(Doubt).filter(
        Doubt.teacher_id == tp.id, Doubt.created_at >= dt0, Doubt.created_at < dt1).count()
    doubts_resolved = db.query(Doubt).filter(
        Doubt.teacher_id == tp.id, Doubt.resolved_at != None,
        Doubt.resolved_at >= dt0, Doubt.resolved_at < dt1).count()

    tasks = db.query(VideoTask).filter(
        VideoTask.teacher_id == tp.id,
        VideoTask.created_at >= dt0, VideoTask.created_at < dt1).all()
    tasks_assigned = len(tasks)
    tasks_on_time = sum(1 for t in tasks if t.on_time)

    resched = db.query(RescheduleRequest).filter(
        RescheduleRequest.teacher_id == tp.id,
        RescheduleRequest.status == RescheduleStatus.approved,
        RescheduleRequest.created_at >= dt0, RescheduleRequest.created_at < dt1).count()

    # --- v83: DPP-per-chapter coverage (1 chapter = 1 DPP mandatory; part-wise OK) ---
    # Chapter "covered" hai agar: class report me dpp_given=True, YA us chapter se
    # match karta chapterwise DPP / material-dpp upload hai (month se pehle ka bhi chalega).
    # Jo completed chapter cover nahi hua -> dpp_pending me teacher ko dikhega.
    from models import DPPType as _DPPType
    def _norm(s):
        return re.sub(r"\s+", " ", (s or "").strip().lower())
    chap_refs = []
    for x in db.query(DPP).filter(DPP.teacher_id == tp.id, DPP.is_active == True).all():
        if x.dpp_type == _DPPType.chapterwise or str(x.dpp_type) in ("chapterwise", "DPPType.chapterwise"):
            if _norm(x.reference):
                chap_refs.append((_norm(x.subject), _norm(x.reference)))
    for x in db.query(Material).options(defer(Material.content_b64)).filter(Material.teacher_id == tp.id,
                                       Material.material_type == "dpp",
                                       Material.chapter != None).all():
        if _norm(x.chapter):
            chap_refs.append((_norm(x.subject), _norm(x.chapter)))

    def _dpp_covered(subject, chapter):
        ns, nc = _norm(subject), _norm(chapter)
        if not nc:
            return False
        for s, r in chap_refs:
            if s and ns and s != ns:
                continue
            if r and (r in nc or nc in r):
                return True
        return False

    dpp_covered = 0
    dpp_pending = []
    for e in done:
        if bool(getattr(e, "dpp_given", False)) or _dpp_covered(e.subject, e.chapter):
            dpp_covered += 1
        else:
            dpp_pending.append({
                "chapter": e.chapter or "", "subject": e.subject or "",
                "class_name": e.class_name or "",
                "date": e.entry_date.isoformat() if e.entry_date else ""})

    return {
        "classes_scheduled": scheduled, "classes_conducted": conducted, "late_classes": late,
        "extra_reschedules": max(0, resched - 1),
        "notes_uploaded": notes, "dpp_uploaded": dpps + mat_dpp,
        "dpp_covered": dpp_covered, "dpp_pending": dpp_pending,
        "tests_created": tests + mat_test,
        "videos_made": videos, "live_sessions": live, "shorts_made": shorts,
        "doubts_assigned": doubts_assigned, "doubts_resolved": doubts_resolved,
        "tasks_assigned": tasks_assigned, "tasks_on_time": tasks_on_time,
    }


def earnings_payload(db, tp, month):
    """Teacher + pay config + month activity + earnings — slip/letter dono ka data."""
    _ensure_v86(db)
    cfg = get_pay_config(db, tp.id)
    act = _month_activity(db, tp, month)
    pay = {k: int(getattr(cfg, k) or 0) for k in EARNINGS_PAY_FIELDS}
    targets = {k: int(getattr(cfg, k) or 0) for k in EARNINGS_TARGET_FIELDS}
    # v95: editable target names + admin ke custom extra targets (letter/display only)
    _DEF_LABELS = {"tests": "Weekly Tests", "videos": "Videos (One Shot/Revision)",
                   "live": "YouTube Live", "shorts": "Shorts"}
    _saved_lab = getattr(cfg, "target_labels", None) or {}
    targets["labels"] = {k: (str(_saved_lab.get(k) or "").strip()[:60] or v)
                         for k, v in _DEF_LABELS.items()}
    _cust = getattr(cfg, "custom_targets", None) or []
    targets["custom"] = [{"name": str(c.get("name") or "").strip()[:60],
                          "count": max(0, int(c.get("count") or 0))}
                         for c in _cust
                         if isinstance(c, dict) and str(c.get("name") or "").strip()][:12]
    e = calc_earnings(act, {**pay, **targets})
    # ---- v86: UNPAID approved leave ka per-day deduction (complete salary se) ----
    # PAID leave (admin ne approve karte waqt 'Paid' chuna) pe koi ktaunti nahi.
    start, end = _month_range(month)
    unpaid_lv = sum(_leave_days_map(db, tp.id, start, end, unpaid_only=True).values())
    total_lv = sum(_leave_days_map(db, tp.id, start, end).values())
    dim = max(1, (end - start).days)
    per_day_sal = e["max_potential"] / dim
    lv_ded = round(per_day_sal * unpaid_lv)
    e["leave_unpaid_days"] = round(unpaid_lv, 1)
    e["leave_paid_days"] = round(total_lv - unpaid_lv, 1)
    e["leave_per_day"] = round(per_day_sal)
    e["leave_deduction"] = lv_ded
    e["net_payable"] = max(0, e["gross_earned"] - lv_ded)
    # ---- v104: target-only teachers (Attendance System = Disabled) ----
    # Inka punch nahi hota, isliye assigned hours-target har mahine FULLY MET maana
    # jaata hai. Payout amounts kabhi display nahi hote — sirf ESTIMATED % of salary:
    # base 100% minus common deduction rules (same as every teacher) ka estimated cut.
    pol = _policy_dict(db, tp.id)
    target_only = bool(pol.get("disabled"))
    if target_only:
        mx = e["max_potential"] or 0
        pcts = e["pcts"]

        def _share(amt):
            return (float(amt) / mx * 100.0) if mx else 0.0

        ded = []

        def _add(pct, label):
            pct = round(pct, 1)
            if pct >= 0.5:
                ded.append({"pct": pct, "label": label})

        sched, cond = act["classes_scheduled"], act["classes_conducted"]
        if sched > 0:
            miss = sched - cond
            if miss > 0:
                _add((1 - pcts["class"]) * _share(pay["class_retainer"]),
                     "%d scheduled class%s not conducted" % (miss, "es" if miss != 1 else ""))
            if act["late_classes"] > 0:
                _add((1 - pcts["quality"]) * _share(pay["class_quality"]),
                     "%d late start%s (15-min grace crossed)"
                     % (act["late_classes"], "s" if act["late_classes"] != 1 else ""))
        if cond > 0:
            if act["notes_uploaded"] < cond:
                _add((1 - pcts["notes"]) * _share(pay["notes_dpp"] * 0.40),
                     "class notes missing on %d conducted class%s"
                     % (cond - act["notes_uploaded"], "es" if cond - act["notes_uploaded"] != 1 else ""))
            pend = len(act.get("dpp_pending") or [])
            if pend > 0:
                _add((1 - pcts["dpp"]) * _share(pay["notes_dpp"] * 0.40),
                     "DPP pending on %d chapter%s" % (pend, "s" if pend != 1 else ""))
        if targets["tests_target"] > 0 and act["tests_created"] < targets["tests_target"]:
            _add((1 - pcts["tests"]) * _share(pay["notes_dpp"] * 0.20),
                 "weekly tests short (%d/%d)" % (act["tests_created"], targets["tests_target"]))
        if act["doubts_assigned"] > act["doubts_resolved"]:
            _add((1 - pcts["doubt"]) * _share(pay["doubt_resolution"]),
                 "%d doubt%s unresolved" % (act["doubts_assigned"] - act["doubts_resolved"],
                                            "s" if act["doubts_assigned"] - act["doubts_resolved"] != 1 else ""))
        if act["tasks_assigned"] > act["tasks_on_time"]:
            _add((1 - pcts["task"]) * _share(pay["project_delivery"] * 0.50),
                 "%d task%s not delivered on time" % (act["tasks_assigned"] - act["tasks_on_time"],
                                                      "s" if act["tasks_assigned"] - act["tasks_on_time"] != 1 else ""))
        if pcts["content"] < 1:
            _add((1 - pcts["content"]) * _share(pay["project_delivery"] * 0.50),
                 "monthly content targets short (videos %d/%d, live %d/%d, shorts %d/%d, tests %d/%d)"
                 % (act["videos_made"], targets["videos_target"],
                    act["live_sessions"], targets["live_target"],
                    act["shorts_made"], targets["shorts_target"],
                    act["tests_created"], targets["tests_target"]))
        if lv_ded > 0:
            _add((float(lv_ded) / mx * 100.0) if mx else 0.0,
                 "%s unpaid-leave day%s" % (e["leave_unpaid_days"],
                                            "s" if e["leave_unpaid_days"] != 1 else ""))
        ded.sort(key=lambda x: -x["pct"])
        e["salary_pct"] = round(e["net_payable"] / mx * 100.0, 1) if mx else 0.0
        e["est_deductions"] = ded[:8]
        e["est_ded_total"] = round(sum(x["pct"] for x in ded), 1)
        hrs = _policy_required(pol)
        tlabel = ("min %gh/day" % hrs) if pol.get("mode") == "flexible" else ("%gh/day" % hrs)
        pol_target = {"label": tlabel, "hours": hrs, "mode": pol.get("mode")}
    else:
        pol_target = None
    name = tp.user.name if tp.user else "Teacher"
    subs = tp.subjects or []
    now = _ist_now()
    try:
        ml = datetime(int(month[:4]), int(month[5:7]), 1).strftime("%B %Y")
    except Exception:
        ml = month
    # ---- v91: digital signature details — letter/slip pe dikhane ke liye ----
    from models import TeacherContract as _TC
    _c = db.query(_TC).filter(_TC.teacher_id == tp.id).first()
    _acc = bool((getattr(tp, "letter_accept_version", 0) or 0) >= LETTER_VERSION)
    _sign = (_c.signature_name or "").strip() if (_c and _acc) else ""
    _sat = _c.accepted_at.strftime("%d %B %Y") if (_c and _acc and _c.accepted_at) else ""
    return {
        "month": month, "month_label": ml,
        "teacher": {
            "id": tp.id, "name": name,
            "employee_code": (cfg.employee_code or "").strip() or ("MVS-T-%03d" % tp.id),
            "designation": (cfg.designation or "").strip() or "Subject Teacher",
            "department": (cfg.department or "").strip() or
                          ("Academic — " + subs[0] if subs else "Academic"),
            "subjects": subs,
            "has_photo": bool(tp.photo_b64),
            "bank": (cfg.bank_name or "").strip(), "account_no": (cfg.account_no or "").strip(),
            "ifsc": (cfg.ifsc or "").strip(),
        },
        "pay": pay, "targets": targets, "activity": act, "earnings": e,
        "target_only": target_only, "policy_target": pol_target,
        "letter": {"ref": "MVS/APT/%d/%03d" % (now.year, tp.id),
                   "date": now.strftime("%d %B %Y"),
                   "configured": cfg.id is not None,
                   "accepted": _acc, "signature_name": _sign, "signed_at": _sat},
    }


@router.get("/earnings")
def teacher_earnings(month: str = "", db: Session = Depends(get_db),
                     current_user=Depends(get_teacher)):
    """Teacher ka apna earnings breakdown (appointment-letter model)."""
    tp = get_teacher_profile(current_user, db)
    month = (month or "").strip() or _ist_now().strftime("%Y-%m")
    if not re.match(r"^\d{4}-\d{2}$", month):
        raise HTTPException(status_code=400, detail="Invalid month (use YYYY-MM).")
    return earnings_payload(db, tp, month)


# Faculty Service Agreement - Table A-0: sirf GROSS salary input hoti hai,
# breakup in fixed % se automatic banta hai (sabhi teachers ke liye same).
SALARY_SPLIT = [("basic", "Basic Pay", 0.50), ("hra", "House Rent Allowance (HRA)", 0.25),
                ("conveyance", "Conveyance / Transport Allowance", 0.05),
                ("medical", "Medical Reimbursement", 0.03125),
                ("lta", "Leave Travel Allowance (LTA)", 0.04375),
                ("special_allowance", "Special Academic Allowance", 0.125)]

_CONTRACT_COLS_READY = False
def _ensure_contract_columns(db):
    """teacher_contracts me salary-breakup columns pehli use me add karta hai."""
    global _CONTRACT_COLS_READY
    if _CONTRACT_COLS_READY:
        return
    from sqlalchemy import text as _text
    cols = ["basic", "hra", "conveyance", "medical", "lta", "special_allowance"]
    for col in cols:
        for ddl in ("ALTER TABLE teacher_contracts ADD COLUMN %s INTEGER NULL" % col,):
            try:
                db.execute(_text(ddl)); db.commit(); break
            except Exception:
                db.rollback()
    _CONTRACT_COLS_READY = True

def _salary_breakup(gross):
    """Gross se agreement-ke-% me breakup. Rounding ke baad total gross ke
    barabar rahe, isliye last component adjust hota hai."""
    gross = int(gross or 0)
    out, used = {}, 0
    for i, (k, _lbl, pct) in enumerate(SALARY_SPLIT):
        if i == len(SALARY_SPLIT) - 1:
            v = gross - used
        else:
            v = round(gross * pct)
            used += v
        out[k] = max(0, v)
    return out

# Annexure A (Penalty & Deduction Schedule) se auto rules - naye contract me pre-fill
DEFAULT_CONTRACT_RULES = """Har class scheduled time par shuru hogi; bina prior intimation ke 15 minute se zyada der ho to class delayed/missed mani jayegi.
Month me sirf 1 approved class re-scheduling allowed hai; uske baad har approved re-schedule par Rs 300/-, aur bina intimation/approval ke re-schedule par Rs 600/- per class deduction lagega.
Class notes, DPP aur lecture report har class ke baad prescribed interval me upload karna compulsory hai; delay par 1st instance Rs 200/-, 2nd Rs 400/-, 3rd aur uske baad har instance par Rs 700/- auto-deduction hoga.
Portal ke student doubts 24 hours me resolve karo; doubt pending >1 din Rs 100/-, >2 din Rs 300/-, >5 din Rs 600/- per doubt auto-deduction hoga.
Shorts, strategy, promotional aur recording tasks deadline tak submit karo; 1st delay par warning, 2nd delay se Rs 100/- per day deduction submission tak lagega.
Har Sunday doubt class + DPP solutions discussion compulsory hai.
Monthly payout portal verification ke baad next month ki first week me process hoga; salary confidential hai aur payout ke liye sirf designated Account Manager se contact karna hai."""

def _contract_out(c, teacher_name=""):
    gross = (c.base_salary or 0) + (c.allowances or 0)
    if c.basic is not None:
        brk = {"basic": c.basic, "hra": c.hra, "conveyance": c.conveyance,
               "medical": c.medical, "lta": c.lta, "special_allowance": c.special_allowance}
    else:
        brk = _salary_breakup(gross)
    brk_lbl = [{"key": k, "label": lbl, "amount": brk[k]} for k, lbl, _p in SALARY_SPLIT]
    return {
        "exists": True, "teacher_name": teacher_name,
        "designation": c.designation or "Subject Teacher",
        "joining_date": str(c.joining_date) if c.joining_date else None,
        "base_salary": c.base_salary or 0, "allowances": c.allowances or 0,
        "gross_salary": gross, "breakup": brk_lbl,
        "working_days": c.working_days or 26,
        "per_day_rate": round((c.base_salary or 0) / (c.working_days or 26)),
        "rules": [r.strip() for r in (c.rules_text or "").splitlines() if r.strip()],
        "accepted": bool(c.accepted),
        "accepted_at": c.accepted_at.strftime("%d %b %Y, %I:%M %p") if c.accepted_at else None,
        "signature_name": c.signature_name
    }

# ===== ATTENDANCE =====
@router.get("/attendance/today")
def attendance_today(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherAttendance
    _ensure_geofence(db)
    now = _ist_now(); today = now.date()
    a = db.query(TeacherAttendance).filter(
        TeacherAttendance.teacher_id == tp.id, TeacherAttendance.att_date == today).first()
    pol = _policy_dict(db, tp.id)
    req = _policy_required(pol)
    # live progress: abhi tak kitne counted hours (punch-out se pehle bhi)
    live_net = None
    if a and a.punch_in:
        end_ref = a.punch_out or now
        gross = (end_ref - a.punch_in).total_seconds() / 3600.0
        brk = (pol.get("break_minutes") or 0) / 60.0 if pol.get("work_type") == "full_time" else 0.0
        live_net = round(max(0.0, gross - brk), 2)
    return {
        "date": str(today), "day": today.strftime("%A"),
        "server_time": now.strftime("%I:%M:%S %p"),
        "punch_in": _fmt_t(a.punch_in if a else None),
        "punch_out": _fmt_t(a.punch_out if a else None),
        "hours": _att_hours(a),
        "net_hours": live_net, "required_hours": req,
        "extra_hours": round(max(0.0, (live_net or 0) - req), 2) if (live_net and pol.get("mode") != "flexible") else 0.0,
        "on_track": bool(live_net is not None and live_net >= req),
        "policy": {**pol, "required": req, "label": _policy_label(pol)}
    }

@router.get("/work-policy")
def teacher_work_policy(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher ko apni admin-set work timing dikhane ke liye (read-only)."""
    tp = get_teacher_profile(current_user, db)
    pol = _policy_dict(db, tp.id)
    return {"policy": {**pol, "required": _policy_required(pol), "label": _policy_label(pol)}}

# ===== v86 MIGRATION (leave paid flag, notif image, resched tracking) =====
_V86_READY = False
def _ensure_v86(db):
    """v86 ke naye columns pehli use me add karta hai (MySQL/SQLite dono safe)."""
    global _V86_READY
    if _V86_READY:
        return
    from sqlalchemy import text as _text
    for stmt in [
        "ALTER TABLE teacher_leaves ADD COLUMN paid BOOLEAN DEFAULT FALSE",
        "ALTER TABLE notifications ADD COLUMN image_url VARCHAR(500) NULL",
        "ALTER TABLE timetable_entries ADD COLUMN resched_by VARCHAR(20) NULL",
        "ALTER TABLE reschedule_requests ADD COLUMN tt_entry_id INTEGER NULL",
    ]:
        try:
            db.execute(_text(stmt)); db.commit()
        except Exception:
            db.rollback()
    _V86_READY = True

# ===== GEOFENCE (punch sirf office ke radius me) =====
_GEOFENCE_READY = False
def _ensure_geofence(db):
    """app_settings table + attendance ke location columns pehli use me bana/add karta hai."""
    global _GEOFENCE_READY
    if _GEOFENCE_READY:
        return
    from sqlalchemy import text as _text
    try:
        db.execute(_text("CREATE TABLE IF NOT EXISTS app_settings (key VARCHAR(50) PRIMARY KEY, value TEXT NULL)"))
        db.commit()
    except Exception:
        db.rollback()
    for col in ["in_lat FLOAT NULL", "in_lng FLOAT NULL", "in_dist INTEGER NULL", "in_office VARCHAR(80) NULL",
                "out_lat FLOAT NULL", "out_lng FLOAT NULL", "out_dist INTEGER NULL", "out_office VARCHAR(80) NULL"]:
        try:
            db.execute(_text("ALTER TABLE teacher_attendance ADD COLUMN %s" % col)); db.commit()
        except Exception:
            db.rollback()
    _GEOFENCE_READY = True

def _office_list(db):
    """Admin ke saare office branches: [{'name','lat','lng','radius'}, ...]. Empty list = geofence off."""
    from models import AppSetting
    import json as _json
    try:
        rows = {r.key: (r.value or "") for r in db.query(AppSetting).filter(
            AppSetting.key.in_(["offices", "office_lat", "office_lng", "office_radius"])).all()}
    except Exception:
        return []
    out = []
    raw = rows.get("offices") or ""
    if raw:
        try:
            data = _json.loads(raw)
            if isinstance(data, list):
                for o in data:
                    try:
                        lat = float(o.get("lat")); lng = float(o.get("lng"))
                        radius = float(o.get("radius") or 30)
                        name = str(o.get("name") or "Office").strip()[:80] or "Office"
                    except Exception:
                        continue
                    if -90 <= lat <= 90 and -180 <= lng <= 180 and radius > 0:
                        out.append({"name": name, "lat": lat, "lng": lng, "radius": radius})
        except Exception:
            out = []
    if not out and rows.get("office_lat"):
        # purana single-office setup -> auto migrate (naam "Main Office")
        try:
            lat = float(rows.get("office_lat")); lng = float(rows.get("office_lng") or "")
            radius = float(rows.get("office_radius") or 30)
            if -90 <= lat <= 90 and -180 <= lng <= 180 and radius > 0:
                out = [{"name": "Main Office", "lat": lat, "lng": lng, "radius": radius}]
        except Exception:
            out = []
    return out

def _office_ips(db):
    """Office ke broadband/WiFi ke public IPs — in se aaye punch GPS ke bina allowed."""
    from models import AppSetting
    import json as _json
    try:
        row = db.query(AppSetting).filter(AppSetting.key == "office_ips").first()
        data = _json.loads(row.value) if row and row.value else []
        # v121: IP count limit hata di (19+ office WiFis). Reader ka [:10] cap bug tha —
        # 10 se aage waale IPs padhe hi nahi jaate the, isliye naya IP save karne par bhi
        # recognize nahi hota tha aur wapas "New IPs" me aa jaata tha. Ab poori list padho.
        return [str(ip).strip() for ip in data if str(ip).strip()] if isinstance(data, list) else []
    except Exception:
        return []

def _log_unknown_ip(db, ip):
    """Geofence ke bahar/GPS-fail punch try karne wala IP (office WiFi list me nahi) — admin ko dikhane ke liye."""
    from models import AppSetting
    import json as _json
    ip = (ip or "").strip()
    if not ip:
        return
    try:
        row = db.query(AppSetting).filter(AppSetting.key == "unknown_ips").first()
        if not row:
            row = AppSetting(key="unknown_ips", value="[]"); db.add(row)
        try:
            data = _json.loads(row.value or "[]")
            if not isinstance(data, list):
                data = []
        except Exception:
            data = []
        data = [x for x in data if isinstance(x, dict) and x.get("ip") != ip]
        data.insert(0, {"ip": ip, "at": _ist_now().strftime("%d %b %Y, %I:%M %p")})
        row.value = _json.dumps(data[:15])
        db.commit()
    except Exception:
        try:
            db.rollback()
        except Exception:
            pass

def _client_ip(request):
    """Railway/proxy ke peeche real client IP (X-Forwarded-For ka pehla)."""
    try:
        xff = request.headers.get("x-forwarded-for", "") if request else ""
        if xff:
            return xff.split(",")[0].strip()
        return request.client.host if request and request.client else ""
    except Exception:
        return ""

def _nearest_office(offices, lat, lng):
    """Sabse nazdeek branch aur uska distance (m)."""
    best, best_d = None, None
    for o in offices:
        d = _haversine_m(lat, lng, o["lat"], o["lng"])
        if best is None or d < best_d:
            best, best_d = o, d
    return best, (int(round(best_d)) if best_d is not None else None)

def _haversine_m(lat1, lng1, lat2, lng2):
    """Do GPS points ke beech ka distance, meters me."""
    from math import radians, sin, cos, asin, sqrt
    R = 6371000.0
    p1, p2 = radians(lat1), radians(lat2)
    dp = radians(lat2 - lat1)
    dl = radians(lng2 - lng1)
    h = sin(dp / 2) ** 2 + cos(p1) * cos(p2) * sin(dl / 2) ** 2
    return 2 * R * asin(sqrt(h))

def _geofence_check(db, lat, lng, accuracy, client_ip=""):
    """Koi bhi branch set hai to nearest se distance validate karta hai. Returns (office|None, dist_m).
    Office ke WiFi/broadband IP se aaya punch GPS ke bina bhi allowed hai (office dict me wifi=True).
    Fail ho to HTTPException raise (403 'outside|<branch>|<dist>|<radius>' = geofence breach)."""
    offices = _office_list(db)
    if not offices:
        return None, None      # geofence off - purana behavior
    ips = _office_ips(db)
    if client_ip and ips and client_ip not in ips:
        _log_unknown_ip(db, client_ip)   # naya WiFi/IP — admin ko one-click add ka option milega
    if client_ip and client_ip in ips:
        # office ka WiFi/broadband — trusted network, GPS optional (PC ke liye)
        acc = None
        try:
            acc = float(accuracy) if accuracy is not None else None
        except Exception:
            acc = None
        if lat is not None and lng is not None and acc is not None and acc <= 80:
            o, d = _nearest_office(offices, float(lat), float(lng))
            return {"name": o["name"], "wifi": True}, d
        return {"name": offices[0]["name"], "wifi": True}, None
    if lat is None or lng is None:
        raise HTTPException(status_code=400, detail="Location is required. Allow location permission in your browser to punch.")
    try:
        acc = float(accuracy or 0)
    except Exception:
        acc = 0
    if acc > 80:
        raise HTTPException(status_code=400, detail="Could not get an accurate location (\u00b1%dm error). PCs/laptops have no GPS - punch from Chrome/Safari on your mobile, in an open area." % int(acc))
    office, dist = _nearest_office(offices, float(lat), float(lng))
    if dist > office["radius"] + min(acc, 20):
        raise HTTPException(status_code=403, detail="outside|%s|%d|%d" % (office["name"], int(dist), int(office["radius"])))
    return office, dist

@router.get("/geofence")
def teacher_geofence_status(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Punch card pe dikhane ke liye: geofence on hai ya nahi + radius."""
    _ensure_geofence(db)
    offices = _office_list(db)
    if not offices:
        return {"active": False, "offices": []}
    return {"active": True, "offices": [{"name": o["name"], "radius": int(o["radius"])} for o in offices]}

@router.post("/attendance/punch-in")
def punch_in(request: Request, payload: dict = Body(default={}), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherAttendance
    if _policy_dict(db, tp.id).get("disabled"):
        raise HTTPException(status_code=403, detail="Punch attendance is disabled for you on this portal — target-only mode")
    _ensure_geofence(db)
    office, dist = _geofence_check(db, payload.get("lat"), payload.get("lng"), payload.get("accuracy"), _client_ip(request))
    now = _ist_now(); today = now.date()
    a = db.query(TeacherAttendance).filter(
        TeacherAttendance.teacher_id == tp.id, TeacherAttendance.att_date == today).first()
    if a and a.punch_in:
        raise HTTPException(status_code=400, detail=f"You already punched in today at {_fmt_t(a.punch_in)}")
    if not a:
        a = TeacherAttendance(teacher_id=tp.id, att_date=today)
        db.add(a)
    a.punch_in = now
    wifi = bool(office and office.get("wifi"))
    if office:
        if payload.get("lat") is not None and payload.get("lng") is not None and not wifi:
            a.in_lat = float(payload.get("lat")); a.in_lng = float(payload.get("lng"))
        a.in_dist = dist if dist is not None else 0
        a.in_office = office["name"]
    db.commit()
    msg = f"Punched in at {_fmt_t(now)}"
    if office:
        msg += " (%s - Office WiFi)" % office["name"] if wifi else " (%s - office se %dm)" % (office["name"], dist)
    return {"message": msg, "punch_in": _fmt_t(now), "distance": dist,
            "office": office["name"] if office else None, "wifi": wifi}

@router.post("/attendance/punch-out")
def punch_out(request: Request, payload: dict = Body(default={}), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherAttendance
    if _policy_dict(db, tp.id).get("disabled"):
        raise HTTPException(status_code=403, detail="Punch attendance is disabled for you on this portal — target-only mode")
    _ensure_geofence(db)
    office, dist = _geofence_check(db, payload.get("lat"), payload.get("lng"), payload.get("accuracy"), _client_ip(request))
    now = _ist_now(); today = now.date()
    a = db.query(TeacherAttendance).filter(
        TeacherAttendance.teacher_id == tp.id, TeacherAttendance.att_date == today).first()
    if not a or not a.punch_in:
        raise HTTPException(status_code=400, detail="Please punch in first")
    if a.punch_out:
        raise HTTPException(status_code=400, detail=f"You already punched out today at {_fmt_t(a.punch_out)}")
    gap_h = (now - a.punch_in).total_seconds() / 3600.0
    if gap_h < MIN_PRESENT_HOURS and not payload.get("confirm"):
        # 1 ghante se pehle punch-out: accidental press se bachne ke liye confirm maango.
        mins = int(gap_h * 60)
        return {"need_confirm": True, "hours": round(gap_h, 2),
                "message": f"Only {mins} min worked — you will be marked Present (Short) today. "
                           f"To confirm, press PUNCH OUT again."}
    a.punch_out = now
    wifi = bool(office and office.get("wifi"))
    if office:
        if payload.get("lat") is not None and payload.get("lng") is not None and not wifi:
            a.out_lat = float(payload.get("lat")); a.out_lng = float(payload.get("lng"))
        a.out_dist = dist if dist is not None else 0
        a.out_office = office["name"]
    db.commit()
    # policy-aware result: net = gap - lunch break (full time); short = required se kam
    pol = _policy_dict(db, tp.id)
    req = _policy_required(pol)
    net = _net_hours(a, pol)
    extra = _extra_hours(a, pol)
    short = (net or 0) < req
    msg = f"Punched out at {_fmt_t(now)}"
    if office:
        msg += " (%s - Office WiFi)" % office["name"] if wifi else " (%s - office se %dm)" % (office["name"], dist)
    if short:
        msg += " — Present (Short): %sh of %sh" % (net, req)
    elif extra > 0:
        msg += " — Extra %sh today" % extra
    return {"message": msg, "punch_out": _fmt_t(now), "hours": _att_hours(a),
            "net_hours": net, "required_hours": req, "extra_hours": extra,
            "short": short, "present": True, "distance": dist,
            "office": office["name"] if office else None, "wifi": wifi}

@router.get("/attendance/history")
def attendance_history(month: str = "", db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherAttendance
    _ensure_geofence(db)
    start, end = _month_range(month)
    rows = db.query(TeacherAttendance).filter(
        TeacherAttendance.teacher_id == tp.id,
        TeacherAttendance.att_date >= start, TeacherAttendance.att_date < end
    ).order_by(TeacherAttendance.att_date.desc()).all()
    lvmap = _leave_days_map(db, tp.id, start, end)

    pol = _policy_dict(db, tp.id)
    req = _policy_required(pol)

    def _st(r):
        return _day_status(r, pol) if (r and (r.punch_in or r.punch_out)) else ""

    out = [{"date": str(r.att_date), "day": r.att_date.strftime("%A"),
            "punch_in": _fmt_t(r.punch_in), "punch_out": _fmt_t(r.punch_out),
            "hours": _att_hours(r), "net_hours": _net_hours(r, pol),
            "extra_hours": _extra_hours(r, pol), "required_hours": req,
            "status": _st(r)} for r in rows]
    total_hours = round(sum(x["net_hours"] or 0 for x in out), 1)
    extra_hours = round(sum(x["extra_hours"] or 0 for x in out), 1)
    full_days = sum(1 for r in rows if _day_status(r, pol) == "present")
    short_days = sum(1 for r in rows if _day_status(r, pol) == "short")
    present_days = full_days + short_days      # short day bhi PRESENT hi count hota hai
    today = _ist_now().date()
    leave_days = round(sum(v for d, v in lvmap.items()), 1)
    leave_elapsed = round(sum(v for d, v in lvmap.items() if d <= today), 1)
    elapsed = _elapsed_days(start, end)
    absent_days = max(0, round(elapsed - present_days - leave_elapsed))
    # leave dates list (calendar me blue dikhane ke liye)
    leave_dates = sorted(str(d) for d in lvmap.keys())
    return {"month": start.strftime("%Y-%m"), "rows": out,
            "present_days": present_days, "full_days": full_days, "short_days": short_days,
            "leave_days": leave_days, "absent_days": absent_days,
            "leave_dates": leave_dates, "elapsed_days": elapsed,
            "min_hours": MIN_PRESENT_HOURS, "required_hours": req,
            "extra_hours": extra_hours,
            "policy": {**pol, "required": req, "label": _policy_label(pol)},
            "total_hours": total_hours}


# ===== LEAVE REQUESTS (apply -> admin approve/reject) =====
@router.post("/leaves/apply")
def leave_apply(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherLeave, User
    try:
        sd = date.fromisoformat((payload.get("start_date") or "")[:10])
        ed = date.fromisoformat((payload.get("end_date") or "")[:10])
    except Exception:
        raise HTTPException(400, "Valid start and end dates are required (YYYY-MM-DD)")
    if ed < sd:
        raise HTTPException(400, "End date cannot be before the start date")
    if (ed - sd).days > 60:
        raise HTTPException(400, "Leave cannot be longer than 60 days")
    ltype = (payload.get("leave_type") or "full").strip().lower()
    if ltype not in ("full", "half"):
        ltype = "full"
    if ltype == "half" and sd != ed:
        raise HTTPException(400, "Half day leave must start and end on the same date")
    reason = (payload.get("reason") or "").strip()
    if len(reason) < 3:
        raise HTTPException(400, "Please write a short reason for the leave")
    overlap = db.query(TeacherLeave).filter(
        TeacherLeave.teacher_id == tp.id,
        TeacherLeave.status.in_(["pending", "approved"]),
        TeacherLeave.start_date <= ed, TeacherLeave.end_date >= sd).first()
    if overlap:
        raise HTTPException(400, "A leave request already exists for overlapping dates "
                                 f"({overlap.start_date} to {overlap.end_date} — {overlap.status})")
    lv = TeacherLeave(teacher_id=tp.id, start_date=sd, end_date=ed,
                      leave_type=ltype, reason=reason, status="pending")
    db.add(lv)
    uname = current_user.name if current_user else "A teacher"
    admins = db.query(User).filter(User.role == "admin", User.is_active == True).all()
    for a in admins:
        db.add(Notification(user_id=a.id, title="🌴 New Leave Request",
                            message=f'{uname} requested leave from {sd.strftime("%d %b")} to {ed.strftime("%d %b")} '
                                    f'({"Half day" if ltype == "half" else "Full day"}): {reason}. '
                                    f'Review it in Attendance > Leave Requests.',
                            notif_type="leave"))
    db.commit()
    return {"ok": True, "id": lv.id}


@router.get("/leaves/my")
def leave_my(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherLeave
    rows = (db.query(TeacherLeave).filter(TeacherLeave.teacher_id == tp.id)
            .order_by(TeacherLeave.created_at.desc()).limit(50).all())
    return {"leaves": [{
        "id": r.id, "start_date": str(r.start_date), "end_date": str(r.end_date),
        "leave_type": r.leave_type, "reason": r.reason or "",
        "status": r.status, "admin_remark": r.admin_remark or "",
        "reviewed_at": r.reviewed_at.strftime("%d %b %Y, %I:%M %p") if r.reviewed_at else "",
        "created_at": r.created_at.strftime("%d %b %Y") if r.created_at else "",
    } for r in rows]}

# ===== CONTRACT (APPOINTMENT LETTER) =====
@router.get("/contract")
def my_contract(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherContract
    c = db.query(TeacherContract).filter(TeacherContract.teacher_id == tp.id).first()
    if not c:
        return {"exists": False}
    return _contract_out(c, current_user.name)

@router.post("/contract/accept")
def accept_contract(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    from models import TeacherContract
    c = db.query(TeacherContract).filter(TeacherContract.teacher_id == tp.id).first()
    if not c:
        raise HTTPException(status_code=404, detail="No appointment letter found for you")
    sig = (payload.get("signature_name") or "").strip()
    if len(sig) < 3:
        raise HTTPException(status_code=400, detail="Please type your full name as your digital signature")
    if not c.accepted:
        c.accepted = True
        c.accepted_at = _ist_now()
        c.signature_name = sig
        db.commit()
    return {"message": "Appointment letter accepted", "accepted_at": c.accepted_at.strftime("%d %b %Y, %I:%M %p")}

# ===== v89: PAYOUT GATE — letter accept (v2) + passcode lock + admin-approved reset =====
LETTER_VERSION = 2  # naya payout system — purane accept (v1/NULL) invalid, dubara sign hoga
_V89_READY = False

def _ensure_v89(db):
    """v89 ke naye columns pehli use me add karta hai (MySQL/SQLite dono safe)."""
    global _V89_READY
    if _V89_READY:
        return
    from sqlalchemy import text as _text
    for stmt in [
        "ALTER TABLE teacher_profiles ADD COLUMN payout_passcode VARCHAR(255) NULL",
        "ALTER TABLE teacher_profiles ADD COLUMN letter_accept_version INTEGER DEFAULT 0",
        "ALTER TABLE teacher_profiles ADD COLUMN passcode_reset_pending BOOLEAN DEFAULT FALSE",
    ]:
        try:
            db.execute(_text(stmt)); db.commit()
        except Exception:
            db.rollback()
    _V89_READY = True

try:
    from security import verify_password as _sec_verify
except Exception:
    _sec_verify = None

def _passcode_ok(plain, hashed):
    if not hashed:
        return False
    if _sec_verify is not None:
        try:
            return bool(_sec_verify(plain, hashed))
        except Exception:
            pass
    from security import hash_password as _hp
    return _hp(plain) == hashed

def _payout_status(tp):
    return {
        "letter_version": LETTER_VERSION,
        "letter_accepted": bool((tp.letter_accept_version or 0) >= LETTER_VERSION),
        "passcode_set": bool(tp.payout_passcode),
        "reset_pending": bool(tp.passcode_reset_pending),
        "remark": tp.letter_remark or "",
        "remark_status": tp.letter_remark_status or "",
        "remark_reply": tp.letter_remark_reply or "",
    }

@router.get("/payout/status")
def payout_gate_status(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Payout section ka lock status — letter accept hua? passcode set? reset pending? remark?"""
    _ensure_v89(db)
    _ensure_v90(db)
    tp = get_teacher_profile(current_user, db)
    return _payout_status(tp)

@router.post("/payout/accept-letter")
def payout_accept_letter(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Payout kholne se pehle appointment letter ka digital sign (current version)."""
    _ensure_v89(db)
    tp = get_teacher_profile(current_user, db)
    sig = (payload.get("signature_name") or "").strip()
    if len(sig) < 3:
        raise HTTPException(status_code=400, detail="Please type your full name as your digital signature")
    tp.letter_accept_version = LETTER_VERSION
    from models import TeacherContract
    c = db.query(TeacherContract).filter(TeacherContract.teacher_id == tp.id).first()
    if not c:
        # v91: contract row na ho to bana do — signature kabhi lose nahi hona chahiye
        c = TeacherContract(teacher_id=tp.id)
        db.add(c)
    c.accepted = True
    c.accepted_at = _ist_now()
    c.signature_name = sig
    db.commit()
    return {"message": "Appointment letter accepted", **_payout_status(tp)}

@router.post("/payout/set-passcode")
def payout_set_passcode(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Letter accept ke baad 4-6 digit passcode create."""
    _ensure_v89(db)
    tp = get_teacher_profile(current_user, db)
    if (tp.letter_accept_version or 0) < LETTER_VERSION:
        raise HTTPException(status_code=400, detail="Please accept the appointment letter first.")
    if tp.passcode_reset_pending:
        raise HTTPException(status_code=400, detail="A reset request is already pending with the admin — please wait for approval.")
    code = (payload.get("passcode") or "").strip()
    if not re.fullmatch(r"\d{4,6}", code or ""):
        raise HTTPException(status_code=400, detail="Passcode must be 4–6 digits.")
    from security import hash_password as _hp
    tp.payout_passcode = _hp(code)
    db.commit()
    return {"message": "Passcode created successfully — payout is now unlocked.", **_payout_status(tp)}

@router.post("/payout/verify-passcode")
def payout_verify_passcode(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Payout unlock — passcode check. Galat pe 403."""
    _ensure_v89(db)
    tp = get_teacher_profile(current_user, db)
    if not tp.payout_passcode:
        raise HTTPException(status_code=400, detail="No passcode has been set yet.")
    code = (payload.get("passcode") or "").strip()
    if not _passcode_ok(code, tp.payout_passcode):
        raise HTTPException(status_code=403, detail="Incorrect passcode — please try again.")
    return {"ok": True, **_payout_status(tp)}

@router.post("/payout/request-passcode-reset")
def payout_request_passcode_reset(db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Reset request ADMIN ke paas jaati hai — approve hone pe hi reset hota hai
    (admin pehle teacher se confirm karega ki request usi ne bheji)."""
    _ensure_v89(db)
    tp = get_teacher_profile(current_user, db)
    if tp.passcode_reset_pending:
        return {"message": "A reset request is already pending with the admin.", **_payout_status(tp)}
    tp.passcode_reset_pending = True
    for adm in db.query(User).filter(User.role == "admin").all():
        db.add(Notification(
            user_id=adm.id,
            title="Passcode Reset Request",
            message=f"{current_user.name} has requested a payout passcode reset. Please confirm with the teacher, then approve or reject it in the Approvals section.",
            notif_type="passcode_reset_request"))
    db.commit()
    return {"message": "Reset request sent to the admin — once approved, you can set a new passcode.", **_payout_status(tp)}

# ===== v90: LETTER REMARKS — accept se pehle doubt bhejo, admin reply kare =====
_V90_READY = False

def _ensure_v90(db):
    """v90 ke remark columns pehli use me add karta hai (MySQL/SQLite dono safe)."""
    global _V90_READY
    if _V90_READY:
        return
    from sqlalchemy import text as _text
    for stmt in [
        "ALTER TABLE teacher_profiles ADD COLUMN letter_remark TEXT NULL",
        "ALTER TABLE teacher_profiles ADD COLUMN letter_remark_status VARCHAR(20) NULL",
        "ALTER TABLE teacher_profiles ADD COLUMN letter_remark_reply TEXT NULL",
        "ALTER TABLE teacher_profiles ADD COLUMN letter_remark_at DATETIME NULL",
    ]:
        try:
            db.execute(_text(stmt)); db.commit()
        except Exception:
            db.rollback()
    _V90_READY = True


@router.post("/payout/letter-remark")
def payout_letter_remark(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Letter padhne ke baad accept karne se PEHLE teacher apna doubt bhejta hai.
    Admin Approvals me check karke reply/resolve karta hai — tab teacher accept kare."""
    _ensure_v90(db)
    tp = get_teacher_profile(current_user, db)
    remark = (payload.get("remark") or "").strip()
    if len(remark) < 5:
        raise HTTPException(status_code=400, detail="Please describe your question in a little more detail (minimum 5 characters).")
    if tp.letter_remark_status == "pending":
        raise HTTPException(status_code=400, detail="Your previous remark is still pending with the admin — please wait for the reply.")
    tp.letter_remark = remark
    tp.letter_remark_status = "pending"
    tp.letter_remark_reply = None
    tp.letter_remark_at = _ist_now()
    for adm in db.query(User).filter(User.role == "admin").all():
        db.add(Notification(
            user_id=adm.id,
            title="Appointment Letter Remark",
            message=f"{current_user.name} has sent a remark on the appointment letter. Please review and reply in the Approvals section.",
            notif_type="letter_remark"))
    db.commit()
    return {"message": "Remark sent to the admin — you will be notified as soon as they reply.", **_payout_status(tp)}

# ===== PAYOUT =====
@router.get("/payout")
def my_payout(month: str = "", db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    tp = get_teacher_profile(current_user, db)
    p = compute_payout(db, tp.id, month)
    if not p:
        return {"exists": False}
    p["exists"] = True
    p["teacher_name"] = current_user.name
    return p

@router.get("/payout-tasks")
def my_payout_tasks(month: str = "", db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher ke manual-category tasks (marked work) us month ke liye."""
    tp = get_teacher_profile(current_user, db)
    perf = compute_performance(db, tp.id, month)
    return {"tasks": perf["tasks"], "categories": [
        {"key": c["key"], "label": c["label"], "source": c["source"], "target": c["target"]}
        for c in perf["categories"] if c["source"] == "manual" and c["target"] > 0]}

@router.post("/payout-task")
def mark_payout_task(payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher off-portal kaam (YouTube video, Short, Live session...) done mark
    karta hai. Admin approve karega tab count hoga. done_date us din ka jab kaam
    HUA - same month me hua to count, next month me hua to 'delayed' (policy)."""
    from models import PayoutTask, PayoutTemplate
    tp = get_teacher_profile(current_user, db)
    mk = (payload.get("month") or "").strip() or _ist_now().strftime("%Y-%m")
    if mk < PAYOUT_PERF_START:
        raise HTTPException(400, "Performance payout %s se shuru hoga" % PAYOUT_PERF_START)
    key = (payload.get("key") or "").strip()
    tpl = db.query(PayoutTemplate).filter(
        PayoutTemplate.teacher_id == tp.id, PayoutTemplate.key == key,
        PayoutTemplate.source == "manual").first()
    if not tpl or (tpl.target or 0) <= 0:
        raise HTTPException(400, "Ye category aapke monthly target me nahi hai")
    title = (payload.get("title") or "").strip()
    if not title:
        raise HTTPException(400, "Kaam ka naam likhna zaroori hai")
    dd = None
    if payload.get("done_date"):
        try:
            dd = datetime.strptime(str(payload["done_date"])[:10], "%Y-%m-%d").date()
        except Exception:
            raise HTTPException(400, "Date format galat hai")
    t = PayoutTask(teacher_id=tp.id, month=mk, key=key, title=title[:200],
                   status="pending", done_date=dd, note=(payload.get("note") or "")[:300])
    db.add(t); db.commit()
    return {"message": "Marked! It will count once the admin approves it.", "id": t.id}

@router.delete("/payout-task/{tid}")
def delete_payout_task(tid: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    from models import PayoutTask
    tp = get_teacher_profile(current_user, db)
    t = db.query(PayoutTask).filter(PayoutTask.id == tid, PayoutTask.teacher_id == tp.id).first()
    if not t:
        raise HTTPException(404, "Task nahi mila")
    if t.status == "approved":
        raise HTTPException(400, "Approved task delete nahi ho sakta - admin se bolo")
    db.delete(t); db.commit()
    return {"message": "Task removed"}

# ===== TEACHER: CHANGE CLASS SLOT (subject ka time — aage ki saari classes) =====
@router.post("/change-slot")
def change_slot(payload: dict, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Teacher apne subject ka slot (time) change karta hai. Aaj se aage ki saari
    incomplete classes naye time pe shift ho jaati hain. Us subject ke students
    aur admins ko notification jaati hai. Completed/purani classes untouched."""
    tp = get_teacher_profile(current_user, db)
    from models import TimetableEntry, StudentProfile, User, UserRole
    subject = (payload.get("subject") or "").strip()
    new_time = (payload.get("new_time") or "").strip()
    if not subject:
        raise HTTPException(status_code=400, detail="Subject is required")
    if not new_time or len(new_time) > 20:
        raise HTTPException(status_code=400, detail="Please enter a valid time, e.g. 9:30 AM")
    if subject not in (tp.subjects or []):
        owned = db.query(TimetableEntry).filter(
            TimetableEntry.subject == subject, TimetableEntry.teacher_id == tp.id).first()
        if not owned:
            raise HTTPException(status_code=403, detail="This subject is not assigned to you")
    today = _ist_now().date()
    _sc = _subj_scope_for(db, TimetableEntry, [subject])
    entries = db.query(TimetableEntry).filter(
        TimetableEntry.subject.in_(list(_sc)),
        TimetableEntry.entry_date.isnot(None),
        TimetableEntry.entry_date >= today,
        TimetableEntry.completed == False
    ).all()
    if not entries:
        raise HTTPException(status_code=400, detail="No upcoming classes found for this subject")
    old_times = sorted({e.time_text for e in entries if e.time_text})
    for e in entries:
        e.time_text = new_time
    db.commit()
    # ---- notifications: subject ke students (fallback: sab students) + admins ----
    eff = today.strftime("%d %b")
    old_str = f" (earlier {', '.join(old_times)})" if old_times else ""
    s_title = "Class Timing Changed"
    s_msg = f"{subject} classes will now start at {new_time} effective {eff}{old_str}. Your timetable has been updated."
    students = db.query(StudentProfile).join(User, StudentProfile.user_id == User.id).filter(User.is_active == True).all()
    matched = [sp for sp in students if subject in (sp.subjects or [])]
    targets = matched if matched else students
    for sp in targets:
        notify(db, sp.user_id, s_title, s_msg, "timetable")
    a_msg = f"{current_user.name} moved {subject} to {new_time}{old_str}. {len(entries)} upcoming classes updated from {eff}."
    for admin in db.query(User).filter(User.role == UserRole.admin, User.is_active == True).all():
        notify(db, admin.id, "Slot Changed - " + subject, a_msg, "timetable")
    db.commit()
    return {"updated": len(entries), "students_notified": len(targets), "new_time": new_time}


# ---------- test editing / status flow (portal v2) ----------
@router.patch("/exam/{exam_id}")
def update_exam(exam_id: int, payload: dict = Body(...), db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Edit a test. Metadata is always updated; questions are fully replaced
    only when a non-empty "questions" list is sent."""
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    _see = _teacher_sees_students(tp, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Test not found")
    for f in ("title", "subject", "chapter", "medium", "test_type"):
        if payload.get(f) is not None:
            setattr(ex, f, payload.get(f))
    if payload.get("class_name") is not None:
        ex.class_name = (payload.get("class_name") or "").strip()
    if payload.get("duration_min") is not None:
        ex.duration_min = _parse_dur(payload.get("duration_min"))
    if "scheduled_at" in payload:
        ex.scheduled_at = _exam_parse_dt(payload.get("scheduled_at"))
    qs = payload.get("questions")
    if isinstance(qs, list) and qs:
        ttype = ex.test_type or "subjective"
        db.query(ExamQuestion).filter(ExamQuestion.exam_id == exam_id).delete()
        total = 0
        for i, q in enumerate(qs, start=1):
            try:
                mm = int(q.get("max_marks", 1) or 1)
            except Exception:
                mm = 1
            total += mm
            co = q.get("correct_option")
            opts_hi = q.get("options_hi") if ttype == "mcq" else None
            db.add(ExamQuestion(exam_id=ex.id, q_no=i,
                   question_text=q.get("question_text", ""),
                   max_marks=mm,
                   model_answer=q.get("model_answer"),
                   options=q.get("options") if ttype == "mcq" else None,
                   correct_option=(str(co) if co not in (None, "") else None),
                   image_b64=_r2img(q.get("image_b64")),
                   question_text_hi=(q.get("question_text_hi") or None),
                   model_answer_hi=(q.get("model_answer_hi") or None),
                   options_hi=(opts_hi if opts_hi else None),
                   model_answer_image=_r2img(q.get("model_answer_image")),
                   alt_image_b64=_r2img(q.get("alt_image_b64")),
                   explanation=(q.get("explanation") or None),
                   explanation_hi=(q.get("explanation_hi") or None)))
        ex.total_marks = total
    db.commit()
    return {"id": ex.id, "title": ex.title, "total_marks": ex.total_marks,
            "scheduled_at": ex.scheduled_at.isoformat() if getattr(ex, "scheduled_at", None) else None}


@router.delete("/exam/{exam_id}")
def delete_exam(exam_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Soft-delete a test (keeps attempts/marks, hides it everywhere)."""
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    _see = _teacher_sees_students(tp, db)
    ex = db.query(Exam).filter(Exam.id == exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(404, "Test not found")
    ex.is_active = False
    db.commit()
    return {"status": "deleted", "id": exam_id}


@router.post("/attempt/{attempt_id}/marking")
def attempt_marking(attempt_id: int, db: Session = Depends(get_db), current_user=Depends(get_teacher)):
    """Flip an attempt from 'checking soon' to 'being checked by teacher'."""
    _ensure_exam_columns(db)
    tp = get_teacher_profile(current_user, db)
    att = db.query(ExamAttempt).filter(ExamAttempt.id == attempt_id).first()
    if not att:
        raise HTTPException(404, "Attempt not found")
    ex = db.query(Exam).filter(Exam.id == att.exam_id, Exam.teacher_id == tp.id).first()
    if not ex:
        raise HTTPException(403, "Not your test")
    if (att.status or "") == "grading":
        att.status = "marking"
        db.commit()
    return {"status": att.status}


# ---------------------------------------------------------------------------
# Teacher ranking - podium board fed by real activity
# ---------------------------------------------------------------------------
# Reports, DPPs, class notes, tests, doubts and attendance all count toward a
# teacher's rank. Scores are normalised against the best teacher in each
# metric so the board stays meaningful as activity grows.

# ---------------------------------------------------------------------------
# TEACHER RANKING v2 — quality/target based (NOT "counting against the best")
# ---------------------------------------------------------------------------
# Owner ki soch (kyun redesign):
#  - DPP: jitne chapters teacher ke syllabus me hain, un sab me 1-1 DPP dena
#    zaroori — score = coverage %. 8 chapter wala 8 DPP = 100%, 18 wala 18 = 100%.
#  - Test: same coverage soch (kitne chapters me test diya).
#  - Class/Lecture: on-time class = full, delayed = aadha. Sirf count nahi.
#  - Attendance: PRESENT DAYS se (hours se nahi — har teacher ke hours alag).
#  - Doubts: koi doubt aaya hi nahi to FULL. Sirf UNRESOLVED doubt se kam hota
#    hai; resolve karte hi wapas full. (Teacher ko lagta tha "doubt aaya nahi
#    to score kyun kam" — ye us ko theek karta hai.)
#  - Tasks/Projects (production team): on-time delivery = up, delayed = down;
#    teacher khud task propose kare aur approve ho to bonus (self-drive).
# Har part apne aap me 0-100 hai (absolute), best ke against normalise NAHI hota.

TEACHER_RANK_WEIGHTS = {
    "dpp": 0.25, "test": 0.15, "lectures": 0.20,
    "attendance": 0.10, "doubts": 0.15, "tasks": 0.15,
}


def _tr_doubt(received, pending):
    """PENALTY-ONLY: sirf pending (unresolved) doubts ghatate hain. Koi pending
    nahi (ya doubt aaya hi nahi) -> None = score me count nahi (display 'All clear')."""
    if pending <= 0 or received <= 0:
        return None
    return round(max(0, received - pending) / received * 100, 1)


def _tr_lectures(ontime, delayed):
    """SIRF scheduled + reported classes: on-time full, delayed aadha. Koi reported
    scheduled class nahi -> None (N/A). Recorded/bina-report class 100% nahi banti."""
    done = ontime + delayed
    if done <= 0:
        return None
    return round((ontime * 1.0 + delayed * 0.5) / done * 100, 1)


def _tr_present(present_days, window_days):
    """Present DAYS ke hisaab se (hours se nahi). ~5 working days/week expected."""
    expected = max(1, round(window_days * 5 / 7))
    return round(min(present_days, expected) / expected * 100, 1)


def _tr_task(ontime, delayed, proposed_ok, assigned):
    """Koi task assign hi nahi -> None (N/A). Warna on-time rate + self-proposed bonus."""
    if assigned <= 0:
        return None
    done = ontime + delayed
    base = (ontime / done * 100) if done > 0 else 0.0
    bonus = min(proposed_ok * 5, 20)
    return round(min(base + bonus, 100), 1)


def _tr_coverage(covered, total):
    """chapters_with_content / total_PE_chapters. 0 diya -> 0% (kabhi free 100 nahi).
    Syllabus mapped nahi (total 0) -> count-based: ~8 chapters full credit."""
    denom = total if total > 0 else 8
    return round(min(covered, denom) / denom * 100, 1)


def _tr_score(parts):
    """None (N/A) categories EXCLUDE + baaki weights renormalise. DPP/test/attendance
    hamesha count (0 bhi) isliye score kabhi undefined nahi. Kuch na karne pe score
    apne aap kam ho jaata hai — koi free 100 nahi."""
    acc = tot = 0.0
    for k, w in TEACHER_RANK_WEIGHTS.items():
        v = parts.get(k)
        if v is None:
            continue
        acc += v * (w * 100)
        tot += (w * 100)
    return round(acc / tot, 1) if tot > 0 else 0.0


def _tr_teacher_pe_and_coverage(db, tp, since):
    """(total_PE_chapters, dpp_covered, test_covered) — teacher ke subjects ke
    verified syllabus PE chapters vs jinme DPP/Test diya. Sab defensive."""
    # teacher ke (subject, class) jode
    pairs = []
    for sc in (getattr(tp, "subject_classes", None) or []):
        try:
            s = (sc.get("subject") or "").strip()
            c = (sc.get("class") or sc.get("class_name") or "").strip()
            if s:
                pairs.append((s, c))
        except Exception:
            continue
    if not pairs:
        for s in (tp.subjects or []):
            if str(s).strip():
                pairs.append((str(s).strip(), ""))
    total_pe = 0
    try:
        import syllabus_routes as _sr
        import syllabus_data as _sd
        _sr._ensure_syllabus(db)
        seen_subj = set()
        for s, c in pairs:
            key = s.strip().lower()
            if key in seen_subj:
                continue
            cands = []
            cl = _sr.class_level_from_name(c) if c else None
            cands = [cl] if cl else ["12", "10"]
            for cc in cands:
                try:
                    code = _sr.subject_code_for_name(db, cc, s)
                    subj = _sr.get_subject(db, cc, code) if code else None
                except Exception:
                    subj = None
                if subj and subj.get("status") == "ready":
                    try:
                        pe = [r for r in _sd.flatten(subj) if r.get("kind") == "PE"]
                        total_pe += len(pe)
                        seen_subj.add(key)
                    except Exception:
                        pass
                    break
    except Exception:
        total_pe = 0
    # distinct chapters jinme DPP diya (DppPack.chapter + purana DPP.reference)
    dpp_ch = set()
    try:
        from models import DppPack
        for pk in db.query(DppPack).filter(DppPack.teacher_id == tp.id,
                                           DppPack.created_at >= since).all():
            ch = (pk.chapter or "").strip().lower()
            if ch:
                dpp_ch.add(((pk.subject or "").strip().lower(), ch))
    except Exception:
        pass
    try:
        for d in db.query(DPP).filter(DPP.teacher_id == tp.id, DPP.is_active == True,
                                      DPP.created_at >= since).all():
            ref = (d.reference or "").strip().lower()
            if ref:
                dpp_ch.add(((d.subject or "").strip().lower(), ref))
    except Exception:
        pass
    # distinct chapters jinme Test/Exam diya
    test_ch = set()
    try:
        for ex in db.query(Exam).filter(Exam.teacher_id == tp.id,
                                        Exam.created_at >= since).all():
            ch = (getattr(ex, "chapter", "") or "").strip().lower()
            # "... \u27E6S:...\u27E7" schedule tag hata do
            ch = re.sub(r"\s*\u27e6s:[^\u27e7]*\u27e7\s*", "", ch).strip()
            if ch:
                test_ch.add(((ex.subject or "").strip().lower(), ch))
    except Exception:
        pass
    return total_pe, len(dpp_ch), len(test_ch)


def _teacher_rank_rows(db, days=90):
    since = datetime.utcnow() - timedelta(days=days)
    from models import TeacherWorkPolicy
    disabled_ids = {r.teacher_id for r in db.query(TeacherWorkPolicy).filter(
        TeacherWorkPolicy.disabled == True).all()}
    try:
        from models import VideoTask
    except Exception:
        VideoTask = None
    rows = []
    for tp in db.query(TeacherProfile).all():
        u = db.query(User).filter(User.id == tp.user_id).first()
        if not u:
            continue

        # ---- DPP / Test coverage (syllabus PE chapters) ----
        total_pe, dpp_cov_n, test_cov_n = _tr_teacher_pe_and_coverage(db, tp, since)
        dpp_cov = _tr_coverage(dpp_cov_n, total_pe)
        test_cov = _tr_coverage(test_cov_n, total_pe)

        # ---- Lectures: SIRF scheduled+reported classes ka on-time (timetable link) ----
        ontime = delayed = 0
        try:
            from models import TimetableEntry
            lecs = db.query(Lecture).options(defer(Lecture.pdf_b64), defer(Lecture.dpp_b64)).filter(Lecture.teacher_id == tp.id,
                                            Lecture.lecture_date >= since.date()).all()
            for l in lecs:
                te = None
                if getattr(l, "timetable_entry_id", None):
                    te = db.query(TimetableEntry).filter(
                        TimetableEntry.id == l.timetable_entry_id).first()
                if not te or not te.entry_date:
                    continue   # bina schedule-link wali class on-time me count NAHI hoti
                cd = l.lecture_date or (te.completed_at.date() if te.completed_at else None)
                if cd and cd > te.entry_date:
                    delayed += 1
                else:
                    ontime += 1
        except Exception:
            pass
        lec_score = _tr_lectures(ontime, delayed)   # None if no scheduled+reported class

        # ---- Attendance: present DAYS ----
        if tp.id in disabled_ids:
            present = len({l.lecture_date for l in db.query(Lecture).options(defer(Lecture.pdf_b64), defer(Lecture.dpp_b64)).filter(
                Lecture.teacher_id == tp.id, Lecture.is_active == True,
                Lecture.lecture_date != None,
                Lecture.lecture_date >= since.date()).all()})
            att_src = "reports"
        else:
            att_rows = db.query(TeacherAttendance).filter(
                TeacherAttendance.teacher_id == tp.id,
                TeacherAttendance.att_date >= since.date()).all()
            present = len({a.att_date for a in att_rows if a.punch_in})
            att_src = "punch"
        att_score = _tr_present(present, days)

        # ---- Doubts: PENALTY-ONLY (sirf pending ghatate hain; koi doubt/pending nahi -> N/A) ----
        received = db.query(Doubt).filter(Doubt.teacher_id == tp.id,
                                          Doubt.created_at >= since).count()
        pending = db.query(Doubt).filter(Doubt.teacher_id == tp.id,
                                         Doubt.status == DoubtStatus.pending,
                                         Doubt.created_at >= since).count()
        doubt_score = _tr_doubt(received, pending)          # None when nothing pending
        doubt_display = 100 if pending <= 0 else (doubt_score if doubt_score is not None else 0)

        # ---- Tasks / projects (production): on-time + proposal bonus; none assigned -> N/A ----
        t_ontime = t_delayed = t_proposed_ok = t_assigned = 0
        if VideoTask is not None:
            try:
                tasks = db.query(VideoTask).filter(VideoTask.teacher_id == tp.id).all()
                for t in tasks:
                    if (t.proposal_ok or "") != "pending":
                        t_assigned += 1
                    if t.submitted_at and t.on_time is True:
                        t_ontime += 1
                    elif t.submitted_at and t.on_time is False:
                        t_delayed += 1
                    if (t.proposal_ok == "approved") and \
                       (str(t.proposed_by or "").lower() not in ("admin", "system", "")):
                        t_proposed_ok += 1
            except Exception:
                pass
        task_score = _tr_task(t_ontime, t_delayed, t_proposed_ok, t_assigned)

        parts = {"dpp": dpp_cov, "test": test_cov, "lectures": lec_score,
                 "attendance": att_score, "doubts": doubt_score, "tasks": task_score}
        rows.append({
            "teacher_id": tp.id, "name": u.name or "", "user_id": u.user_id or "",
            "photo": getattr(u, "photo_b64", None) or "",
            "subjects": tp.subjects or [], "batch": tp.batch or "",
            "attendance_source": att_src,
            "score": _tr_score(parts),
            # display metrics: % (0-100) except present-days (count). None -> UI shows "—".
            "metrics": {
                "dpp_cov": dpp_cov, "test_cov": test_cov,
                "lectures": lec_score, "attendance": present,
                "doubts": doubt_display, "tasks": task_score,
            },
            "detail": {
                "dpp_covered": dpp_cov_n, "test_covered": test_cov_n,
                "chapters_total": total_pe, "coverage_mapped": total_pe > 0,
                "lectures_ontime": ontime, "lectures_delayed": delayed,
                "present_days": present,
                "doubts_received": received, "doubts_pending": pending,
                "tasks_ontime": t_ontime, "tasks_delayed": t_delayed,
                "tasks_assigned": t_assigned, "tasks_proposed_ok": t_proposed_ok,
            },
        })
    rows.sort(key=lambda r: (-r["score"], r["name"].lower()))
    for i, r in enumerate(rows, 1):
        r["rank"] = i
    return rows


@router.get("/rankings")
def teacher_rankings(days: int = 90, db: Session = Depends(get_db),
                     _admin=Depends(get_admin)):
    """Admin board: every teacher ranked by activity (podium UI)."""
    days = max(7, min(int(days or 90), 365))
    from video_tasks import vt_task_rank_rows
    return {"days": days, "results": _teacher_rank_rows(db, days),
            "weights": TEACHER_RANK_WEIGHTS,
            "task_ranking": vt_task_rank_rows(db)}


@router.get("/my-rank")
def teacher_my_rank(days: int = 90, db: Session = Depends(get_db),
                    current_user=Depends(get_teacher)):
    """The logged-in teacher's own position on the board."""
    days = max(7, min(int(days or 90), 365))
    tp = get_teacher_profile(current_user, db)
    rows = _teacher_rank_rows(db, days)
    mine = next((r for r in rows if r["teacher_id"] == tp.id), None)
    # full board bhi bhejo — teacher apne Performance page par podium + list dekhe
    # aur kisi bhi rank par click karke comparison + suggestions paaye
    from video_tasks import vt_task_rank_rows
    return {"days": days, "total": len(rows), "me": mine, "results": rows,
            "weights": TEACHER_RANK_WEIGHTS,
            "task_ranking": vt_task_rank_rows(db),
            "top3": [{k: r[k] for k in ("rank", "name", "score", "photo")} for r in rows[:3]]}
